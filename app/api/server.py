"""Small FastAPI surface for status + manual triggers.

Endpoints:
  GET  /health           — liveness
  GET  /stats            — counts by status
  POST /run/discovery    — kick off discovery
  POST /run/matching     — kick off matching
  POST /run/tailor       — tailor all SHORTLISTED
  POST /run/autofill/{id} — autofill one application
"""
from __future__ import annotations

import logging
from typing import Optional, Optional as _Opt

from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, Request, UploadFile

# Module-level settings so every route can read it. Individual functions also
# import it locally in places; this makes the bare `settings.` references in the
# trust/recruiter/intro endpoints resolve too.
from app.config import settings

log = logging.getLogger(__name__)
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from sqlmodel import select
from sqlalchemy import func, desc

from app.autofill.agent import autofill, preview
from app.db.init_db import get_session
from app.db.models import (
    Application, ApplicationStatus, Job, JobSource,
    UserSubscription, UserUsage, PlanTier, PLAN_LIMITS,
)
from app.discovery.pipeline import run_discovery
from app.matching.pipeline import run_matching
from app.tailoring.tailor import tailor_all_shortlisted
from app.api.demo import PublicDemoRequest

app = FastAPI(title="JobAgent")

# ── Rate limiting ─────────────────────────────────────────────────────────────
# slowapi wraps the same limits as Flask-Limiter; uses the client IP as the key.
# Sensitive / expensive endpoints get tighter per-minute caps.

try:
    from slowapi import Limiter, _rate_limit_exceeded_handler
    from slowapi.util import get_remote_address
    from slowapi.errors import RateLimitExceeded
    _limiter = Limiter(key_func=get_remote_address, default_limits=["200/minute"])
    app.state.limiter = _limiter
    app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)
    _RATE_LIMIT_AVAILABLE = True
except ImportError:
    _RATE_LIMIT_AVAILABLE = False
    log.warning("slowapi not installed — rate limiting disabled. Run: pip install slowapi")

def _rate_limit(limit: str):
    """Decorator factory that applies a slowapi rate limit if available, else no-op."""
    def decorator(fn):
        if _RATE_LIMIT_AVAILABLE:
            return _limiter.limit(limit)(fn)
        return fn
    return decorator


# Serve static files (PWA manifest, icons, service worker)
from fastapi.staticfiles import StaticFiles
import os as _os
_static_dir = _os.path.join(_os.path.dirname(__file__), "..", "static")
if _os.path.isdir(_static_dir):
    app.mount("/static", StaticFiles(directory=_static_dir), name="static")


# ── Supabase session-refresh middleware ─────────────────────────────────────
# When SUPABASE_URL is configured, every response gets a refreshed access
# token in the X-Supabase-Token header so the frontend can keep the session
# alive without the user having to re-login every hour.

from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as StarletteRequest
from starlette.responses import Response as StarletteResponse

class SupabaseSessionMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: StarletteRequest, call_next):
        response = await call_next(request)
        from app.config import settings
        if not settings.use_supabase:
            return response
        # Only attempt refresh if a bearer token is present
        auth_header = request.headers.get("Authorization", "")
        if not auth_header.startswith("Bearer "):
            return response
        token = auth_header.split(" ", 1)[1]
        try:
            from app.db.supabase_client import anon_client
            sb = anon_client()
            result = sb.auth.refresh_session(token)
            if result and result.session:
                response.headers["X-Supabase-Token"] = result.session.access_token
        except Exception:
            pass
        return response

from fastapi.middleware.cors import CORSMiddleware

# CORS Configuration
allowed_origins = [o.strip() for o in settings.cors_allowed_origins.split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(SupabaseSessionMiddleware)


# ── Security audit logging ────────────────────────────────────────────────────
# Logs every 401/403/429 response with IP + path so failed-auth patterns are
# visible in production log streams (Railway / Supabase logs).

import json as _json

_sec_log = logging.getLogger("security")

class SecurityAuditMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: StarletteRequest, call_next):
        response = await call_next(request)
        if response.status_code in (401, 403, 429):
            ip = request.headers.get("x-forwarded-for", request.client.host if request.client else "unknown").split(",")[0].strip()
            event = {
                "event": "auth_failure" if response.status_code in (401, 403) else "rate_limited",
                "status": response.status_code,
                "method": request.method,
                "path": request.url.path,
                "ip": ip,
                "ua": request.headers.get("user-agent", "")[:120],
            }
            _sec_log.warning(_json.dumps(event))
        return response

app.add_middleware(SecurityAuditMiddleware)


# ── Auth helpers ─────────────────────────────────────────────────────────────

def _get_user_id(request: Request) -> str | None:
    """Extract user_id from the Supabase JWT. Returns None if not authenticated.

    Token sources, in order:
    1. Authorization: Bearer header — used by all fetch()/extension API calls.
    2. sb_token cookie — needed for full-page navigations (e.g. GET /dashboard),
       since the browser does NOT attach the Authorization header to those, only
       to fetch(). Without this, server-rendered pages can't identify the user.
    """
    from app.config import settings
    if not settings.use_supabase:
        return "local"   # single-user mode — no auth
    auth = request.headers.get("Authorization", "")
    token = auth.split(" ", 1)[1] if auth.startswith("Bearer ") else None
    if not token:
        token = request.cookies.get("sb_token")
    if not token:
        return None
    from app.db.supabase_client import get_user_id_from_token
    return get_user_id_from_token(token)


def _require_user(request: Request) -> str:
    """Like _get_user_id but raises 401 if not authenticated."""
    uid = _get_user_id(request)
    if not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return uid


def _user_has_resume(uid: str | None) -> bool:
    """True if the user has uploaded (or synthesized) a resume.

    Supabase mode → check the ``resumes/{uid}/`` storage folder.
    Local mode    → check for any ``./data/resume_master.*`` file.
    """
    from app.config import settings
    if settings.use_supabase and uid and uid != "local":
        try:
            from app.db.supabase_client import service_client
            sb = service_client()
            files = sb.storage.from_("resume").list(uid)
            return any((f.get("name") or "").startswith("resume.") for f in (files or []))
        except Exception:
            return False
    import glob
    return bool(glob.glob("./data/resume_master.*"))


def _require_owned_application(request: Request, application_id: int):
    """Load an Application, enforcing that it belongs to the requesting user.

    Raises 401 if unauthenticated, 404 if missing or owned by someone else
    (404 not 403 — don't leak that the ID exists). In SQLite single-user mode
    (uid == "local") ownership is not enforced.
    """
    uid = _require_user(request)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        if uid != "local" and application.user_id != uid:
            raise HTTPException(status_code=404, detail="Application not found")
    return uid


@app.on_event("startup")
async def startup_event():
    import asyncio
    from app.autofill.agent import set_main_loop
    set_main_loop(asyncio.get_running_loop())
    # Create all DB tables at runtime (after env vars are injected by Railway)
    from app.db.init_db import init_db
    init_db()
    # Start background scheduler — runs discovery + matching every 24 hours (once a day)
    asyncio.create_task(_scheduler())


async def _scheduler():
    """Run discovery → matching once a day (every 24 hours) for each user who has a resume."""
    import asyncio
    import logging
    from app.config import settings
    _log = logging.getLogger("scheduler")
    INTERVAL = 24 * 60 * 60
    await asyncio.sleep(120)  # let server fully boot first
    while True:
        try:
            from app.db.models import UserProfile
            with get_session() as session:
                users = session.exec(select(UserProfile)).all()
            user_ids = [u.user_id for u in users if u.user_id and _user_has_resume(u.user_id)]
            if not user_ids:
                _log.info("Scheduler: no users with resumes, skipping run")
            for uid in user_ids:
                try:
                    _log.info("Scheduler: running discovery + matching for user %s", uid)
                    await asyncio.to_thread(_discover_then_match, uid)
                except Exception as e:
                    _log.exception("Scheduler error for user %s: %s", uid, e)
        except Exception as e:
            _log.exception("Scheduler outer error: %s", e)
        await asyncio.sleep(INTERVAL)


try:
    templates = Jinja2Templates(directory="app/templates")
except Exception:
    import os
    templates = Jinja2Templates(directory=os.path.join(os.path.dirname(__file__), "..", "templates"))


def _fromjson_filter(value):
    """Parse a JSON string into a Python object for templates; tolerant of
    None / already-parsed lists / malformed strings (returns [] on failure)."""
    if value is None or value == "":
        return []
    if isinstance(value, (list, dict)):
        return value
    try:
        import json as _json
        return _json.loads(value)
    except Exception:
        return []


def _cleantext_filter(value):
    """Turn scraped HTML / entity-laden text into clean, readable plain text:
    block tags → newlines, <li> → bullets, strip remaining tags, unescape
    entities, normalize whitespace. Job descriptions often arrive as raw HTML."""
    if not value:
        return ""
    import re as _re, html as _html
    text = str(value)
    # Preserve structure: turn list / line-break / block-close tags into newlines
    text = _re.sub(r'(?i)<\s*li[^>]*>', '\n• ', text)
    text = _re.sub(r'(?i)<\s*(br|/p|/div|/li|/ul|/h[1-6]|/tr)\s*/?>', '\n', text)
    # Drop everything else that looks like a tag
    text = _re.sub(r'<[^>]+>', '', text)
    # Decode entities (run twice to catch double-escaped &amp;nbsp;)
    text = _html.unescape(_html.unescape(text)).replace('\xa0', ' ')
    # Tidy each line; common scraper bullets (~, -, *, •) → "• "
    out, blank = [], 0
    for ln in text.splitlines():
        ln = _re.sub(r'[ \t]+', ' ', ln).strip()
        ln = _re.sub(r'^[~\-\*•]\s+', '• ', ln)
        if ln == '':
            blank += 1
            if blank > 1:
                continue
        else:
            blank = 0
        out.append(ln)
    return '\n'.join(out).strip()


templates.env.filters["fromjson"] = _fromjson_filter
templates.env.filters["cleantext"] = _cleantext_filter


def _humanize_signal_filter(value):
    """Turn raw hire-probability signal tokens into short, human-readable
    phrases for the UI. e.g. 'fresh_posting_4d' -> 'Posted 4 days ago',
    'funding_language:series a' -> 'Series A funded',
    'low_velocity_2_openings' -> 'Only 2 openings (selective)'."""
    if not value:
        return ""
    import re as _re
    s = str(value).strip()
    low = s.lower()

    # key:value style → "Series A funded", "Fast-growing team", ...
    if ":" in s:
        key, _, val = s.partition(":")
        key = key.strip().lower()
        val = val.strip()
        val_t = val.title()
        if "funding" in key:
            return f"{val_t} funded"
        if "growth" in key:
            return f"{val_t} team"
        return val_t or key.replace("_", " ").title()

    # fresh_posting_4d → "Posted 4 days ago" (0 → "Posted today", 1 → "yesterday")
    m = _re.match(r"fresh_posting_(\d+)\s*d", low)
    if m:
        n = int(m.group(1))
        if n == 0:
            return "Posted today"
        if n == 1:
            return "Posted yesterday"
        return f"Posted {n} days ago"

    # low/high/med _velocity_N_openings → "6 openings (actively hiring)".
    # The bucket goes in the parenthetical — never a lead word + a number
    # ("A few 6 openings" reads wrong).
    m = _re.match(r"(low|high|med|medium)_velocity_(\d+)_opening", low)
    if m:
        n = int(m.group(2))
        bucket = m.group(1)
        tail = {"low": " (selective)", "high": " (actively hiring)"}.get(bucket, "")
        return f"{n} opening{'s' if n != 1 else ''}{tail}".strip()

    # generic *_N_openings → "N open roles"
    m = _re.match(r".*?(\d+)_opening", low)
    if m:
        n = int(m.group(1))
        return f"{n} open role{'s' if n != 1 else ''}"

    # fallback: snake_case → readable sentence case
    out = _re.sub(r"\s+", " ", s.replace("_", " ").replace(":", " ")).strip()
    return (out[:1].upper() + out[1:]) if out else out


templates.env.filters["humanize_signal"] = _humanize_signal_filter


def _sponsorship_of(job):
    """Jinja global: legal sponsorship assessment for a job (cap-exempt aware)."""
    try:
        from app.intelligence.sponsorship import assess
        return assess(company=getattr(job, "company", "") or "",
                      description=getattr(job, "description", "") or "",
                      url=getattr(job, "url", "") or "")
    except Exception:
        return None


templates.env.globals["sponsorship_of"] = _sponsorship_of


def _urgency_of(job):
    """Jinja global: timing/urgency assessment for a job (fresh / hard-to-fill)."""
    try:
        from app.intelligence.urgency import assess
        return assess(job)
    except Exception:
        return None


templates.env.globals["urgency_of"] = _urgency_of


# ── Public / marketing pages ─────────────────────────────────────────────────

@app.get("/robots.txt")
def serve_robots():
    import os
    from fastapi.responses import FileResponse
    file_path = os.path.join(os.path.dirname(__file__), "..", "static", "robots.txt")
    return FileResponse(file_path)


@app.get("/sitemap.xml")
def serve_sitemap():
    import os
    from fastapi.responses import FileResponse
    file_path = os.path.join(os.path.dirname(__file__), "..", "static", "sitemap.xml")
    return FileResponse(file_path, media_type="application/xml")


@app.get("/api/notifications")
def get_notifications(request: Request) -> dict:
    """Fetch notifications for the current authenticated user."""
    uid = _require_user(request)
    from app.db.models import UserNotification
    with get_session() as session:
        notifs = session.exec(
            select(UserNotification)
            .where(UserNotification.user_id == uid)
            .order_by(desc(UserNotification.created_at))
            .limit(50)
        ).all()
        return {
            "notifications": [
                {
                    "id": n.id,
                    "title": n.title,
                    "message": n.message,
                    "type": n.type,
                    "read": n.read,
                    "link": n.link,
                    "created_at": n.created_at.isoformat() if n.created_at else None,
                }
                for n in notifs
            ]
        }


@app.post("/api/notifications/{notification_id}/read")
def mark_notification_read(request: Request, notification_id: int) -> dict:
    """Mark a specific notification as read."""
    uid = _require_user(request)
    from app.db.models import UserNotification
    with get_session() as session:
        notif = session.get(UserNotification, notification_id)
        if not notif or notif.user_id != uid:
            raise HTTPException(status_code=404, detail="Notification not found")
        notif.read = True
        session.add(notif)
        session.commit()
        return {"ok": True}


@app.post("/api/notifications/read-all")
def mark_all_notifications_read(request: Request) -> dict:
    """Mark all unread notifications for the user as read."""
    uid = _require_user(request)
    from app.db.models import UserNotification
    with get_session() as session:
        unread = session.exec(
            select(UserNotification)
            .where(UserNotification.user_id == uid, UserNotification.read == False)
        ).all()
        for notif in unread:
            notif.read = True
            session.add(notif)
        session.commit()
        return {"ok": True}


@app.post("/api/notifications/clear")
def clear_all_notifications(request: Request) -> dict:
    """Delete all notifications for the current user."""
    uid = _require_user(request)
    from app.db.models import UserNotification
    with get_session() as session:
        notifs = session.exec(
            select(UserNotification).where(UserNotification.user_id == uid)
        ).all()
        for notif in notifs:
            session.delete(notif)
        session.commit()
        return {"ok": True}


@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    return templates.TemplateResponse(request=request, name="landing.html", context={})


@app.get("/pricing", response_class=HTMLResponse)
def pricing_page(request: Request):
    return templates.TemplateResponse(request=request, name="pricing.html", context={})


@app.get("/privacy", response_class=HTMLResponse)
def privacy_page(request: Request):
    return templates.TemplateResponse(request=request, name="privacy.html", context={})


@app.get("/terms", response_class=HTMLResponse)
def terms_page(request: Request):
    return templates.TemplateResponse(request=request, name="terms.html", context={})


@app.post("/api/public/demo-match")
@_rate_limit("10/minute")
async def public_demo_match(
    request: Request,
    target_role: str = Form(...),
    resume_text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    """Public unauthenticated endpoint to run a mock matching cascade and outreach generation."""
    from app.api.demo import run_demo_match, extract_text_from_file, PublicDemoRequest
    try:
        text = resume_text or ""
        if file:
            text = extract_text_from_file(file)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="Please paste resume text or upload a file.")
            
        payload = PublicDemoRequest(resume_text=text, target_role=target_role)
        return run_demo_match(payload)
    except Exception as e:
        log.exception("Public demo matching failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


# ── Auth pages ───────────────────────────────────────────────────────────────

@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    from app.config import settings
    return templates.TemplateResponse(request=request, name="auth.html", context={
        "supabase_url": settings.supabase_url,
        "supabase_anon_key": settings.supabase_anon_key,
    })


@app.get("/auth/callback", response_class=HTMLResponse)
def auth_callback(request: Request):
    from app.config import settings
    return templates.TemplateResponse(request=request, name="auth_callback.html", context={
        "supabase_url": settings.supabase_url,
        "supabase_anon_key": settings.supabase_anon_key,
    })


@app.get("/auth/reset", response_class=HTMLResponse)
def auth_reset_page(request: Request):
    """Password-reset landing page. Supabase sends the recovery link here with
    a token in the URL hash; this page lets the user set a new password."""
    from app.config import settings
    return templates.TemplateResponse(request=request, name="auth_reset.html", context={
        "supabase_url": settings.supabase_url,
        "supabase_anon_key": settings.supabase_anon_key,
    })


@app.post("/auth/logout")
def logout():
    return {"success": True}


# ── Resume upload ─────────────────────────────────────────────────────────────

@app.post("/api/resume/upload")
async def upload_resume(request: Request):
    """Upload resume file. Stores in Supabase Storage (production) or local disk (dev)."""
    from fastapi import UploadFile, File
    uid = _require_user(request)
    form = await request.form()
    file: UploadFile = form.get("file")
    if not file:
        raise HTTPException(status_code=400, detail="No file provided")

    content = await file.read()
    filename = file.filename
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "pdf"
    if ext not in ("pdf", "docx", "md", "txt"):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, MD, TXT allowed")

    # Clear the cached experience/education extraction when a new resume is uploaded
    from app.db.models import AnswerMemory
    from app.db.init_db import get_session
    from sqlmodel import delete as sql_delete
    with get_session() as session:
        session.exec(
            sql_delete(AnswerMemory).where(
                AnswerMemory.label_normalized == "__resume_extracted_experience_education",
                AnswerMemory.user_id == (uid if uid != "local" else None)
            )
        )
        session.commit()

    from app.config import settings
    if settings.use_supabase:
        try:
            from app.db.supabase_client import service_client
            sb = service_client()
            path = f"{uid}/resume.{ext}"
            sb.storage.from_("resume").upload(path, content, {"upsert": "true", "content-type": file.content_type})
            public_url = sb.storage.from_("resume").get_public_url(path)
            return {"success": True, "url": public_url, "path": path}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Storage upload failed: {e}")
    else:
        # Local dev — save to data/
        import os
        os.makedirs("./data", exist_ok=True)
        local_path = f"./data/resume_master.{ext}"
        with open(local_path, "wb") as f:
            f.write(content)
        return {"success": True, "path": local_path}


@app.post("/api/resume/extract-profile")
@_rate_limit("5/minute")
async def extract_profile_from_resume(request: Request, background_tasks: BackgroundTasks) -> dict:
    """Parse the user's uploaded resume and auto-fill their profile fields using Claude."""
    import re as _re
    uid = _require_user(request)

    # Load resume text
    try:
        from app.matching.pipeline import _load_resume
        resume_text = _load_resume(user_id=uid)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not load resume: {exc}")

    if not resume_text or len(resume_text.strip()) < 20:
        raise HTTPException(status_code=400, detail="Resume appears empty")

    # Ask Claude to extract structured info
    import anthropic as _anthropic
    from app.config import settings as _settings
    client = _anthropic.Anthropic(api_key=_settings.anthropic_api_key)

    # Education (university/degree/grad year) usually sits near the END of a resume,
    # so a naive head-only truncation drops it. Send the head plus the tail when the
    # resume is long so the model always sees the Education section.
    _MAX = 16000
    if len(resume_text) <= _MAX:
        resume_for_prompt = resume_text
    else:
        head = resume_text[:11000]
        tail = resume_text[-5000:]
        resume_for_prompt = head + "\n\n...[middle truncated]...\n\n" + tail

    prompt = f"""Extract the following fields from this resume. Return ONLY a JSON object with these exact keys (use null for missing fields):
first_name, last_name, email, phone, location, current_title, years_experience (integer),
linkedin_url, github_url, portfolio_url, degree, university, graduation_year (integer),
key_skills (comma-separated string), professional_summary (2-3 sentence summary of their background),
suggested_target_roles (an array of 4-6 specific job titles this candidate is genuinely
well-qualified for and should apply to right now, ordered best-fit first. Use real,
searchable titles like "Senior Backend Engineer" or "Data Scientist" — base them on the
candidate's actual experience, seniority and skills, not generic guesses).

IMPORTANT: Read the ENTIRE resume, including the Education section (often at the end).
- "university" = the school/college/institution name from Education (e.g. "University of Cincinnati"). Do not leave it null if any school is listed.
- "degree" = the highest/most-recent degree (e.g. "Master of Science in Computer Science").
- "graduation_year" = the most recent graduation year as a 4-digit integer.

Resume:
{resume_for_prompt}

Return only valid JSON, no markdown, no explanation."""

    msg = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=800,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = msg.content[0].text.strip()

    # Strip markdown fences if present
    raw = _re.sub(r"^```(?:json)?\s*", "", raw)
    raw = _re.sub(r"\s*```$", "", raw)

    import json as _json
    try:
        extracted = _json.loads(raw)
    except Exception:
        raise HTTPException(status_code=500, detail="Could not parse extraction response")

    # Save to profile — reuse same single-session logic. Map the sentinel
    # "local" user to None so single-user (SQLite) and SaaS (Supabase) modes
    # read & write the SAME profile row the dashboard later loads.
    from app.db.models import UserProfile
    import datetime as _datetime
    user_id_arg = uid if uid != "local" else None
    with get_session() as session:
        q = select(UserProfile)
        if user_id_arg:
            q = q.where(UserProfile.user_id == user_id_arg)
        else:
            q = q.where(UserProfile.user_id == None)  # noqa: E711 (SQL IS NULL)
        db_profile = session.exec(q).first()
        if not db_profile:
            db_profile = UserProfile(user_id=user_id_arg)
            session.add(db_profile)
            session.flush()

        field_map = [
            "first_name", "last_name", "email", "phone", "location", "current_title",
            "years_experience", "linkedin_url", "github_url", "portfolio_url",
            "degree", "university", "graduation_year", "key_skills", "professional_summary"
        ]
        for field in field_map:
            val = extracted.get(field)
            if val is not None and val != "":
                setattr(db_profile, field, val)
        db_profile.updated_at = _datetime.datetime.utcnow()
        session.add(db_profile)
        session.commit()

    # Auto-seed target_roles from the extracted data — only when the user
    # hasn't already set any roles (so we never overwrite deliberate choices).
    user_id_arg = uid if uid != "local" else None
    seeded_roles: list[str] = []
    with get_session() as session:
        q = select(UserProfile)
        if user_id_arg:
            q = q.where(UserProfile.user_id == user_id_arg)
        p = session.exec(q).first()
        if p and not (p.target_roles or "").strip():
            seen_r: set[str] = set()
            def _add_role(r: str):
                r = (r or "").strip()
                if r and r.lower() not in seen_r and len(seeded_roles) < 6:
                    seen_r.add(r.lower())
                    seeded_roles.append(r)
            # Prefer the roles Claude chose for this candidate.
            ai_roles = extracted.get("suggested_target_roles") or []
            if isinstance(ai_roles, str):
                ai_roles = [r for r in ai_roles.split(",")]
            if isinstance(ai_roles, list):
                for r in ai_roles:
                    _add_role(r if isinstance(r, str) else str(r))
            # Fallback heuristic if Claude returned nothing usable: current_title
            # first, then skill tokens that read like roles.
            if not seeded_roles:
                _add_role(extracted.get("current_title") or "")
                _role_tokens = ("engineer", "scientist", "developer", "manager",
                                "analyst", "designer", "architect", "lead")
                for s in (extracted.get("key_skills") or "").split(","):
                    s = s.strip()
                    if s and any(tok in s.lower() for tok in _role_tokens):
                        _add_role(s)
            if seeded_roles:
                p.target_roles = ", ".join(seeded_roles)
                p.updated_at = _datetime.datetime.utcnow()
                session.add(p)
                session.commit()

        # Trigger background memory harvesting if GitHub/LinkedIn urls are present
        if db_profile.github_url or db_profile.linkedin_url:
            from app.intelligence.harvester import run_harvest
            background_tasks.add_task(
                run_harvest,
                user_id=(uid if uid != "local" else None),
                notify=False
            )

        # Trigger background LinkedIn/Resume alignment review
        from app.intelligence.harvester import check_cross_profile_alignment
        background_tasks.add_task(
            check_cross_profile_alignment,
            user_id=(uid if uid != "local" else None)
        )

        # Run resume analysis and create notifications for "Needs work" suggestions
        try:
            analysis = analyze_resume_text(resume_text, uid)
            from app.db.models import UserNotification
            with get_session() as session:
                for f in analysis.get("findings", []):
                    if not f.get("ok"):
                        existing_notif = session.exec(
                            select(UserNotification)
                            .where(
                                UserNotification.user_id == uid,
                                UserNotification.title == f"Résumé Suggestion: {f['label']} 📄",
                                UserNotification.read == False
                            )
                        ).first()
                        if not existing_notif:
                            notif = UserNotification(
                                user_id=uid,
                                title=f"Résumé Suggestion: {f['label']} 📄",
                                message=f.get("detail") or "Consider improving this section of your resume.",
                                type="resume_suggestions",
                                link="/dashboard",
                            )
                            session.add(notif)
                session.commit()
        except Exception as ae:
            log.warning("Failed to run resume suggestions: %s", ae)

    return {"success": True, "extracted": {k: extracted.get(k) for k in field_map},
            "seeded_roles": seeded_roles}


@app.get("/api/resume/status")
def resume_status(request: Request) -> dict:
    """Whether the current user has a resume on file. Drives the Discover gate."""
    uid = _get_user_id(request)
    return {"has_resume": _user_has_resume(uid)}


def analyze_resume_text(text: str, uid: str) -> dict:
    import re as _re
    low = text.lower()
    words = _re.findall(r"\b\w+\b", text)
    wc = len(words)
    findings: list[dict] = []

    def _add(label, ok, detail):
        findings.append({"label": label, "ok": bool(ok), "detail": detail})

    # 1. Contact info
    has_email = bool(_re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text))
    has_phone = bool(_re.search(r"(\+?\d[\d\s().-]{7,}\d)", text))
    _add("Contact details", has_email and has_phone,
         "Email & phone found" if has_email and has_phone else "Add a clear email and phone number")

    # 2. Standard sections
    sections = {
        "experience": any(s in low for s in ("experience", "employment", "work history")),
        "education": "education" in low,
        "skills": any(s in low for s in ("skills", "technologies", "technical")),
    }
    missing_sec = [k for k, v in sections.items() if not v]
    _add("Standard sections", not missing_sec,
         "Experience, education & skills present" if not missing_sec
         else f"Missing section(s): {', '.join(missing_sec)}")

    # 3. Quantified achievements
    metric_hits = len(_re.findall(r"\b\d+%|\$\s?\d|\b\d{2,}\b", text))
    _add("Quantified impact", metric_hits >= 3,
         f"{metric_hits} measurable results" if metric_hits >= 3
         else "Add metrics (%, $, counts) to your bullets")

    # 4. Action verbs
    action_verbs = ("built", "led", "designed", "developed", "launched", "created",
                    "improved", "reduced", "increased", "managed", "shipped", "drove",
                    "owned", "architected", "delivered", "implemented", "optimized")
    verb_hits = sum(low.count(v) for v in action_verbs)
    _add("Action verbs", verb_hits >= 5,
         f"{verb_hits} strong action verbs" if verb_hits >= 5
         else "Start bullets with action verbs (Built, Led, Shipped…)")

    # 5. Length
    good_len = 350 <= wc <= 1100
    _add("Length", good_len,
         f"{wc} words — good" if good_len
         else (f"{wc} words — too short, add detail" if wc < 350 else f"{wc} words — consider trimming"))

    # 6. Target-role alignment
    roles = _get_target_roles(uid)
    if not roles:
        _add("Target-role match", False, "Add target roles to check résumé alignment")
    else:
        role_terms = set()
        for r in roles:
            for tok in _re.findall(r"\b\w+\b", r.lower()):
                if len(tok) > 2:
                    role_terms.add(tok)
        covered = [t for t in role_terms if t in low]
        role_ok = bool(role_terms) and (len(covered) / len(role_terms) >= 0.5)
        _add("Target-role match", role_ok,
             f"Résumé reflects your target roles ({len(covered)}/{len(role_terms)} terms)" if role_ok
             else "Résumé weakly matches your target roles — weave in relevant keywords")

    score = round(sum(1 for f in findings if f["ok"]) / len(findings) * 100)
    grade = "Strong" if score >= 80 else "Good" if score >= 60 else "Needs work"
    return {"has_resume": True, "score": score, "grade": grade,
            "word_count": wc, "findings": findings}


@app.get("/api/resume/analysis")
def resume_analysis(request: Request) -> dict:
    """General ATS-readiness analysis of the user's résumé — no specific job needed.

    Scores deterministic parse-ability signals (contact info, sections, quantified
    impact, action verbs, length) plus how well the résumé reflects the user's
    saved target roles. Fully local: no LLM / network calls, so it's instant + free.
    """
    uid = _get_user_id(request)
    from app.config import settings
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if not _user_has_resume(uid):
        return {"has_resume": False}
    try:
        from app.matching.pipeline import _load_resume
        text = _load_resume(user_id=uid)
    except Exception as e:
        return {"has_resume": True, "error": f"Could not read résumé: {e}"}
    if not text or len(text.strip()) < 30:
        return {"has_resume": True, "error": "Résumé appears empty or unreadable."}

    return analyze_resume_text(text, uid)


@app.get("/api/discovery/last-run")
def discovery_last_run(request: Request) -> dict:
    """Return the most recent discovery run's per-source summary for this user."""
    import json
    from app.db.models import DiscoveryRun
    uid = _get_user_id(request)
    with get_session() as session:
        q = select(DiscoveryRun).order_by(desc(DiscoveryRun.id))
        if uid and uid != "local":
            q = q.where(DiscoveryRun.user_id == uid)
        run = session.exec(q).first()
        if not run:
            return {"run": None, "can_run": True, "gate_detail": ""}
        try:
            counts = json.loads(run.source_counts or "{}")
        except Exception:
            counts = {}
        allowed, detail = _discovery_gate(uid)
        return {"can_run": allowed, "gate_detail": detail, "run": {
            "id": run.id,
            "status": run.status,
            "started_at": run.started_at.isoformat() if run.started_at else None,
            "finished_at": run.finished_at.isoformat() if run.finished_at else None,
            "total_fetched": run.total_fetched,
            "total_inserted": run.total_inserted,
            "total_shortlisted": run.total_shortlisted,
            "error": run.error,
            "sources": counts,
        }}


@app.get("/api/resume/view")
def view_resume(request: Request) -> dict:
    """Return a temporary signed URL to view/download the user's uploaded resume."""
    from app.config import settings
    uid = _require_user(request)
    if settings.use_supabase and uid and uid != "local":
        try:
            from app.db.supabase_client import service_client
            sb = service_client()
            files = sb.storage.from_("resume").list(uid)
            names = [f.get("name", "") for f in (files or []) if f.get("name", "").startswith("resume.")]
            if names:
                signed = sb.storage.from_("resume").create_signed_url(f"{uid}/{names[0]}", 3600)
                url = (signed or {}).get("signedURL") or (signed or {}).get("signedUrl")
                return {"url": url, "filename": names[0]}
        except Exception as e:
            log.warning("Could not sign resume URL: %s", e)
        return {"url": None}
    # Local mode — serve the file directly
    import glob
    matches = glob.glob("./data/resume_master.*")
    if matches:
        return {"url": "/api/resume/file", "filename": matches[0].split("/")[-1]}
    return {"url": None}


@app.get("/api/resume/file")
def resume_file(request: Request):
    """Serve the locally-stored resume (local/dev mode only)."""
    from fastapi.responses import FileResponse
    import glob
    matches = glob.glob("./data/resume_master.*")
    if not matches:
        raise HTTPException(status_code=404, detail="No resume on file")
    return FileResponse(matches[0])


@app.post("/api/resume/synthesize")
def synthesize_resume(request: Request) -> dict:
    """Build a minimal markdown resume from the user's profile fields.

    Used when the user fills their details in manually instead of uploading a
    file — gives the matcher something to work with so discovery isn't blocked.
    """
    uid = _require_user(request)
    from app.autofill.answer_pack import _get_or_create_profile
    profile = _get_or_create_profile(user_id=uid if uid != "local" else None)
    name = f"{profile.first_name} {profile.last_name}".strip() or "Candidate"
    md = (
        f"# {name}\n\n"
        f"**Email:** {profile.email}\n"
        f"**Location:** {profile.location}\n"
        f"**Current Title:** {profile.current_title}\n\n"
        f"## Summary\n{profile.professional_summary or ''}\n\n"
        f"## Skills\n{profile.key_skills or ''}\n\n"
        f"## Experience\n{profile.current_title or ''}\n"
    )
    from app.config import settings
    if settings.use_supabase and uid != "local":
        try:
            from app.db.supabase_client import service_client
            sb = service_client()
            sb.storage.from_("resume").upload(
                f"{uid}/resume.md",
                md.encode("utf-8"),
                {"upsert": "true", "content-type": "text/markdown"},
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Could not save resume: {e}")
    else:
        import os
        os.makedirs("./data", exist_ok=True)
        with open("./data/resume_master.md", "w", encoding="utf-8") as f:
            f.write(md)
    return {"success": True}


@app.post("/api/jobs/clear")
def clear_jobs(request: Request) -> dict:
    """Delete all jobs + applications belonging to the current user.

    Used to reset a pool that was populated before the resume gate existed.
    In SQLite single-user mode (uid == "local") this clears the whole DB.
    """
    uid = _require_user(request)
    _scoped = uid and uid != "local"
    deleted_apps = 0
    deleted_jobs = 0
    with get_session() as session:
        aq = select(Application)
        jq = select(Job)
        if _scoped:
            aq = aq.where(Application.user_id == uid)
            jq = jq.where(Job.user_id == uid)
        for a in session.exec(aq).all():
            session.delete(a)
            deleted_apps += 1
        for j in session.exec(jq).all():
            session.delete(j)
            deleted_jobs += 1
        session.commit()
    return {"success": True, "deleted_jobs": deleted_jobs, "deleted_applications": deleted_apps}


@app.get("/health")
def health() -> dict:
    return {"ok": True}


@app.get("/api/debug/tenancy")
def debug_tenancy(request: Request) -> dict:
    """Quick check that multi-tenant mode is live and the backend sees the caller
    as a real user (not the shared 'local' tenant). Auth required.

    Healthy production response looks like:
        { "use_supabase": true, "your_uid": "<uuid>", "is_local": false, ... }
    If use_supabase is false or your_uid is "local", DATABASE_URL/SUPABASE_URL
    aren't set on the backend and all users share one tenant.
    """
    from app.config import settings
    auth = request.headers.get("Authorization", "")
    token = auth.split(" ", 1)[1] if auth.startswith("Bearer ") else None
    uid = _get_user_id(request)
    # If a token was sent but didn't resolve to a user, surface WHY — this is the
    # usual "authenticated:false while logged in" cause (bad/missing service key,
    # or a token from a different Supabase project than the backend keys).
    verify_error = None
    if token and uid is None and settings.use_supabase:
        try:
            from app.db.supabase_client import service_client
            service_client().auth.get_user(token)
        except Exception as e:
            verify_error = f"{type(e).__name__}: {str(e)[:180]}"
    return {
        "use_supabase": settings.use_supabase,
        "database_url_set": bool(settings.database_url),
        "supabase_url_set": bool(settings.supabase_url),
        "service_role_key_set": bool(settings.supabase_service_role_key),
        "anon_key_set": bool(settings.supabase_anon_key),
        "token_present": bool(token),
        "authenticated": uid is not None,
        "your_uid": uid,
        "is_local": uid == "local",
        "verify_error": verify_error,
    }


@app.get("/stats")
def stats(request: Request) -> dict:
    uid = _get_user_id(request)
    with get_session() as session:
        q = select(Job).where(Job.is_closed == False)
        if uid and uid != "local":
            q = q.where(Job.user_id == uid)
        total_jobs = len(session.exec(q).all())
        by_status = {}
        for st in ApplicationStatus:
            # Exclude orphan applications (Job row deleted) by joining Job.
            aq = (
                select(func.count(Application.id))
                .join(Job, Application.job_id == Job.id)
                .where(Application.status == st)
            )
            if uid and uid != "local":
                aq = aq.where(Application.user_id == uid)
            count = session.exec(aq).first() or 0
            if count:
                by_status[st.value] = count
    return {"total_jobs": total_jobs, "applications": by_status}


@app.get("/shortlist")
def shortlist(request: Request):
    """Top-scored jobs not yet processed — scoped to the authenticated user."""
    uid = _require_user(request)
    with get_session() as session:
        q = select(Job).where(Job.rerank_score >= 70)
        if uid != "local":
            q = q.where(Job.user_id == uid)
        jobs = session.exec(q.order_by(Job.rerank_score.desc())).all()
    return [
        {
            "id": j.id,
            "company": j.company,
            "title": j.title,
            "location": j.location,
            "url": j.url,
            "similarity": j.similarity_score,
            "rerank": j.rerank_score,
            "hire_probability": j.hire_probability_score,
            "blended": j.blended_score,
            "reason": j.rerank_reasoning,
        }
        for j in jobs
    ]


@app.get("/api/stats")
def api_stats(request: Request) -> dict:
    uid = _get_user_id(request)
    _uid_filter = (uid and uid != "local")
    with get_session() as session:
        # Total jobs in db
        jq = select(func.count(Job.id))
        if _uid_filter:
            jq = jq.where(Job.user_id == uid)
        total_jobs = session.exec(jq).first() or 0

        # Unique companies in Job db
        cq = select(func.count(func.distinct(Job.company)))
        if _uid_filter:
            cq = cq.where(Job.user_id == uid)
        total_companies = session.exec(cq).first() or 0

        # Funnel metrics
        def _jcount(extra=None):
            q = select(func.count(Job.id))
            if _uid_filter:
                q = q.where(Job.user_id == uid)
            if extra is not None:
                q = q.where(extra)
            return session.exec(q).first() or 0

        cross_encoder_passed = _jcount(Job.similarity_score.is_not(None))
        reranker_scored = _jcount(Job.rerank_score.is_not(None))

        # Application counts by status — JOIN Job so orphan applications (whose
        # Job row was deleted) are excluded; otherwise counts here disagree with
        # the dashboard kanban, which inner-joins Job and never shows orphans.
        app_counts = {}
        for status in ApplicationStatus:
            aq = (
                select(func.count(Application.id))
                .join(Job, Application.job_id == Job.id)
                .where(Application.status == status)
                .where(
                    Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
                )
            )
            if _uid_filter:
                aq = aq.where(Application.user_id == uid)
            count = session.exec(aq).first() or 0
            app_counts[status.value] = count

        shortlisted = app_counts[ApplicationStatus.SHORTLISTED.value] + app_counts[ApplicationStatus.TAILORED.value]

        # Score distribution
        band_85_100 = _jcount(Job.rerank_score >= 85)
        band_60_84 = _jcount((Job.rerank_score >= 60) & (Job.rerank_score < 85))
        band_40_59 = _jcount((Job.rerank_score >= 40) & (Job.rerank_score < 60))
        band_0_39 = _jcount((Job.rerank_score >= 0) & (Job.rerank_score < 40))
        unranked = _jcount(Job.rerank_score.is_(None))

        # Top companies
        top_q = select(Job.company, func.count(Job.id)).group_by(Job.company).order_by(desc(func.count(Job.id))).limit(10)
        if _uid_filter:
            top_q = top_q.where(Job.user_id == uid)
        top_companies_res = session.exec(top_q).all()
        top_companies = [{"company": company, "count": count} for company, count in top_companies_res]

        # Per-source job counts (for source breakdown bar in UI)
        from app.db.models import JobSource as JS
        source_counts: dict[str, int] = {}
        for src in JS:
            sq = select(func.count(Job.id)).where(Job.source == src)
            if _uid_filter:
                sq = sq.where(Job.user_id == uid)
            cnt = session.exec(sq).first() or 0
            if cnt:
                source_counts[src.value] = cnt

        # Company registry stats (global — not per-user)
        from app.db.models import CompanyRegistry
        total_boards = session.exec(select(func.count(CompanyRegistry.id))).first() or 0
        active_boards = session.exec(select(func.count(CompanyRegistry.id)).where(CompanyRegistry.is_active == True)).first() or 0
        total_validated_jobs = session.exec(select(func.sum(CompanyRegistry.job_count))).first() or 0

    return {
        "total_jobs": total_jobs,
        "total_companies": total_companies,
        "funnel": {
            "total_pool": total_jobs,
            "cross_encoder_passed": cross_encoder_passed,
            "reranker_scored": reranker_scored,
            "shortlisted": shortlisted,
        },
        "applications": app_counts,
        "scores": {
            "band_85_100": band_85_100,
            "band_60_84": band_60_84,
            "band_40_59": band_40_59,
            "band_0_39": band_0_39,
            "unranked": unranked,
        },
        "sources": source_counts,
        "top_companies": top_companies,
        "registry": {
            "total_boards": total_boards,
            "active_boards": active_boards,
            "total_validated_jobs": total_validated_jobs,
        }
    }


_AUTOFILL_SOURCES = {"greenhouse", "lever", "ashby", "workday", "smartrecruiters"}

@app.get("/api/jobs")
def api_jobs(
    request: Request,
    page: int = 1,
    limit: int = 50,
    search: str = None,
    company: str = None,
    status: str = None,
    min_score: int = None,
    max_score: int = None,
    remote: str = None,
    track: str = None,   # "autofill" | "manual"
    hide_aggregators: str = None,  # "1" = exclude aggregator-redirect jobs
) -> dict:
    uid = _get_user_id(request)
    _uid_filter = uid and uid != "local"
    offset = (page - 1) * limit

    with get_session() as session:
        # Base query — exclude closed/purged jobs from the Jobs table.
        query = select(Job, Application).outerjoin(Application, Application.job_id == Job.id).where(Job.is_closed == False)
        if _uid_filter:
            query = query.where(Job.user_id == uid)

        # Apply filters
        if search:
            search_pattern = f"%{search}%"
            query = query.where(Job.title.like(search_pattern) | Job.company.like(search_pattern) | Job.location.like(search_pattern))

        if company:
            query = query.where(Job.company == company)

        if track == "autofill":
            query = query.where(Job.source.in_(list(_AUTOFILL_SOURCES)))
        elif track == "manual":
            query = query.where(Job.source.not_in(list(_AUTOFILL_SOURCES)))
            
        if status:
            if status == "unprocessed":
                query = query.where(Application.id.is_(None))
            else:
                query = query.where(Application.status == status)
                
        if min_score is not None:
            query = query.where(Job.rerank_score >= min_score)
            
        if max_score is not None:
            query = query.where(Job.rerank_score <= max_score)
            
        if remote is not None:
            is_remote = remote.lower() == "true"
            query = query.where(Job.remote == is_remote)

        if hide_aggregators == "1":
            query = query.where(
                Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
            )

        # Get total count (for pagination) — also exclude closed jobs.
        count_query = select(func.count(Job.id)).where(Job.is_closed == False)
        if _uid_filter:
            count_query = count_query.where(Job.user_id == uid)
        if search:
            search_pattern = f"%{search}%"
            count_query = count_query.where(Job.title.like(search_pattern) | Job.company.like(search_pattern) | Job.location.like(search_pattern))
        if company:
            count_query = count_query.where(Job.company == company)
        if status:
            if status == "unprocessed":
                count_query = count_query.select_from(Job).outerjoin(Application, Application.job_id == Job.id).where(Application.id.is_(None))
            else:
                count_query = count_query.select_from(Job).join(Application, Application.job_id == Job.id).where(Application.status == status)
        if min_score is not None:
            count_query = count_query.where(Job.rerank_score >= min_score)
        if max_score is not None:
            count_query = count_query.where(Job.rerank_score <= max_score)
        if remote is not None:
            is_remote = remote.lower() == "true"
            count_query = count_query.where(Job.remote == is_remote)

        if hide_aggregators == "1":
            count_query = count_query.where(
                Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
            )

        total = session.exec(count_query).first() or 0
        
        # Apply pagination and sorting
        query = query.order_by(desc(Job.blended_score), desc(Job.rerank_score), desc(Job.similarity_score), desc(Job.id)).offset(offset).limit(limit)
        
        results = session.exec(query).all()
        
        jobs_list = []
        for job, app in results:
            jobs_list.append({
                "id": job.id,
                "source": job.source.value if job.source else "manual",
                "company": job.company,
                "title": job.title,
                "location": job.location,
                "remote": job.remote,
                "url": job.url,
                "similarity": job.similarity_score,
                "rerank": job.rerank_score,
                "hire_probability": job.hire_probability_score,
                "blended": job.blended_score,
                "reason": job.rerank_reasoning,
                "application": {
                    "id": app.id,
                    "status": app.status.value,
                    "apply_track": app.apply_track,
                    "created_at": app.created_at.isoformat() if app.created_at else None,
                    "updated_at": app.updated_at.isoformat() if app.updated_at else None,
                } if app else None
            })
            
        import math
        pages = math.ceil(total / limit) if total else 0
        
        return {
            "jobs": jobs_list,
            "total": total,
            "page": page,
            "pages": pages,
            "limit": limit
        }


@app.get("/dashboard", response_class=HTMLResponse)
def dashboard(request: Request, all_submitted: bool = False):
    """Kanban board UI for tracking application progress."""
    from app.config import settings
    uid = _get_user_id(request)
    _uid_filter = uid and uid != "local"
    # SSR auth: whether THIS page navigation was authenticated (via sb_token
    # cookie). If not, fail closed and render no pipeline data — never leak other
    # tenants' applications. The client auth-guard sets the cookie and reloads.
    ssr_authed = bool(uid) or not settings.use_supabase
    
    shortlisted = []
    submitted = []
    interviewing = []
    rejected = []
    skipped = []
    bot_filled = []
    manual_queue = []
    total_submitted_count = 0

    if not (settings.use_supabase and not uid):
        _AUTOFILL_REVIEW_STATUSES = [
            ApplicationStatus.AUTOFILLED,
            ApplicationStatus.AWAITING_USER,
            ApplicationStatus.READY_TO_SUBMIT,
        ]
        
        with get_session() as session:
            # 1. Fetch Shortlisted (limit to 100 to keep cache and parsing fast)
            q_short = select(Application, Job).join(Job).where(
                Application.status.in_([ApplicationStatus.SHORTLISTED, ApplicationStatus.TAILORED] + _AUTOFILL_REVIEW_STATUSES)
            ).where(
                Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
            ).order_by(Application.updated_at.desc()).limit(100)
            if _uid_filter:
                q_short = q_short.where(Application.user_id == uid)
            shortlisted = list(session.exec(q_short).all())

            # 2. Fetch Submitted (limit to 20 by default unless all_submitted=True)
            q_sub = select(Application, Job).join(Job).where(
                Application.status == ApplicationStatus.SUBMITTED
            ).where(
                Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
            ).order_by(Application.submitted_at.desc())
            if _uid_filter:
                q_sub = q_sub.where(Application.user_id == uid)

            # Get total count of submitted for the UI toggle button
            q_sub_count = select(func.count(Application.id)).join(Job).where(
                Application.status == ApplicationStatus.SUBMITTED
            ).where(
                Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
            )
            if _uid_filter:
                q_sub_count = q_sub_count.where(Application.user_id == uid)
            total_submitted_count = session.exec(q_sub_count).first() or 0

            if not all_submitted:
                q_sub = q_sub.limit(20)
            submitted = list(session.exec(q_sub).all())

            # 3. Fetch Interviewing (uncapped since active interviews are few)
            q_int = select(Application, Job).join(Job).where(
                Application.status.in_([ApplicationStatus.INTERVIEWING, ApplicationStatus.OFFER, ApplicationStatus.ACCEPTED])
            ).where(
                Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
            ).order_by(Application.updated_at.desc())
            if _uid_filter:
                q_int = q_int.where(Application.user_id == uid)
            interviewing = list(session.exec(q_int).all())

            # 4. Fetch Rejected (limit to 20 by default)
            q_rej = select(Application, Job).join(Job).where(
                Application.status == ApplicationStatus.REJECTED
            ).where(
                Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
            ).order_by(Application.updated_at.desc()).limit(20)
            if _uid_filter:
                q_rej = q_rej.where(Application.user_id == uid)
            rejected = list(session.exec(q_rej).all())

            # 5. Fetch Skipped (limit to 20)
            q_skip = select(Application, Job).join(Job).where(
                Application.status == ApplicationStatus.SKIPPED
            ).where(
                Job.ghost_flags.is_(None) | ~Job.ghost_flags.contains("aggregator_redirect")
            ).order_by(Application.updated_at.desc()).limit(20)
            if _uid_filter:
                q_skip = q_skip.where(Application.user_id == uid)
            skipped = list(session.exec(q_skip).all())

    from datetime import datetime as _dt

    # Legal work-authorization framing for this user (drives the visa-fit panel
    # and the sponsorship-aware ranking boost below).
    visa_framing = None
    try:
        from app.intelligence.work_auth import assess_profile
        from app.autofill.answer_pack import _get_or_create_profile
        _prof = _get_or_create_profile(user_id=uid if uid and uid != "local" else None)
        visa_framing = assess_profile(_prof)
    except Exception as _e:
        log.debug("visa framing unavailable: %s", _e)

    # For users who need sponsorship, float no-lottery (cap-exempt) and known
    # sponsors to the top — those are the jobs that can actually hire them.
    _boost_sponsorship = bool(visa_framing and getattr(visa_framing, "needs_future_sponsorship", False))

    def _priority(job) -> float:
        """Rank by blended score (fit + hiring intent), plus an urgency/timing
        tiebreak, plus a sponsorship-aware boost for visa users."""
        base = job.blended_score if job.blended_score is not None else (job.rerank_score or 0)
        # Urgency is a strong tiebreak (fresh / hard-to-fill float up) but stays
        # secondary to fit: up to ~+14 on a 0-100 scale.
        try:
            urg = _urgency_of(job)
            if urg:
                base += urg.score * 0.15
        except Exception:
            pass
        if _boost_sponsorship:
            try:
                spons = _sponsorship_of(job)
                if spons and spons.cap_exempt:
                    base += 1000          # no-lottery → absolute top
                elif spons and spons.tone == "good":
                    base += 200           # established sponsor → boosted
                elif spons and spons.explicitly_refuses:
                    base -= 500           # explicitly won't sponsor → sink
            except Exception:
                pass
        return base

    # Highest-priority roles first; recency breaks ties so fresh postings float up.
    shortlisted.sort(
        key=lambda x: (_priority(x[1]), x[1].posted_at or x[1].discovered_at or _dt.min),
        reverse=True,
    )
    manual_queue.sort(key=lambda x: _priority(x[1]), reverse=True)
    # Sort submitted by when they were actually submitted (newest first).
    # Use submitted_at which is set once at submission — background pipeline
    # updates to updated_at (score rescoring, etc.) won't reorder the column.
    submitted.sort(key=lambda x: x[0].submitted_at or x[0].updated_at or _dt.min, reverse=True)
    interviewing.sort(key=lambda x: x[0].updated_at or _dt.min, reverse=True)
    rejected.sort(key=lambda x: x[0].updated_at or _dt.min, reverse=True)

    # Job-based, not company-based: show at most 2 roles per company in the
    # shortlist so a single company can't dominate the view. Keeps the
    # highest-priority 2 (list is already sorted best-first).
    # NOTE: jobs with an empty/unknown company must NOT be collapsed together —
    # otherwise dozens of distinct aggregator postings (RemoteOK/HN often have a
    # blank company field) all share key "" and get capped to 2 globally, which
    # is why the shortlist count showed 19 but only 2 rendered.
    def _cap_per_company(items, cap: int = 2):
        seen: dict[str, int] = {}
        capped = []
        for app_model, job_model in items:
            company = (job_model.company or "").strip().lower()
            # Distinct key per row when company is unknown → never grouped.
            key = company if company else f"__unknown__:{job_model.id}"
            if seen.get(key, 0) >= cap:
                continue
            seen[key] = seen.get(key, 0) + 1
            capped.append((app_model, job_model))
        return capped

    shortlisted = _cap_per_company(shortlisted)
    manual_queue = _cap_per_company(manual_queue)

    return templates.TemplateResponse(
        request=request,
        name="dashboard.html",
        context={
            "shortlisted": shortlisted,
            "bot_filled": bot_filled,
            "manual_queue": manual_queue,
            "submitted": submitted,
            "interviewing": interviewing,
            "skipped": skipped,
            "rejected": rejected,
            "ssr_authed": ssr_authed,
            "visa_framing": visa_framing,
            "total_submitted_count": total_submitted_count,
            "all_submitted": all_submitted,
            "supabase_url": settings.supabase_url,
            "supabase_anon_key": settings.supabase_anon_key,
        }
     )


@app.get("/api/pipeline/live")
def pipeline_live(request: Request) -> dict:
    """Lightweight JSON snapshot of the pipeline for live (poll-driven) updates —
    lets the dashboard surface freshly-ranked jobs without a full page reload."""
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    _uid_filter = uid and uid != "local"

    _SHORTLIST = {ApplicationStatus.SHORTLISTED, ApplicationStatus.TAILORED}
    _INPROGRESS = {ApplicationStatus.AUTOFILLED, ApplicationStatus.AWAITING_USER,
                   ApplicationStatus.READY_TO_SUBMIT}
    _SUBMITTED = {ApplicationStatus.SUBMITTED, ApplicationStatus.INTERVIEWING}

    counts = {"pool": 0, "shortlisted": 0, "submitted": 0, "rejected": 0}
    shortlist: list[dict] = []
    with get_session() as session:
        pq = select(func.count(Job.id))
        if _uid_filter:
            pq = pq.where(Job.user_id == uid)
        counts["pool"] = _scalar(session.exec(pq).one())

        q = select(Application, Job).join(Job).order_by(Job.rerank_score.desc())
        if _uid_filter:
            q = q.where(Application.user_id == uid)
        for app_model, job_model in session.exec(q).all():
            st = app_model.status
            if st in _SHORTLIST or st in _INPROGRESS:
                counts["shortlisted"] += 1
                shortlist.append({
                    "app_id": app_model.id,
                    "title": job_model.title,
                    "company": job_model.company,
                    "location": job_model.location,
                    "remote": bool(job_model.remote),
                    "score": round(job_model.rerank_score) if job_model.rerank_score is not None else None,
                    "track": app_model.apply_track,
                    "url": app_model.apply_url or job_model.url,
                    "status": st.value,
                })
            elif st in _SUBMITTED:
                counts["submitted"] += 1
            elif st == ApplicationStatus.REJECTED:
                counts["rejected"] += 1

    # Is a discovery/ranking run still in flight? (drives client polling cadence)
    running = False
    try:
        from app.db.models import DiscoveryRun
        with get_session() as session:
            rq = select(DiscoveryRun).order_by(DiscoveryRun.id.desc())
            if _uid_filter:
                rq = rq.where(DiscoveryRun.user_id == uid)
            last = session.exec(rq).first()
            if last and (last.status or "") in ("discovering", "ranking", "running", "pending"):
                running = True
    except Exception:
        running = False

    return {"counts": counts, "shortlist": shortlist, "running": running}


@app.get("/application/{application_id}/details")
def application_details(application_id: int, request: Request) -> dict:
    """Return tailored resume + cover letter text for modal preview."""
    from pathlib import Path
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        job = session.get(Job, application.job_id)

    resume_text = ""
    cover_text = ""

    if application.tailored_resume_path:
        try:
            from docx import Document
            doc = Document(application.tailored_resume_path)
            resume_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            resume_text = f"(Could not read resume: {e})"

    if application.cover_letter_path:
        try:
            cover_text = Path(application.cover_letter_path).read_text(encoding="utf-8")
        except Exception as e:
            cover_text = f"(Could not read cover letter: {e})"

    rejection_data = None
    if application.rejection_analysis:
        try:
            import json
            rejection_data = json.loads(application.rejection_analysis)
        except Exception:
            rejection_data = None

    return {
        "id": application_id,
        "company": job.company,
        "title": job.title,
        "apply_url": application.apply_url or job.url,
        "status": application.status.value,
        "source": job.source.value,
        "resume": resume_text,
        "cover_letter": cover_text,
        "rejection_analysis": rejection_data,
    }


@app.get("/application/{application_id}/match")
def application_match(application_id: int, request: Request) -> dict:
    """Why this job matched: overall score, plain-English reason, and the
    per-factor breakdown (skills / experience / location / work_auth)."""
    import json as _json
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        job = session.get(Job, application.job_id)
    breakdown = {}
    if job and job.rerank_breakdown:
        try:
            breakdown = _json.loads(job.rerank_breakdown)
        except (ValueError, TypeError):
            breakdown = {}

    reason = (job.rerank_reasoning if job else "") or ""
    # Reconcile work-auth with Sponsorship Reality: if the employer is a verified
    # H-1B sponsor (strong public filing record), a "needs sponsorship" penalty
    # contradicts the data — lift the work_auth factor and drop the concern so the
    # panel doesn't flag a concern the H-1B record already answers.
    if job:
        try:
            from app.intelligence.h1b_data import lookup as _h1b_lookup
            from app.intelligence.sponsorship import assess as _spons_assess
            rec = _h1b_lookup(job.company or "")
            spons = _spons_assess(company=job.company or "", description=job.description or "", url=job.url or "")
            strong_sponsor = bool((rec and (rec.get("approvals", 0) or 0) >= 50)
                                  or (spons and (spons.cap_exempt or spons.tone == "good")))
            if strong_sponsor and not (spons and spons.explicitly_refuses):
                wa = breakdown.get("work_auth")
                if isinstance(wa, dict) and (wa.get("score", 0) or 0) < 70:
                    wa["score"] = 80
                    note = "Employer is a verified H-1B sponsor"
                    if rec and rec.get("approvals"):
                        note += f" ({rec['approvals']} approvals on record)"
                    wa["note"] = note
                # Drop the sponsorship line from the concerns text.
                if reason and "\nConcerns:" in reason:
                    head, _, tail = reason.partition("\nConcerns:")
                    kept = [c.strip() for c in tail.split(";")
                            if c.strip() and "sponsor" not in c.lower()]
                    reason = head + (("\nConcerns: " + "; ".join(kept)) if kept else "")
        except Exception as _e:
            log.debug("work-auth sponsorship reconcile skipped: %s", _e)

    return {
        "id": application_id,
        "company": job.company if job else "",
        "title": job.title if job else "",
        "location": job.location if job else "",
        "remote": bool(job.remote) if job else False,
        "score": round(job.rerank_score) if (job and job.rerank_score is not None) else None,
        "reason": reason,
        "breakdown": breakdown,
    }


@app.get("/application/{application_id}/senior-review")
def application_senior_review(application_id: int, request: Request) -> dict:
    """A senior-engineer's independent take on this job (fit score + verdict).

    Computed on demand the first time the user opens a job, then cached on the
    Application — moved off the matching loop so matching doesn't pay a second
    serial LLM call per shortlisted job.
    """
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        cached = application.senior_verdict
        job_id = application.job_id

    # Compute + cache on first open.
    if not cached:
        try:
            from app.intelligence.senior_reviewer import SeniorReviewer
            from app.matching.pipeline import _run_senior_review
            _run_senior_review(SeniorReviewer(), job_id, application_id)
        except Exception as e:
            log.warning("On-demand senior review failed for app %d: %s", application_id, e)

    with get_session() as session:
        application = session.get(Application, application_id)
        return {
            "id": application_id,
            "fit_score": application.senior_fit_score,
            "verdict": application.senior_verdict or "",
            "highlight_block": application.custom_highlight_block or "",
            "resume_variant": application.profile_variant or "",
        }


@app.get("/application/{application_id}/autopsy")
def application_autopsy(application_id: int, request: Request) -> dict:
    """Rejection Autopsy: reverse-engineer the hidden bar from people who actually
    hold this role, and tell the user whether they are aiming at the wrong door
    (and the right door to aim at instead). Additive/read-only — never hides or
    edits any job. Degrades gracefully (JD-only) when SERPAPI_KEY is unset."""
    uid = _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        job = session.get(Job, application.job_id) if application else None
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        company, role, jd = job.company, job.title, (job.description or "")
    try:
        from app.autofill.answer_pack import _get_or_create_profile
        from app.intelligence.door_match import CandidateProfile
        from app.intelligence.autopsy import run_autopsy
        profile = _get_or_create_profile(user_id=uid if uid and uid != "local" else None)
        candidate = CandidateProfile.from_user_profile(profile)
        return run_autopsy(company, role, jd, candidate)
    except HTTPException:
        raise
    except Exception as e:
        log.exception("Autopsy failed for app %d: %s", application_id, e)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/fill-pack/{application_id}")
def get_fill_pack(application_id: int, request: Request) -> dict:
    """Returns all data the browser extension needs to fill a job application form."""
    import io, zipfile as _zf
    from pathlib import Path as _P
    _require_owned_application(request, application_id)
    uid = _get_user_id(request)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        job = session.get(Job, application.job_id)
        from app.db.models import UserProfile
        profile = session.exec(select(UserProfile).where(UserProfile.user_id == uid)).first() if uid else None
        needs_tailoring = not (application.tailored_resume_path and application.cover_letter_path)

    # Auto-tailor in background — don't block the fill-pack response.
    # The extension gets the master resume immediately; tailored version
    # will be ready if the user refreshes or applies again later.
    if needs_tailoring:
        try:
            from app.tailoring.tailor import tailor_for_application
            import threading
            threading.Thread(
                target=tailor_for_application,
                args=(application_id,),
                daemon=True,
            ).start()
        except Exception as e:
            log.warning("Auto-tailor background start failed for app %d: %s", application_id, e)

    cover_text = ""
    if application.cover_letter_path:
        try:
            cover_text = _P(application.cover_letter_path).read_text(encoding="utf-8")
        except Exception:
            pass

    resume_text = ""
    if application.tailored_resume_path:
        try:
            resume_text = _P(application.tailored_resume_path).read_text(encoding="utf-8")
        except Exception:
            pass

    p = profile
    pack = {
        "app_id": application_id,
        "job_title": job.title if job else "",
        "company": job.company if job else "",
        "apply_url": application.apply_url or (job.url if job else ""),
        "first_name": p.first_name if p else "",
        "last_name": p.last_name if p else "",
        "email": p.email if p else "",
        "phone": p.phone if p else "",
        "location": p.location if p else "",
        "linkedin_url": p.linkedin_url if p else "",
        "github_url": p.github_url if p else "",
        "portfolio_url": p.portfolio_url if p else "",
        "current_title": p.current_title if p else "",
        "years_experience": p.years_experience if p else 0,
        "salary_min": p.salary_min if p else 0,
        "work_authorization": p.work_authorization if p else "",
        "requires_sponsorship": p.requires_sponsorship if p else False,
        "gender": p.gender if p else "Decline to self-identify",
        "ethnicity": p.ethnicity if p else "Decline to self-identify",
        "veteran_status": p.veteran_status if p else "I am not a protected veteran",
        "disability_status": p.disability_status if p else "No, I do not have a disability, or history/record of having a disability",
        "cover_letter": cover_text,
        "resume_text": resume_text,
    }

    # Add AI-generated essay answers
    try:
        from app.autofill.answer_pack import get_essay_answers
        essay_answers = get_essay_answers(application_id, user_id=uid if uid != "local" else None)
        pack["ai_answers"] = essay_answers
    except Exception as e:
        log.warning("Failed to generate essay answers for app %d: %s", application_id, e)
        pack["ai_answers"] = {}

    # Add hirepath_url and auth_token so extension can save answers back
    from app.config import settings
    # Use request.base_url so local dev hits 127.0.0.1:8000, prod hits hirepath.dev
    _base = str(request.base_url).rstrip("/")
    pack["hirepath_url"] = getattr(settings, "hirepath_url", None) or _base
    token = request.headers.get("Authorization", "").split(" ", 1)[-1]
    pack["auth_token"] = token

    # Add work experience & education (extracted from resume via LLM, cached)
    try:
        from app.autofill.answer_pack import _get_or_extract_experience_education
        exp_edu = _get_or_extract_experience_education(application, profile, user_id=uid if uid != "local" else None)
        pack["work_experience"] = exp_edu.get("work_experience", [])
        pack["education"] = exp_edu.get("education", [])
    except Exception as e:
        log.warning("Failed to extract experience/education for app %d: %s", application_id, e)
        pack["work_experience"] = []
        pack["education"] = []

    return pack


@app.get("/api/fill-pack/{application_id}/resume")
def get_tailored_resume(application_id: int, request: Request) -> dict:
    """Return the tailored resume .docx as base64 so the extension can attach it
    to a form's file input. Auto-tailors first if no resume exists yet."""
    import base64
    from pathlib import Path as _P
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        path = application.tailored_resume_path

    if not path or not _P(path).exists():
        try:
            from app.tailoring.tailor import tailor_for_application
            resume_path, _ = tailor_for_application(application_id)
            path = str(resume_path)
        except Exception as e:
            log.warning("Resume tailoring failed for app %d: %s", application_id, e)
            raise HTTPException(status_code=503, detail="Could not generate resume")

    p = _P(path)
    if not p.exists():
        raise HTTPException(status_code=404, detail="Resume file not found")

    data = p.read_bytes()
    mime = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if p.suffix == ".docx" else "application/octet-stream")
    return {
        "filename": p.name,
        "mime": mime,
        "base64": base64.b64encode(data).decode("ascii"),
    }


class SaveAnswerBody(BaseModel):
    question: str
    answer: str
    app_id: _Opt[int] = None


@app.post("/api/save-answer")
def save_answer(request: Request, body: SaveAnswerBody) -> dict:
    """Save a user-typed answer back to AnswerMemory so it's used next time."""
    from datetime import datetime
    from app.db.models import AnswerMemory
    uid = _require_user(request)
    question = body.question.strip()
    answer = body.answer.strip()
    if not question or not answer:
        raise HTTPException(status_code=400, detail="question and answer required")
    norm = question.lower().strip()
    user_id_arg = uid if uid != "local" else None
    with get_session() as session:
        q = select(AnswerMemory).where(AnswerMemory.label_normalized == norm)
        if user_id_arg:
            q = q.where(AnswerMemory.user_id == user_id_arg)
        existing = session.exec(q).first()
        if existing:
            existing.answer = answer
            existing.use_count += 1
            existing.last_used_at = datetime.utcnow()
            session.add(existing)
        else:
            session.add(AnswerMemory(
                user_id=user_id_arg,
                label_normalized=norm,
                label_original=question,
                answer=answer,
            ))
        session.commit()
    return {"ok": True}


class RecallAnswersBody(BaseModel):
    labels: list[str]


@app.post("/api/recall-answers")
def recall_answers(request: Request, body: RecallAnswersBody) -> dict:
    """Return remembered answers for any of the given field labels.

    Lets the extension pre-fill fields on a NEW application using answers the
    user typed by hand on PREVIOUS applications. Pure cache lookup — free.
    """
    from app.db.models import AnswerMemory
    uid = _require_user(request)
    user_id_arg = uid if uid != "local" else None
    labels = [l.strip() for l in (body.labels or []) if l and l.strip()]
    if not labels:
        return {"answers": {}}
    # Map normalized label -> original label so we can return by the caller's key
    norm_to_orig = {l.lower().strip(): l for l in labels}
    answers: dict[str, str] = {}
    with get_session() as session:
        q = select(AnswerMemory).where(
            AnswerMemory.label_normalized.in_(list(norm_to_orig.keys()))
        )
        if user_id_arg:
            q = q.where(AnswerMemory.user_id == user_id_arg)
        for mem in session.exec(q).all():
            orig = norm_to_orig.get(mem.label_normalized)
            if orig and mem.answer:
                answers[orig] = mem.answer
    return {"answers": answers}


class AskQuestionBody(BaseModel):
    question: str
    app_id: int


@app.post("/api/answer-question")
def answer_question_endpoint(request: Request, body: AskQuestionBody) -> dict:
    """Generate (or retrieve cached) answer for a single essay question.

    Called by the extension only when it finds an unanswered textarea on the
    live form. Cache-hit = free. Cache-miss = ~$0.002 (Haiku, stored forever).
    """
    uid = _require_user(request)
    question = body.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="question required")
    _require_owned_application(request, body.app_id)
    from app.autofill.answer_pack import answer_question
    user_id_arg = uid if uid != "local" else None
    answer = answer_question(question, body.app_id, user_id=user_id_arg)
    return {"answer": answer, "cached": bool(answer)}


@app.get("/api/extension/download")
def download_extension():
    """Bundle the extension/ folder as a downloadable zip for Chrome."""
    import io, zipfile as _zf, os as _os, traceback as _tb
    from fastapi.responses import StreamingResponse
    try:
        # extension/ is two levels up from app/api/server.py  →  repo_root/extension/
        ext_dir = _os.path.abspath(_os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "..", "..", "extension"))
        if not _os.path.isdir(ext_dir):
            raise HTTPException(status_code=404, detail=f"Extension folder not found: {ext_dir}")
        buf = io.BytesIO()
        with _zf.ZipFile(buf, "w", _zf.ZIP_DEFLATED) as zf:
            for root, _, files in _os.walk(ext_dir):
                for fname in sorted(files):
                    fpath = _os.path.join(root, fname)
                    arcname = "hirepath-extension/" + _os.path.relpath(fpath, ext_dir)
                    zf.write(fpath, arcname)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=hirepath-extension.zip"},
        )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to build zip: {exc}\n{_tb.format_exc()}")


@app.get("/extension")
def extension_page(request: Request):
    return templates.TemplateResponse(request=request, name="extension.html", context={})


@app.post("/application/{application_id}/submit")
def mark_as_submitted(application_id: int, request: Request) -> dict:
    from datetime import datetime
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        application.status = ApplicationStatus.SUBMITTED
        application.submitted_at = datetime.utcnow()
        application.updated_at = datetime.utcnow()
        session.add(application)
        session.commit()
    return {"success": True, "application_id": application_id}


@app.post("/application/{application_id}/skip")
def skip_application(application_id: int, request: Request) -> dict:
    from datetime import datetime
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        application.status = ApplicationStatus.SKIPPED
        application.updated_at = datetime.utcnow()
        session.add(application)
        session.commit()
    return {"success": True, "application_id": application_id}



@app.post("/application/{application_id}/reject")
def mark_as_rejected(application_id: int, request: Request) -> dict:
    """Manually mark an application as rejected (you received a rejection)."""
    from datetime import datetime
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        application.status = ApplicationStatus.REJECTED
        application.updated_at = datetime.utcnow()
        application.notes = (application.notes or "") + f"\nMarked rejected on {datetime.utcnow():%Y-%m-%d}."
        session.add(application)
        session.commit()
    return {"success": True, "application_id": application_id}


# Phase 1.5 — outcome tracking. Candidates mark real outcomes so HirePath
# learns which profiles actually get interviews/offers (not just matches).
_OUTCOME_STATUSES = {
    "interviewing": ApplicationStatus.INTERVIEWING,
    "offer": ApplicationStatus.OFFER,
    "accepted": ApplicationStatus.ACCEPTED,
    "rejected": ApplicationStatus.REJECTED,
}


@app.post("/application/{application_id}/outcome")
def mark_outcome(application_id: int, request: Request, outcome: str) -> dict:
    """Record a real hiring outcome (interviewing | offer | accepted | rejected)."""
    from datetime import datetime
    _require_owned_application(request, application_id)
    status = _OUTCOME_STATUSES.get((outcome or "").lower())
    if not status:
        raise HTTPException(status_code=400, detail=f"Invalid outcome: {outcome}")
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        application.status = status
        application.updated_at = datetime.utcnow()
        application.notes = (application.notes or "") + f"\nOutcome '{outcome}' on {datetime.utcnow():%Y-%m-%d}."
        session.add(application)
        session.commit()
    return {"success": True, "application_id": application_id, "status": status.value}


@app.get("/api/funnel")
def application_funnel(request: Request) -> dict:
    """Outcome funnel for the signed-in user: applied -> interview -> offer ->
    accepted, with conversion rates. The data that lets HirePath eventually
    optimize for real hiring outcomes, not just match similarity."""
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    _submitted_like = [ApplicationStatus.SUBMITTED, ApplicationStatus.INTERVIEWING,
                       ApplicationStatus.OFFER, ApplicationStatus.ACCEPTED, ApplicationStatus.REJECTED]
    _interview_like = [ApplicationStatus.INTERVIEWING, ApplicationStatus.OFFER, ApplicationStatus.ACCEPTED]
    _offer_like = [ApplicationStatus.OFFER, ApplicationStatus.ACCEPTED]
    with get_session() as session:
        def _count(statuses):
            q = select(func.count(Application.id)).where(Application.status.in_(statuses))
            if user_id_arg:
                q = q.where(Application.user_id == user_id_arg)
            return int(session.exec(q).first() or 0)
        applied = _count(_submitted_like)
        interviewing = _count(_interview_like)
        offers = _count(_offer_like)
        accepted = _count([ApplicationStatus.ACCEPTED])
    rate = lambda a, b: round(a / b * 100) if b else 0
    return {
        "applied": applied, "interviewing": interviewing, "offers": offers, "accepted": accepted,
        "applied_to_interview": rate(interviewing, applied),
        "interview_to_offer": rate(offers, interviewing),
        "offer_to_accepted": rate(accepted, offers),
    }


def _discover_then_match(user_id) -> None:
    """Discover → rank, tracking staged status (discovering → ranking → done)
    in a DiscoveryRun row so the UI can show live progress + a final summary."""
    from app.discovery.pipeline import create_discovery_run, finish_discovery_run
    run_id = create_discovery_run(user_id)
    # Tailor the keyword search to the user's saved Target Roles when set.
    roles = _get_target_roles(user_id) or None
    # Load the profile once for keyword augmentation (cap-exempt + internships +
    # department fallback). Bounded so discovery cost never balloons.
    _prof = None
    try:
        from app.autofill.answer_pack import _get_or_create_profile
        _prof = _get_or_create_profile(user_id=user_id if user_id and user_id != "local" else None)
    except Exception:
        _prof = None

    def _add_kw(lst, term):
        if term and term.lower() not in [r.lower() for r in lst]:
            lst.append(term)

    try:
        # If the user has no explicit roles, fall back to department-aware keywords.
        if not roles:
            roles = _department_keywords(_prof) or None
        if roles:
            # Sponsorship-needing users: nudge toward cap-exempt (no-lottery) roles.
            if _user_needs_sponsorship(user_id):
                primary = roles[0]
                _add_kw(roles, f"research {primary}")
                _add_kw(roles, f"{primary} university")
            # Internship toggle: append internship variants of the primary roles.
            if _prof is not None and getattr(_prof, "include_internships_in_discovery", False):
                for base in roles[:3]:
                    _add_kw(roles, f"{base} intern")
                _add_kw(roles, "internship")
    except Exception as _se:
        log.debug("discovery keyword augmentation skipped: %s", _se)
    try:
        run_discovery(user_id, run_id=run_id, keywords=roles)   # marks the row 'ranking' + per-source counts
    except Exception as e:
        log.exception("Discovery failed: %s", e)
        finish_discovery_run(run_id, "error", error=str(e))
        return
    try:
        shortlisted = run_matching(user_id)
        finish_discovery_run(run_id, "done", total_shortlisted=len(shortlisted or []))
    except Exception as e:
        log.exception("Matching failed: %s", e)
        finish_discovery_run(run_id, "error", error=str(e))


@app.post("/run/discovery")
@_rate_limit("10/minute")
def trigger_discovery(request: Request, bg: BackgroundTasks) -> dict:
    uid = _get_user_id(request)
    # Gate: no resume → no discovery. Without a resume there is nothing to match
    # against, so scraping jobs into the user's pool would only show noise.
    if not _user_has_resume(uid):
        raise HTTPException(
            status_code=400,
            detail="Upload your resume (or fill in your profile) before discovering jobs.",
        )
    # Gate: no target roles → no discovery. Without roles we don't know what
    # titles to search for, so the keyword sources would fall back to a generic
    # list that may not match the user at all.
    if not _get_target_roles(uid):
        raise HTTPException(
            status_code=400,
            detail="Add at least one target role before discovering jobs.",
        )
    # Gate: prevent overlapping runs + enforce a cooldown so repeated clicks
    # don't waste API calls / LLM tokens (discovery also auto-runs every 6h).
    allowed, detail = _discovery_gate(uid)
    if not allowed:
        raise HTTPException(status_code=429, detail=detail)
    bg.add_task(_discover_then_match, uid if uid != "local" else None)
    return {"started": "discovery"}


@app.delete("/run/discovery")
def cancel_discovery(request: Request) -> dict:
    """Mark the active discovery run as cancelled so the poller stops tracking it."""
    from app.db.models import DiscoveryRun
    uid = _get_user_id(request)
    with get_session() as session:
        q = select(DiscoveryRun).order_by(desc(DiscoveryRun.id))
        if uid and uid != "local":
            q = q.where(DiscoveryRun.user_id == uid)
        run = session.exec(q).first()
        if run and run.status in ("discovering", "ranking"):
            run.status = "cancelled"
            run.error = "Cancelled by user"
            session.add(run)
            session.commit()
            return {"cancelled": True}
    return {"cancelled": False}


def _discovery_gate(uid) -> tuple[bool, str]:
    """Block a manual discovery if one is in progress or within the cooldown."""
    from datetime import datetime
    from app.config import settings
    from app.db.models import DiscoveryRun
    with get_session() as session:
        q = select(DiscoveryRun).order_by(desc(DiscoveryRun.id))
        if uid and uid != "local":
            q = q.where(DiscoveryRun.user_id == uid)
        run = session.exec(q).first()
    if not run:
        return True, ""
    now = datetime.utcnow()
    # In progress — block unless it looks stale (likely crashed/hung).
    if run.status in ("discovering", "ranking"):
        age = (now - (run.started_at or now)).total_seconds()
        if age < 1800:  # 30 min
            return False, "A discovery is already running — please wait for it to finish."
        return True, ""  # stale, allow a fresh run
    # Cooldown since the last completed run.
    cooldown = max(0, settings.discovery_cooldown_hours) * 3600
    ref = run.finished_at or run.started_at
    if cooldown and ref:
        elapsed = (now - ref).total_seconds()
        if elapsed < cooldown:
            remaining = int(cooldown - elapsed)
            h, m = remaining // 3600, (remaining % 3600) // 60
            when = f"{h}h {m}m" if h else f"{m}m"
            return False, f"Discovery already ran recently. Next run available in ~{when} (it also auto-runs every {settings.discovery_cooldown_hours}h)."
    return True, ""


@app.post("/run/matching")
def trigger_matching(request: Request, bg: BackgroundTasks) -> dict:
    uid = _get_user_id(request)
    bg.add_task(run_matching, uid if uid != "local" else None)
    return {"started": "matching"}


# ── Usage / Plan helpers ─────────────────────────────────────────────────────

# ── Founding-user trial ──────────────────────────────────────────────────────
def _get_trial(uid):
    """Return this user's TrialGrant row, or None. Always return None to grant unlimited access."""
    return None


def _grant_trial_if_eligible(uid):
    """Grant a trial to the first N users (idempotent). Returns the grant or None."""
    if not uid or uid == "local":
        return None
    from app.db.models import TrialGrant
    from app.config import settings
    try:
        with get_session() as session:
            existing = session.exec(select(TrialGrant).where(TrialGrant.user_id == uid)).first()
            if existing:
                return existing
            cnt = session.exec(select(func.count(TrialGrant.id))).one()
            cnt = int(cnt if not isinstance(cnt, (list, tuple)) else cnt[0])
            if cnt >= settings.trial_max_users:
                return None
            g = TrialGrant(user_id=uid, jobs_quota=settings.trial_job_quota, jobs_used=0)
            session.add(g)
            session.commit()
            session.refresh(g)
            log.info("Granted founding-user trial #%d to %s", cnt + 1, uid)
            return g
    except Exception as e:
        log.debug("trial grant skipped: %s", e)
        return None


def _trial_active(uid) -> bool:
    g = _get_trial(uid)
    return bool(g and g.jobs_used < g.jobs_quota)


def _increment_trial(uid, n: int = 1):
    if not uid or uid == "local":
        return
    from app.db.models import TrialGrant
    try:
        with get_session() as session:
            g = session.exec(select(TrialGrant).where(TrialGrant.user_id == uid)).first()
            if g:
                g.jobs_used += n
                session.add(g)
                session.commit()
    except Exception as e:
        log.debug("trial increment skipped: %s", e)


# ── Referral program ─────────────────────────────────────────────────────────
def _get_user_email(request: Request) -> str | None:
    """Authenticated user's email from the Supabase token (for admin gating)."""
    from app.config import settings
    if not settings.use_supabase:
        return None
    auth = request.headers.get("Authorization", "")
    token = auth.split(" ", 1)[1] if auth.startswith("Bearer ") else request.cookies.get("sb_token")
    if not token:
        return None
    try:
        from app.db.supabase_client import verify_jwt
        return (verify_jwt(token) or {}).get("email")
    except Exception:
        return None


def _gen_referral_code() -> str:
    import secrets, string
    return "".join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(8))


def _ensure_referral_code(uid) -> str | None:
    """Return the user's referral code, generating a unique one if missing."""
    if not uid or uid == "local":
        return None
    from app.db.models import UserProfile
    from app.autofill.answer_pack import _get_or_create_profile
    _get_or_create_profile(user_id=uid)   # make sure a profile row exists
    with get_session() as session:
        p = session.exec(select(UserProfile).where(UserProfile.user_id == uid)).first()
        if not p:
            return None
        if p.referral_code:
            return p.referral_code
        for _ in range(8):
            code = _gen_referral_code()
            if not session.exec(select(UserProfile).where(UserProfile.referral_code == code)).first():
                p.referral_code = code
                p.updated_at = _dt.utcnow()
                session.add(p)
                session.commit()
                return code
    return None


def _referral_count(uid) -> int:
    if not uid or uid == "local":
        return 0
    from app.db.models import UserProfile
    with get_session() as session:
        c = session.exec(select(func.count(UserProfile.id)).where(UserProfile.referred_by_id == uid)).one()
    return int(c if not isinstance(c, (list, tuple)) else c[0])


def _grant_referral_reward(referrer_uid: str, count: int) -> bool:
    """Idempotently unlock the referral reward: N days of the reward plan +
    an in-app notification. Returns True if newly granted."""
    from datetime import timedelta
    from app.config import settings
    from app.db.models import (UserReferralReward, UserSubscription,
                               UserNotification, PlanTier)
    try:
        plan = PlanTier(settings.referral_reward_plan)
    except ValueError:
        plan = PlanTier.PRO
    expires = _dt.utcnow() + timedelta(days=settings.referral_reward_days)
    try:
        with get_session() as session:
            if session.exec(select(UserReferralReward).where(UserReferralReward.user_id == referrer_uid)).first():
                return False   # already rewarded
            sub = session.exec(select(UserSubscription).where(UserSubscription.user_id == referrer_uid)).first()
            if not sub:
                sub = UserSubscription(user_id=referrer_uid, plan=plan, current_period_end=expires)
            else:
                sub.plan = plan
                sub.current_period_end = expires
                sub.updated_at = _dt.utcnow()
            session.add(sub)
            session.add(UserReferralReward(
                user_id=referrer_uid, referred_count=count, status="active",
                reward_plan=plan.value, expires_at=expires))
            session.add(UserNotification(
                user_id=referrer_uid, title="Premium Unlocked! 🎁", type="referral_reward",
                message=(f"You referred {count} friends — enjoy {settings.referral_reward_days} "
                         f"days of {plan.value.upper()} on us. Thank you!"),
                read=False))
            session.commit()
            log.info("Referral reward granted to %s (%d referrals)", referrer_uid, count)
            return True
    except Exception as e:
        log.exception("Referral reward grant failed: %s", e)
        return False


@app.get("/api/referral")
def get_referral(request: Request) -> dict:
    """The current user's referral code, link, progress, and reward status."""
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if uid == "local":
        return {"code": "LOCALDEV", "link": "", "count": 0,
                "threshold": settings.referral_threshold, "reward_unlocked": False}
    code = _ensure_referral_code(uid)
    count = _referral_count(uid)
    from app.db.models import UserReferralReward
    with get_session() as session:
        reward = session.exec(select(UserReferralReward).where(UserReferralReward.user_id == uid)).first()
    base = str(request.base_url).rstrip("/")
    return {
        "code": code,
        "link": f"{base}/login?ref={code}" if code else "",
        "count": count,
        "threshold": settings.referral_threshold,
        "reward_unlocked": bool(reward),
        "reward_days": settings.referral_reward_days,
        "reward_plan": settings.referral_reward_plan,
    }


class ReferralClaimBody(BaseModel):
    code: str


@app.post("/api/referral/claim")
@_rate_limit("10/minute")
def claim_referral(request: Request, body: ReferralClaimBody) -> dict:
    """Record that the current (new) user was referred by `code`. One-time;
    can't refer yourself. Unlocks the referrer's reward when they hit the threshold."""
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if uid == "local":
        return {"ok": False, "reason": "local_dev"}
    code = (body.code or "").strip().upper()
    if not code:
        return {"ok": False, "reason": "no_code"}
    from app.db.models import UserProfile
    from app.autofill.answer_pack import _get_or_create_profile
    _get_or_create_profile(user_id=uid)
    ref_uid = None
    with get_session() as session:
        me = session.exec(select(UserProfile).where(UserProfile.user_id == uid)).first()
        if not me:
            return {"ok": False, "reason": "no_profile"}
        if me.referred_by_id:
            return {"ok": False, "reason": "already_referred"}
        referrer = session.exec(select(UserProfile).where(UserProfile.referral_code == code)).first()
        if not referrer or referrer.user_id == uid:
            return {"ok": False, "reason": "invalid_code"}
        me.referred_by_id = referrer.user_id
        me.updated_at = _dt.utcnow()
        session.add(me)
        session.commit()
        ref_uid = referrer.user_id
    count = _referral_count(ref_uid)
    if count >= settings.referral_threshold:
        _grant_referral_reward(ref_uid, count)
    return {"ok": True, "referrer_reached": count}


# ── Promo / coupon codes ───────────────────────────────────────────────────────

class CouponRedeemBody(BaseModel):
    code: str


@app.post("/api/coupon/redeem")
@_rate_limit("10/minute")
def redeem_coupon(request: Request, body: CouponRedeemBody) -> dict:
    """Redeem a promo code. One redemption per user per code."""
    from datetime import timedelta
    from app.db.models import (Coupon, CouponRedemption, UserSubscription,
                                UserNotification, PlanTier)
    uid = _get_user_id(request)
    if not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    code = body.code.strip().upper()
    with get_session() as session:
        coupon = session.exec(select(Coupon).where(Coupon.code == code)).first()
        if not coupon or not coupon.is_active:
            return {"ok": False, "reason": "invalid_code", "message": "Code not found or inactive."}
        if coupon.expires_at and coupon.expires_at < _dt.utcnow():
            return {"ok": False, "reason": "expired", "message": "This code has expired."}
        if coupon.max_uses is not None and coupon.uses_count >= coupon.max_uses:
            return {"ok": False, "reason": "used_up", "message": "This code has reached its usage limit."}
        already = session.exec(
            select(CouponRedemption)
            .where(CouponRedemption.coupon_id == coupon.id)
            .where(CouponRedemption.user_id == uid)
        ).first()
        if already:
            return {"ok": False, "reason": "already_redeemed", "message": "You've already used this code."}
        # Grant plan upgrade
        try:
            plan = PlanTier(coupon.reward_plan)
        except ValueError:
            plan = PlanTier.PRO
        expires = _dt.utcnow() + timedelta(days=coupon.reward_days)
        sub = session.exec(select(UserSubscription).where(UserSubscription.user_id == uid)).first()
        if sub:
            sub.plan = plan
            sub.current_period_end = expires
            session.add(sub)
        else:
            session.add(UserSubscription(user_id=uid, plan=plan, current_period_end=expires))
        coupon.uses_count += 1
        session.add(coupon)
        session.add(CouponRedemption(coupon_id=coupon.id, user_id=uid))
        session.add(UserNotification(
            user_id=uid, title=f"Promo code applied! 🎉", type="coupon_reward",
            message=(f"Code {code} unlocked {coupon.reward_days} days of "
                     f"{plan.value.upper()}. {coupon.description}"),
            read=False,
        ))
        session.commit()
        log.info("Coupon %s redeemed by %s → %s for %dd", code, uid, plan.value, coupon.reward_days)
    return {
        "ok": True,
        "plan": plan.value,
        "days": coupon.reward_days,
        "message": f"🎉 {coupon.reward_days} days of {plan.value.upper()} unlocked! {coupon.description}".strip(),
    }


# ── Coupon admin CRUD ──────────────────────────────────────────────────────────

class CouponCreateBody(BaseModel):
    code: str
    description: str = ""
    reward_plan: str = "pro"
    reward_days: int = 30
    max_uses: Optional[int] = None
    expires_at: Optional[str] = None   # ISO date string or None


class CouponUpdateBody(BaseModel):
    is_active: Optional[bool] = None
    description: Optional[str] = None
    max_uses: Optional[int] = None
    expires_at: Optional[str] = None


@app.get("/api/admin/coupons")
def admin_list_coupons(request: Request) -> dict:
    """List all promo codes with usage stats. Admin-only."""
    _require_admin_user(request)
    from app.db.models import Coupon
    with get_session() as session:
        coupons = session.exec(select(Coupon).order_by(Coupon.created_at.desc())).all()
    return {"coupons": [
        {
            "id": c.id, "code": c.code, "description": c.description,
            "reward_plan": c.reward_plan, "reward_days": c.reward_days,
            "max_uses": c.max_uses, "uses_count": c.uses_count,
            "is_active": c.is_active,
            "expires_at": c.expires_at.isoformat() if c.expires_at else None,
            "created_at": c.created_at.isoformat(),
        }
        for c in coupons
    ]}


@app.post("/api/admin/coupons")
def admin_create_coupon(request: Request, body: CouponCreateBody) -> dict:
    """Create a new promo code. Admin-only."""
    _require_admin_user(request)
    from app.db.models import Coupon
    from datetime import datetime as _dtp
    admin_uid = _get_user_id(request)
    code = body.code.strip().upper()
    if not code:
        raise HTTPException(status_code=400, detail="Code cannot be empty.")
    expires = None
    if body.expires_at:
        try:
            expires = _dtp.fromisoformat(body.expires_at.replace("Z", "+00:00")).replace(tzinfo=None)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid expires_at format.")
    with get_session() as session:
        existing = session.exec(select(Coupon).where(Coupon.code == code)).first()
        if existing:
            raise HTTPException(status_code=409, detail=f"Code '{code}' already exists.")
        coupon = Coupon(
            code=code, description=body.description,
            reward_plan=body.reward_plan, reward_days=body.reward_days,
            max_uses=body.max_uses, expires_at=expires,
            created_by=admin_uid if admin_uid != "local" else None,
        )
        session.add(coupon)
        session.commit()
        session.refresh(coupon)
    return {"ok": True, "id": coupon.id, "code": coupon.code}


@app.patch("/api/admin/coupons/{coupon_id}")
def admin_update_coupon(coupon_id: int, request: Request, body: CouponUpdateBody) -> dict:
    """Toggle active/inactive or update a coupon. Admin-only."""
    _require_admin_user(request)
    from app.db.models import Coupon
    from datetime import datetime as _dtp
    with get_session() as session:
        coupon = session.get(Coupon, coupon_id)
        if not coupon:
            raise HTTPException(status_code=404, detail="Coupon not found.")
        if body.is_active is not None:
            coupon.is_active = body.is_active
        if body.description is not None:
            coupon.description = body.description
        if body.max_uses is not None:
            coupon.max_uses = body.max_uses
        if body.expires_at is not None:
            try:
                coupon.expires_at = _dtp.fromisoformat(body.expires_at.replace("Z", "+00:00")).replace(tzinfo=None)
            except ValueError:
                raise HTTPException(status_code=400, detail="Invalid expires_at format.")
        session.add(coupon)
        session.commit()
    return {"ok": True}


@app.delete("/api/admin/coupons/{coupon_id}")
def admin_delete_coupon(coupon_id: int, request: Request) -> dict:
    """Delete a coupon (also removes redemptions). Admin-only."""
    _require_admin_user(request)
    from app.db.models import Coupon, CouponRedemption
    with get_session() as session:
        coupon = session.get(Coupon, coupon_id)
        if not coupon:
            raise HTTPException(status_code=404, detail="Coupon not found.")
        session.exec(  # type: ignore[call-overload]
            select(CouponRedemption).where(CouponRedemption.coupon_id == coupon_id)
        )
        for r in session.exec(select(CouponRedemption).where(CouponRedemption.coupon_id == coupon_id)).all():
            session.delete(r)
        session.delete(coupon)
        session.commit()
    return {"ok": True}


# ── Owner-only admin dashboard ────────────────────────────────────────────────
_ADMIN_HTML = """<!doctype html><html><head><meta charset=utf-8>
<meta name=viewport content="width=device-width,initial-scale=1"><title>JobAgent · Admin</title><style>
:root{--canvas:#F4F1EA;--surface:#FCFAF5;--surface-2:#EFEADF;--ink:#2E2A24;--muted:#8C857A;--border:#E6E0D3;--sage:#2FB4A6;--sage-700:#16847A}
*{box-sizing:border-box}body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;background:var(--canvas);color:var(--ink);margin:0;padding:28px}
h1{font-size:20px;margin:0 0 2px}.sub{color:var(--muted);font-size:13px;margin:0 0 22px}
.grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(150px,1fr));gap:14px;margin-bottom:26px}
.kpi{background:var(--surface);border:1px solid var(--border);border-radius:18px;padding:18px}
.kpi .n{font-size:26px;font-weight:800;color:var(--sage-700)}.kpi .l{font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:.05em;color:var(--muted);margin-top:4px}
.card{background:var(--surface);border:1px solid var(--border);border-radius:18px;padding:18px;margin-bottom:20px}
h2{font-size:13px;text-transform:uppercase;letter-spacing:.06em;color:var(--muted);margin:0 0 12px}
table{width:100%;border-collapse:collapse;font-size:13px}th,td{text-align:left;padding:8px 6px;border-bottom:1px solid var(--border)}
th{font-size:10px;text-transform:uppercase;color:var(--muted)}.pill{font-size:10px;font-weight:700;padding:2px 8px;border-radius:999px;background:var(--surface-2);color:var(--sage-700)}
.pill-off{background:#fee2e2;color:#b91c1c}.pill-on{background:#d1fae5;color:#065f46}
#err{color:#b91c1c;font-size:14px;margin-bottom:12px}
.row{display:flex;gap:8px;flex-wrap:wrap;align-items:flex-end;margin-bottom:14px}
input,select{font-size:13px;padding:7px 10px;border:1px solid var(--border);border-radius:8px;background:var(--surface);color:var(--ink);outline:none}
input:focus,select:focus{border-color:var(--sage)}
button.btn{font-size:13px;font-weight:700;padding:7px 16px;border:none;border-radius:8px;cursor:pointer;background:var(--sage);color:#fff}
button.btn:hover{background:var(--sage-700)}
button.btn-red{background:#ef4444}button.btn-red:hover{background:#b91c1c}
button.btn-sm{font-size:11px;padding:4px 10px}
#coupon-msg{font-size:13px;margin-top:6px}
</style></head><body>
<h1>📊 JobAgent — Owner Dashboard</h1><p class=sub>Live metrics. Owner-only.</p>
<div id=err></div>
<div class=grid id=kpis></div>

<div class=card><h2>🎟️ Promo codes</h2>
<div class=row>
  <div><label style="font-size:11px;font-weight:700;display:block;margin-bottom:3px;color:var(--muted)">CODE</label><input id=c-code placeholder="LAUNCH50" style="text-transform:uppercase;width:130px"></div>
  <div><label style="font-size:11px;font-weight:700;display:block;margin-bottom:3px;color:var(--muted)">PLAN</label>
    <select id=c-plan><option value=pro>PRO (30d)</option><option value=basic>BASIC</option></select></div>
  <div><label style="font-size:11px;font-weight:700;display:block;margin-bottom:3px;color:var(--muted)">DAYS</label><input id=c-days type=number value=30 style="width:70px"></div>
  <div><label style="font-size:11px;font-weight:700;display:block;margin-bottom:3px;color:var(--muted)">MAX USES</label><input id=c-uses placeholder="∞" style="width:80px"></div>
  <div><label style="font-size:11px;font-weight:700;display:block;margin-bottom:3px;color:var(--muted)">EXPIRES</label><input id=c-exp type=date style="width:140px"></div>
  <div><label style="font-size:11px;font-weight:700;display:block;margin-bottom:3px;color:var(--muted)">NOTE</label><input id=c-desc placeholder="Optional description" style="width:200px"></div>
  <button class=btn onclick=createCoupon()>+ Create</button>
</div>
<p id=coupon-msg></p>
<table id=coupons-tbl><thead><tr><th>Code</th><th>Plan</th><th>Days</th><th>Uses</th><th>Max</th><th>Expires</th><th>Note</th><th>Status</th><th></th></tr></thead><tbody></tbody></table>
</div>

<div class=card><h2>⭐️ Reviews &amp; Testimonials Moderation</h2>
<table id=reviews-tbl><thead><tr><th>Reviewer</th><th>Rating</th><th>Review Content</th><th>Public</th><th>Featured</th><th>Actions</th></tr></thead><tbody></tbody></table>
</div>

<div class=card><h2>👥 Top referrers</h2><table id=reftbl><thead><tr><th>User</th><th>Email</th><th>Code</th><th>Referrals</th><th>Reward</th></tr></thead><tbody></tbody></table></div>
<script>
function H(){const t=localStorage.getItem('sb_token');return t?{'Authorization':'Bearer '+t,'Content-Type':'application/json'}:{'Content-Type':'application/json'}}
async function load(){
  try{
    const m=await fetch('/api/admin/metrics',{headers:H()});
    if(m.status===403){document.getElementById('err').textContent='🔒 Not authorized — sign in with an admin account.';return;}
    if(!m.ok){document.getElementById('err').textContent='Error '+m.status;return;}
    const d=await m.json();
    const cards=[['Total users',d.total_users],['Active (7d)',d.active_users_7d],['Paid subs',d.paid_subscriptions],
      ['MRR','$'+d.mrr_usd],['ARR','$'+d.arr_usd],['Referred signups',d.referred_signups],
      ['Trial users',d.trial_users],['Applications',d.total_applications]];
    document.getElementById('kpis').innerHTML=cards.map(c=>`<div class=kpi><div class=n>${c[1]}</div><div class=l>${c[0]}</div></div>`).join('');
    const r=await fetch('/api/admin/referrals',{headers:H()});const rd=await r.json();
    document.querySelector('#reftbl tbody').innerHTML=(rd.top_referrers||[]).map(x=>
      `<tr><td>${x.name}</td><td>${x.email}</td><td><code>${x.code||''}</code></td><td><b>${x.count}</b></td>
      <td>${x.rewarded?'<span class=pill>🎁 Rewarded</span>':''}</td></tr>`).join('')
      || '<tr><td colspan=5 style="color:#8C857A">No referrals yet.</td></tr>';
    await loadCoupons();
    await loadReviews();
  }catch(e){document.getElementById('err').textContent=''+e;}
}
async function loadCoupons(){
  const res=await fetch('/api/admin/coupons',{headers:H()});
  if(!res.ok)return;
  const {coupons}=await res.json();
  document.querySelector('#coupons-tbl tbody').innerHTML=coupons.length
    ? coupons.map(c=>`<tr>
        <td><code style="font-weight:700">${c.code}</code></td>
        <td>${c.reward_plan.toUpperCase()}</td>
        <td>${c.reward_days}d</td>
        <td>${c.uses_count}</td>
        <td>${c.max_uses??'∞'}</td>
        <td style="font-size:11px">${c.expires_at?c.expires_at.slice(0,10):'—'}</td>
        <td style="font-size:11px;color:#8C857A">${c.description||''}</td>
        <td>${c.is_active?'<span class="pill pill-on">Active</span>':'<span class="pill pill-off">Off</span>'}</td>
        <td style="white-space:nowrap">
          <button class="btn btn-sm" onclick="toggleCoupon(${c.id},${!c.is_active})">${c.is_active?'Disable':'Enable'}</button>
          <button class="btn btn-sm btn-red" style="margin-left:4px" onclick="deleteCoupon(${c.id},'${c.code}')">Del</button>
        </td></tr>`).join('')
    : '<tr><td colspan=9 style="color:#8C857A;padding:16px 0">No coupons yet — create one above.</td></tr>';
}
async function createCoupon(){
  const code=document.getElementById('c-code').value.trim().toUpperCase();
  if(!code){alert('Enter a code.');return;}
  const body={
    code,
    reward_plan:document.getElementById('c-plan').value,
    reward_days:parseInt(document.getElementById('c-days').value)||30,
    max_uses:document.getElementById('c-uses').value?parseInt(document.getElementById('c-uses').value):null,
    expires_at:document.getElementById('c-exp').value||null,
    description:document.getElementById('c-desc').value.trim(),
  };
  const res=await fetch('/api/admin/coupons',{method:'POST',headers:H(),body:JSON.stringify(body)});
  const d=await res.json();
  const msg=document.getElementById('coupon-msg');
  if(res.ok){msg.style.color='#065f46';msg.textContent='✅ Code '+d.code+' created!';
    document.getElementById('c-code').value='';document.getElementById('c-desc').value='';
    await loadCoupons();}
  else{msg.style.color='#b91c1c';msg.textContent='❌ '+(d.detail||'Error');}
}
async function toggleCoupon(id,active){
  await fetch('/api/admin/coupons/'+id,{method:'PATCH',headers:H(),body:JSON.stringify({is_active:active})});
  await loadCoupons();
}
async function deleteCoupon(id,code){
  if(!confirm('Delete coupon '+code+'? This cannot be undone.'))return;
  await fetch('/api/admin/coupons/'+id,{method:'DELETE',headers:H()});
  await loadCoupons();
}
async function loadReviews(){
  const res=await fetch('/api/admin/reviews',{headers:H()});
  if(!res.ok)return;
  const {reviews}=await res.json();
  document.querySelector('#reviews-tbl tbody').innerHTML=reviews.length
    ? reviews.map(r=>`<tr>
        <td><b>${r.user_name}</b><br><span style="font-size:10px;color:#8C857A">${r.created_at?r.created_at.slice(0,10):''}</span></td>
        <td style="color:#d97706;font-size:14px">${'★'.repeat(r.rating)}${'☆'.repeat(5-r.rating)}</td>
        <td style="max-width:350px;word-break:break-word">${r.content}</td>
        <td>${r.is_public?'<span class="pill pill-on">Public</span>':'<span class="pill pill-off">Hidden</span>'}</td>
        <td>${r.is_featured?'<span class="pill pill-on" style="background:#dbeafe;color:#1e40af">★ Featured</span>':'<span style="color:#8C857A">—</span>'}</td>
        <td style="white-space:nowrap">
          <button class="btn btn-sm" onclick="approveReview(${r.id},${!r.is_public})">${r.is_public?'Hide':'Publish'}</button>
          <button class="btn btn-sm" style="margin-left:4px;background:#1e40af" onclick="featureReview(${r.id},${!r.is_featured})">${r.is_featured?'Unfeature':'Feature'}</button>
        </td></tr>`).join('')
    : '<tr><td colspan=6 style="color:#8C857A;padding:16px 0">No reviews submitted yet.</td></tr>';
}
async function approveReview(id,approve){
  await fetch('/api/admin/reviews/'+id+'/approve',{method:'POST',headers:H(),body:JSON.stringify({approve})});
  await loadReviews();
}
async function featureReview(id,feature){
  await fetch('/api/admin/reviews/'+id+'/feature',{method:'POST',headers:H(),body:JSON.stringify({feature})});
  await loadReviews();
}
load();
</script></body></html>"""


def _require_admin_user(request: Request) -> str:
    """Allow only the configured admin emails (local/dev = owner, always allowed)."""
    from app.config import settings
    if not settings.use_supabase:
        return "local-owner"
    email = (_get_user_email(request) or "").lower()
    if not email or email not in settings.admin_emails_list:
        raise HTTPException(status_code=403, detail="Admin access only.")
    return email


def _scalar(v) -> int:
    return int(v if not isinstance(v, (list, tuple)) else v[0])


@app.get("/api/admin/metrics")
def admin_metrics(request: Request) -> dict:
    """Aggregated KPIs for the owner dashboard. Admin-only."""
    _require_admin_user(request)
    from datetime import timedelta
    from app.db.models import (UserProfile, Application, UserSubscription,
                               TrialGrant, PlanTier, PLAN_PRICES)
    cutoff = _dt.utcnow() - timedelta(days=7)
    with get_session() as session:
        total_users = _scalar(session.exec(select(func.count(UserProfile.id))).one())
        active_rows = session.exec(
            select(Application.user_id).where(Application.updated_at >= cutoff)).all()
        active_users = len({u for u in active_rows if u})
        referred = _scalar(session.exec(
            select(func.count(UserProfile.id)).where(UserProfile.referred_by_id.isnot(None))).one())
        trials = _scalar(session.exec(select(func.count(TrialGrant.id))).one())
        total_apps = _scalar(session.exec(select(func.count(Application.id))).one())
        subs = session.exec(select(UserSubscription)).all()
    mrr, paid = 0, 0
    now = _dt.utcnow()
    by_plan = {}
    for s in subs:
        if s.plan and s.plan != PlanTier.FREE and (s.current_period_end is None or s.current_period_end > now):
            mrr += PLAN_PRICES.get(s.plan, 0)
            paid += 1
            by_plan[s.plan.value] = by_plan.get(s.plan.value, 0) + 1
    return {
        "total_users": total_users,
        "active_users_7d": active_users,
        "referred_signups": referred,
        "trial_users": trials,
        "total_applications": total_apps,
        "paid_subscriptions": paid,
        "mrr_usd": mrr,
        "arr_usd": mrr * 12,
        "by_plan": by_plan,
    }


@app.get("/api/admin/referrals")
def admin_referrals(request: Request) -> dict:
    """Top referrers + recent referred signups. Admin-only."""
    _require_admin_user(request)
    from app.db.models import UserProfile, UserReferralReward
    with get_session() as session:
        rows = session.exec(
            select(UserProfile.referred_by_id, func.count(UserProfile.id))
            .where(UserProfile.referred_by_id.isnot(None))
            .group_by(UserProfile.referred_by_id)
        ).all()
        rewarded = {r.user_id for r in session.exec(select(UserReferralReward)).all()}
        top = []
        for ref_uid, cnt in sorted(rows, key=lambda x: x[1], reverse=True)[:30]:
            prof = session.exec(select(UserProfile).where(UserProfile.user_id == ref_uid)).first()
            top.append({
                "user_id": ref_uid,
                "name": (f"{prof.first_name} {prof.last_name}".strip() if prof else "") or "(no name)",
                "email": (prof.email if prof else "") or "",
                "code": prof.referral_code if prof else None,
                "count": int(cnt),
                "rewarded": ref_uid in rewarded,
            })
    return {"top_referrers": top}


@app.get("/api/admin/whoami")
def admin_whoami(request: Request) -> dict:
    """Cheap check so the dashboard can reveal the Admin button only to admins."""
    from app.config import settings
    if not settings.use_supabase:
        return {"is_admin": True}
    email = (_get_user_email(request) or "").lower()
    return {"is_admin": bool(email and email in settings.admin_emails_list)}


# --- User Reviews APIs ---

@app.post("/api/reviews")
def submit_user_review(request: Request, payload: dict) -> dict:
    """Submit a review/feedback from a logged-in candidate."""
    uid = _require_user(request)
    
    from app.db.models import UserProfile, UserReview
    name = "Anonymous"
    with get_session() as session:
        profile = session.exec(
            select(UserProfile).where(UserProfile.user_id == uid)
        ).first()
        if profile and profile.first_name:
            name = profile.first_name
            
    rating = int(payload.get("rating", 5))
    content = str(payload.get("content", "")).strip()
    
    if not content:
        raise HTTPException(status_code=400, detail="Review content cannot be empty")
        
    if rating < 1 or rating > 5:
        raise HTTPException(status_code=400, detail="Rating must be between 1 and 5")
        
    review = UserReview(
        user_id=uid,
        user_name=name,
        rating=rating,
        content=content,
        is_public=True,       # Auto-publish by default
        is_featured=False,    # Admin can feature it later
    )
    
    with get_session() as session:
        session.add(review)
        session.commit()
        session.refresh(review)
        
    return {"status": "success", "review_id": review.id}


@app.get("/api/reviews")
def get_public_reviews() -> dict:
    """Get public featured reviews to show on the landing page."""
    from app.db.models import UserReview
    with get_session() as session:
        reviews = session.exec(
            select(UserReview)
            .where(UserReview.is_public == True)
            .where(UserReview.is_featured == True)
            .order_by(desc(UserReview.created_at))
            .limit(3)
        ).all()
        
        if len(reviews) < 3:
            needed = 3 - len(reviews)
            existing_ids = {r.id for r in reviews}
            fallback = session.exec(
                select(UserReview)
                .where(UserReview.is_public == True)
                .where(UserReview.id.not_in(existing_ids) if existing_ids else True)
                .order_by(desc(UserReview.created_at))
                .limit(needed)
            ).all()
            reviews.extend(fallback)
            
    return {
        "reviews": [
            {
                "id": r.id,
                "user_name": r.user_name,
                "rating": r.rating,
                "content": r.content,
                "created_at": r.created_at.isoformat() if r.created_at else None
            }
            for r in reviews
        ]
    }


# --- Admin Reviews Moderation APIs ---

@app.get("/api/admin/reviews")
def admin_list_reviews(request: Request) -> dict:
    """List all reviews for moderation. Admin only."""
    _require_admin_user(request)
    from app.db.models import UserReview
    with get_session() as session:
        reviews = session.exec(
            select(UserReview).order_by(desc(UserReview.created_at))
        ).all()
    return {
        "reviews": [
            {
                "id": r.id,
                "user_id": r.user_id,
                "user_name": r.user_name,
                "rating": r.rating,
                "content": r.content,
                "is_public": r.is_public,
                "is_featured": r.is_featured,
                "created_at": r.created_at.isoformat() if r.created_at else None
            }
            for r in reviews
        ]
    }


@app.post("/api/admin/reviews/{review_id}/approve")
def admin_approve_review(request: Request, review_id: int, payload: dict) -> dict:
    """Approve/Publish or Unpublish a review. Admin only."""
    _require_admin_user(request)
    from app.db.models import UserReview
    approve = bool(payload.get("approve", True))
    with get_session() as session:
        review = session.get(UserReview, review_id)
        if not review:
            raise HTTPException(status_code=404, detail="Review not found")
        review.is_public = approve
        session.add(review)
        session.commit()
    return {"status": "success", "is_public": approve}


@app.post("/api/admin/reviews/{review_id}/feature")
def admin_feature_review(request: Request, review_id: int, payload: dict) -> dict:
    """Toggle featuring a review on the landing page. Admin only."""
    _require_admin_user(request)
    from app.db.models import UserReview
    feature = bool(payload.get("feature", True))
    with get_session() as session:
        review = session.get(UserReview, review_id)
        if not review:
            raise HTTPException(status_code=404, detail="Review not found")
        review.is_featured = feature
        session.add(review)
        session.commit()
    return {"status": "success", "is_featured": feature}


@app.get("/admin", response_class=HTMLResponse)
def admin_page(request: Request):
    """Owner dashboard shell — data is fetched client-side from the gated APIs."""
    return HTMLResponse(_ADMIN_HTML)


def _get_user_plan(uid: str) -> PlanTier:
    """Return the user's current plan tier. Always return PRO to allow everyone to use all features."""
    return PlanTier.PRO


def _get_or_create_usage(session, uid: str):
    from datetime import date, timedelta
    today = date.today()
    week_start = today - timedelta(days=today.weekday())
    row = session.exec(
        select(UserUsage).where(
            UserUsage.user_id == uid,
            UserUsage.usage_date == today,
        )
    ).first()
    if not row:
        row = UserUsage(user_id=uid, usage_date=today, week_start=week_start)
        session.add(row)
        session.flush()
    return row


def _get_week_autofill_count(session, uid: str) -> int:
    """Sum autofill_count_week across all rows in the current Mon–Sun window."""
    from datetime import date, timedelta
    today = date.today()
    week_start = today - timedelta(days=today.weekday())
    rows = session.exec(
        select(UserUsage).where(
            UserUsage.user_id == uid,
            UserUsage.week_start == week_start,
        )
    ).all()
    return sum(r.autofill_count_week for r in rows)


def _check_tailor_limit(uid: str) -> tuple[bool, str, dict]:
    """Returns (allowed, detail_msg, usage_info)."""
    if uid == "local":
        return True, "", {}
    # Founding-user trial: a hard budget of N fully processed jobs, but no
    # daily/weekly caps until it's used up.
    _trial = _get_trial(uid)
    if _trial is not None:
        if _trial.jobs_used >= _trial.jobs_quota:
            return False, (
                f"Your {_trial.jobs_quota}-job founding trial is complete "
                f"({_trial.jobs_used}/{_trial.jobs_quota}). Upgrade to keep applying."
            ), {"trial": True, "used": _trial.jobs_used, "quota": _trial.jobs_quota}
        return True, "", {"trial": True, "used": _trial.jobs_used, "quota": _trial.jobs_quota}
    plan = _get_user_plan(uid)
    limits = PLAN_LIMITS[plan]
    daily_limit = limits["tailor_daily"]
    if daily_limit is None:
        return True, "", {"plan": plan, "daily_limit": None}
    with get_session() as session:
        row = _get_or_create_usage(session, uid)
        used = row.tailor_count
        session.commit()
    if used >= daily_limit:
        upgrade = "Basic ($19/mo)" if plan == PlanTier.FREE else "Pro ($49/mo)"
        return False, (
            f"Daily tailoring limit reached ({used}/{daily_limit}). "
            f"Resets at midnight UTC. Upgrade to {upgrade} for more."
        ), {"plan": plan, "used": used, "daily_limit": daily_limit}
    return True, "", {"plan": plan, "used": used, "daily_limit": daily_limit}


def _increment_tailor(uid: str):
    if uid == "local":
        return
    # Trial users spend a job from their founding-trial budget (1 tailor = 1 job).
    if _get_trial(uid) is not None:
        _increment_trial(uid, 1)
        return
    with get_session() as session:
        row = _get_or_create_usage(session, uid)
        row.tailor_count += 1
        session.add(row)
        session.commit()


def _check_autofill_limit(uid: str) -> tuple[bool, str, dict]:
    if uid == "local":
        return True, "", {}
    plan = _get_user_plan(uid)
    limits = PLAN_LIMITS[plan]
    weekly_limit = limits["autofill_weekly"]
    if weekly_limit is None:
        return True, "", {"plan": plan, "weekly_limit": None}
    if weekly_limit == 0:
        return False, (
            "Auto-fill is not available on the Free plan. "
            "Upgrade to Basic ($19/mo) to get 10 auto-fills per week."
        ), {"plan": plan, "weekly_limit": 0}
    with get_session() as session:
        used = _get_week_autofill_count(session, uid)
    if used >= weekly_limit:
        upgrade = "Pro ($49/mo)" if plan == PlanTier.BASIC else "Agency ($99/mo)"
        return False, (
            f"Weekly auto-fill limit reached ({used}/{weekly_limit}). "
            f"Resets Monday midnight UTC. Upgrade to {upgrade} for more."
        ), {"plan": plan, "used": used, "weekly_limit": weekly_limit}
    return True, "", {"plan": plan, "used": used, "weekly_limit": weekly_limit}


def _increment_autofill(uid: str):
    if uid == "local":
        return
    with get_session() as session:
        row = _get_or_create_usage(session, uid)
        row.autofill_count_week += 1
        session.add(row)
        session.commit()


@app.get("/api/usage")
def get_usage(request: Request) -> dict:
    """Return current plan + usage counters for the dashboard meter."""
    from datetime import date, timedelta
    from app.config import settings
    uid = _get_user_id(request)
    # Fail closed: an expired/invalid session resolves to uid=None in SaaS mode;
    # return a clean 401 instead of crashing on a NOT NULL user_id constraint.
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if uid == "local":
        return {"plan": "local", "tailor_used": 0, "tailor_daily_limit": None,
                "autofill_used_week": 0, "autofill_weekly_limit": None, "trial": None}
    # Founding-user trial: grant on first sight (first N users), then report it.
    trial_grant = _grant_trial_if_eligible(uid)
    plan = _get_user_plan(uid)
    limits = PLAN_LIMITS[plan]
    today = date.today()
    week_start = today - timedelta(days=today.weekday())
    with get_session() as session:
        row = _get_or_create_usage(session, uid)
        tailor_used = row.tailor_count
        autofill_used = _get_week_autofill_count(session, uid)
        session.commit()
    trial = None
    if trial_grant is not None:
        trial = {
            "jobs_used": trial_grant.jobs_used,
            "jobs_quota": trial_grant.jobs_quota,
            "remaining": max(0, trial_grant.jobs_quota - trial_grant.jobs_used),
            "active": trial_grant.jobs_used < trial_grant.jobs_quota,
        }
    return {
        "plan": plan,
        "tailor_used": tailor_used,
        "tailor_daily_limit": limits["tailor_daily"],
        "autofill_used_week": autofill_used,
        "autofill_weekly_limit": limits["autofill_weekly"],
        "week_start": week_start.isoformat(),
        "trial": trial,
    }


@app.post("/run/tailor")
def trigger_tailor(request: Request, bg: BackgroundTasks) -> dict:
    uid = _get_user_id(request)
    bg.add_task(tailor_all_shortlisted, uid if uid != "local" else None)
    return {"started": "tailoring"}


@app.post("/run/tailor/{application_id}")
@_rate_limit("20/minute")
def trigger_tailor_single(application_id: int, request: Request, bg: BackgroundTasks) -> dict:
    """Tailor resume + cover letter for one specific application."""
    uid = _require_owned_application(request, application_id)
    allowed, detail, usage = _check_tailor_limit(uid)
    if not allowed:
        raise HTTPException(status_code=429, detail=detail)
    from app.tailoring.tailor import tailor_for_application
    bg.add_task(tailor_for_application, application_id)
    _increment_tailor(uid)
    return {"started": "tailoring", "application_id": application_id, "usage": usage}


@app.post("/run/autofill/{application_id}")
@_rate_limit("20/minute")
def trigger_autofill(application_id: int, request: Request, bg: BackgroundTasks) -> dict:
    uid = _require_owned_application(request, application_id)
    allowed, detail, usage = _check_autofill_limit(uid)
    if not allowed:
        raise HTTPException(status_code=429, detail=detail)
    bg.add_task(autofill, application_id, bypass_delay=True)
    _increment_autofill(uid)
    return {"started": "autofill", "application_id": application_id, "usage": usage}


@app.post("/run/preview/{application_id}")
def trigger_preview(application_id: int, request: Request, bg: BackgroundTasks) -> dict:
    """Re-open the filled form in a visible Playwright browser for user review."""
    _require_owned_application(request, application_id)
    bg.add_task(preview, application_id)
    return {"started": "preview", "application_id": application_id}


# ── User Profile endpoints ──────────────────────────────────────────────────

# Full UserProfile column set, kept in lock-step with app/db/models.py.
# (column, sqlite_type, postgres_type)
_USERPROFILE_COLUMNS = [
    ("user_id", "VARCHAR", "VARCHAR"),
    ("first_name", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("last_name", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("email", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("phone", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("location", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("linkedin_url", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("github_url", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("portfolio_url", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("work_authorization", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("requires_sponsorship", "BOOLEAN DEFAULT 0", "BOOLEAN DEFAULT FALSE"),
    ("visa_status", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("current_title", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("years_experience", "INTEGER DEFAULT 0", "INTEGER DEFAULT 0"),
    ("salary_min", "INTEGER DEFAULT 0", "INTEGER DEFAULT 0"),
    ("salary_max", "INTEGER DEFAULT 0", "INTEGER DEFAULT 0"),
    ("salary_currency", "VARCHAR DEFAULT 'USD'", "VARCHAR DEFAULT 'USD'"),
    ("degree", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("university", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("graduation_year", "INTEGER", "INTEGER"),
    ("gender", "VARCHAR DEFAULT 'Decline to self-identify'", "VARCHAR DEFAULT 'Decline to self-identify'"),
    ("ethnicity", "VARCHAR DEFAULT 'Decline to self-identify'", "VARCHAR DEFAULT 'Decline to self-identify'"),
    ("veteran_status", "VARCHAR DEFAULT 'I am not a protected veteran'", "VARCHAR DEFAULT 'I am not a protected veteran'"),
    ("disability_status",
     "VARCHAR DEFAULT 'No, I do not have a disability, or history/record of having a disability'",
     "VARCHAR DEFAULT 'No, I do not have a disability, or history/record of having a disability'"),
    ("professional_summary", "TEXT DEFAULT ''", "TEXT DEFAULT ''"),
    ("key_skills", "TEXT DEFAULT ''", "TEXT DEFAULT ''"),
    ("target_roles", "TEXT DEFAULT ''", "TEXT DEFAULT ''"),
    ("job_type_preference", "VARCHAR DEFAULT 'full_time'", "VARCHAR DEFAULT 'full_time'"),
    ("work_auth_status", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("include_internships_in_discovery", "BOOLEAN DEFAULT 0", "BOOLEAN DEFAULT FALSE"),
    ("industry", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("preferred_country", "VARCHAR DEFAULT 'United States'", "VARCHAR DEFAULT 'United States'"),
    ("remote_ok", "BOOLEAN DEFAULT 1", "BOOLEAN DEFAULT TRUE"),
    ("referral_code", "VARCHAR", "VARCHAR"),
    ("referred_by_id", "VARCHAR", "VARCHAR"),
    # Trust Profile (Phase 0)
    ("email_verified", "BOOLEAN DEFAULT 0", "BOOLEAN DEFAULT FALSE"),
    ("phone_verified", "BOOLEAN DEFAULT 0", "BOOLEAN DEFAULT FALSE"),
    ("public_handle", "VARCHAR", "VARCHAR"),
    ("account_type", "VARCHAR DEFAULT 'candidate'", "VARCHAR DEFAULT 'candidate'"),
    ("availability", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("open_to_relocation", "BOOLEAN DEFAULT 0", "BOOLEAN DEFAULT FALSE"),
    ("articulation_video_url", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("articulation_pr", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("trust_identity_score", "INTEGER DEFAULT 0", "INTEGER DEFAULT 0"),
    ("trust_technical_score", "INTEGER DEFAULT 0", "INTEGER DEFAULT 0"),
    ("trust_consistency_score", "INTEGER DEFAULT 0", "INTEGER DEFAULT 0"),
    ("trust_activity_score", "INTEGER DEFAULT 0", "INTEGER DEFAULT 0"),
    ("trust_completeness_score", "INTEGER DEFAULT 0", "INTEGER DEFAULT 0"),
    ("trust_tier", "VARCHAR DEFAULT ''", "VARCHAR DEFAULT ''"),
    ("trust_evidence", "TEXT", "TEXT"),
    ("resume_grounded_ratio", "FLOAT", "DOUBLE PRECISION"),
    ("trust_computed_at", "DATETIME", "TIMESTAMP"),
    ("created_at", "DATETIME", "TIMESTAMP"),
    ("updated_at", "DATETIME", "TIMESTAMP"),
]


def _repair_userprofile_schema() -> None:
    """Idempotently ensure every UserProfile column exists on the live DB.

    Unlike init_db()'s migration helper (which swallows DDL errors and only
    *prints* them), this raises on a genuine failure — so a real problem
    (permissions, connection pooler, etc.) surfaces in the API response/logs
    instead of silently leaving a column missing and looping us back to a 500.

    Postgres uses ``ADD COLUMN IF NOT EXISTS`` (idempotent). SQLite — which has
    no IF NOT EXISTS for columns — is guarded by an inspector lookup.
    """
    from sqlalchemy import text as _text, inspect as _inspect
    from app.db.init_db import engine as _engine
    from app.config import settings as _settings

    if _settings.use_supabase:
        with _engine.begin() as conn:
            for col, _sq, pg in _USERPROFILE_COLUMNS:
                conn.execute(_text(
                    f'ALTER TABLE userprofile ADD COLUMN IF NOT EXISTS {col} {pg}'))
    else:
        insp = _inspect(_engine)
        if not insp.has_table("userprofile"):
            return
        existing = {c["name"].lower() for c in insp.get_columns("userprofile")}
        with _engine.begin() as conn:
            for col, sq, _pg in _USERPROFILE_COLUMNS:
                if col.lower() not in existing:
                    conn.execute(_text(
                        f'ALTER TABLE userprofile ADD COLUMN {col} {sq}'))


@app.get("/api/profile")
def get_profile(request: Request) -> dict:
    """Return the current user profile (seeds from .env on first call)."""
    from app.autofill.answer_pack import _get_or_create_profile
    from app.config import settings
    uid = _get_user_id(request)
    # In multi-tenant mode an unverified/expired token resolves to uid=None.
    # Return a clean 401 (frontend can prompt re-login) instead of a 500.
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        profile = _get_or_create_profile(user_id=uid if uid != "local" else None)
    except Exception as e:
        # Schema drift self-heal: a not-yet-migrated column makes the read fail.
        log.exception("Profile read failed; repairing schema + retrying: %s", e)
        _repair_userprofile_schema()
        profile = _get_or_create_profile(user_id=uid if uid != "local" else None)
    return {
        "id": profile.id,
        "first_name": profile.first_name,
        "last_name": profile.last_name,
        "email": profile.email,
        "phone": profile.phone,
        "location": profile.location,
        "linkedin_url": profile.linkedin_url,
        "github_url": profile.github_url,
        "portfolio_url": profile.portfolio_url,
        "work_authorization": profile.work_authorization,
        "requires_sponsorship": profile.requires_sponsorship,
        "visa_status": profile.visa_status,
        "current_title": profile.current_title,
        "years_experience": profile.years_experience,
        "salary_min": profile.salary_min,
        "salary_max": profile.salary_max,
        "salary_currency": profile.salary_currency,
        "degree": profile.degree,
        "university": profile.university,
        "graduation_year": profile.graduation_year,
        "gender": profile.gender,
        "ethnicity": profile.ethnicity,
        "veteran_status": profile.veteran_status,
        "disability_status": profile.disability_status,
        "professional_summary": profile.professional_summary,
        "key_skills": profile.key_skills,
        "target_roles": profile.target_roles,
        "job_type_preference": getattr(profile, "job_type_preference", "full_time"),
        "work_auth_status": getattr(profile, "work_auth_status", ""),
        "include_internships_in_discovery": getattr(profile, "include_internships_in_discovery", False),
        "industry": getattr(profile, "industry", ""),
        "preferred_country": getattr(profile, "preferred_country", "United States"),
        "remote_ok": getattr(profile, "remote_ok", True),
    }


class ProfileUpdate(BaseModel):
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin_url: Optional[str] = None
    github_url: Optional[str] = None
    portfolio_url: Optional[str] = None
    work_authorization: Optional[str] = None
    requires_sponsorship: Optional[bool] = None
    visa_status: Optional[str] = None
    current_title: Optional[str] = None
    years_experience: Optional[int] = None
    salary_min: Optional[int] = None
    salary_max: Optional[int] = None
    salary_currency: Optional[str] = None
    degree: Optional[str] = None
    university: Optional[str] = None
    graduation_year: Optional[int] = None
    gender: Optional[str] = None
    ethnicity: Optional[str] = None
    veteran_status: Optional[str] = None
    disability_status: Optional[str] = None
    professional_summary: Optional[str] = None
    key_skills: Optional[str] = None
    target_roles: Optional[str] = None
    job_type_preference: Optional[str] = None
    work_auth_status: Optional[str] = None
    include_internships_in_discovery: Optional[bool] = None
    industry: Optional[str] = None
    preferred_country: Optional[str] = None
    remote_ok: Optional[bool] = None
    availability: Optional[str] = None
    open_to_relocation: Optional[bool] = None
    articulation_video_url: Optional[str] = None
    articulation_pr: Optional[str] = None


from datetime import datetime as _dt


@app.put("/api/profile")
def update_profile(request: Request, update: ProfileUpdate) -> dict:
    """Update user profile fields."""
    from app.db.models import UserProfile
    from app.config import settings

    uid = _get_user_id(request)
    # Fail closed: an unauthenticated/expired token in multi-tenant mode would
    # otherwise resolve to user_id=None, and the query below (no WHERE clause)
    # would update the FIRST profile in the table — i.e. some other user's data.
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None

    def _do_update():
        # Get-or-create inside the same session to avoid detached-instance issues
        with get_session() as session:
            q = select(UserProfile)
            if user_id_arg:
                q = q.where(UserProfile.user_id == user_id_arg)
            db_profile = session.exec(q).first()
            if not db_profile:
                db_profile = UserProfile(user_id=user_id_arg)
                session.add(db_profile)
                session.flush()
            for field, value in update.model_dump(exclude_none=True).items():
                setattr(db_profile, field, value)
            db_profile.updated_at = _dt.utcnow()
            session.add(db_profile)
            session.commit()

    try:
        _do_update()
    except Exception as e:
        # Most common cause is schema drift — a model column (e.g. target_roles)
        # that hasn't been migrated onto this database yet, which makes the
        # SELECT * fail. Repair the schema (surfacing any real DDL error rather
        # than swallowing it) and retry once before giving up.
        log.exception("Profile update failed; repairing schema + retrying: %s", e)
        try:
            _repair_userprofile_schema()
            _do_update()
        except Exception as e2:
            log.exception("Profile update failed after schema repair: %s", e2)
            raise HTTPException(status_code=500, detail=f"Could not save profile: {e2}")

    # Recompute the Trust Profile in the background (GitHub harvest can be slow).
    try:
        from app.intelligence.trust_service import compute_and_store
        import threading
        threading.Thread(target=compute_and_store, args=(user_id_arg,), daemon=True).start()
    except Exception as _te:
        log.debug("trust recompute spawn failed: %s", _te)

    return {"success": True}


class RecruiterRegister(BaseModel):
    full_name: Optional[str] = None
    work_email: Optional[str] = None
    company_name: Optional[str] = None
    company_domain: Optional[str] = None
    title: Optional[str] = None
    specialties: Optional[str] = None


def _verify_recruiter(rp) -> None:
    """Auto-verify on corporate-domain match; attach public H-1B filing cred.
    Sets banned if they ever indicate charging candidates (illegal)."""
    import re as _re
    notes = []
    domain = (rp.company_domain or "").lower().strip().lstrip("@")
    email_domain = (rp.work_email or "").split("@")[-1].lower().strip()
    # Strip common free-mail — corporate domain required to verify.
    _free = {"gmail.com", "yahoo.com", "outlook.com", "hotmail.com", "icloud.com", "proton.me"}
    if rp.charges_candidates:
        rp.banned = True
        rp.verified = False
        rp.verification_notes = "Banned: indicated charging candidates (prohibited)."
        return
    if domain and email_domain and email_domain == domain and domain not in _free:
        rp.verified = True
        notes.append(f"Corporate email matches {domain}")
    elif email_domain in _free:
        notes.append("Free email — corporate domain required to verify")
    else:
        notes.append("Email domain does not match company domain")
    # Public H-1B filing history (sponsorship credibility).
    try:
        from app.intelligence.h1b_data import lookup
        rec = lookup(rp.company_name or "")
        if rec:
            rp.h1b_filings = (rec.get("approvals", 0) or 0) + (rec.get("denials", 0) or 0)
            notes.append(f"{rp.h1b_filings} public H-1B filings on record")
    except Exception:
        pass
    rp.verification_notes = " · ".join(notes)


@app.post("/api/recruiter/register")
def recruiter_register(request: Request, body: RecruiterRegister) -> dict:
    """Create/update a demand-side (recruiter/vendor/client) account and verify it."""
    from app.db.models import RecruiterProfile
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    with get_session() as session:
        rp = session.exec(
            select(RecruiterProfile).where(RecruiterProfile.user_id == user_id_arg)
        ).first()
        if not rp:
            rp = RecruiterProfile(user_id=user_id_arg)
        for f, v in body.model_dump(exclude_none=True).items():
            setattr(rp, f, v)
        rp.updated_at = _dt.utcnow()
        _verify_recruiter(rp)
        # Admin override — verify regardless of domain (for testing/demo).
        try:
            is_admin = not settings.use_supabase or (
                (_get_user_email(request) or "").lower() in settings.admin_emails_list)
        except Exception:
            is_admin = False
        if is_admin and not rp.banned:
            rp.verified = True
            rp.verification_notes = (rp.verification_notes + " · admin-verified").strip(" ·")
        session.add(rp)
        # Mark this user's account as recruiter so they're excluded from the
        # candidate pool and routed to the recruiter portal (role separation).
        from app.db.models import UserProfile as _UP
        up = session.exec(select(_UP).where(_UP.user_id == user_id_arg)).first()
        if up:
            up.account_type = "recruiter"
            session.add(up)
        session.commit()
        return {"verified": rp.verified, "banned": rp.banned,
                "h1b_filings": rp.h1b_filings, "notes": rp.verification_notes}


@app.get("/api/account")
def get_account_type(request: Request) -> dict:
    """The signed-in user's role — drives which portal/UI they see."""
    from app.db.models import UserProfile, RecruiterProfile
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        return {"account_type": "candidate", "authenticated": False}
    user_id_arg = uid if uid != "local" else None
    with get_session() as session:
        up = session.exec(select(UserProfile).where(UserProfile.user_id == user_id_arg)).first()
        rp = session.exec(select(RecruiterProfile).where(RecruiterProfile.user_id == user_id_arg)).first()
        atype = (up.account_type if up else "") or "candidate"
        if rp and atype != "recruiter":
            atype = "recruiter"
        return {"account_type": atype, "authenticated": True,
                "recruiter_verified": bool(rp and rp.verified)}


@app.post("/api/recruiter/search")
def recruiter_search(request: Request, body: dict) -> dict:
    """Reverse search — a verified recruiter pastes a job description and gets an
    AI-ranked list of VERIFIED candidates from the pool. The pull model: demand
    finds supply. Candidate contact details are NOT exposed (intro-gated)."""
    from app.db.models import RecruiterProfile, UserProfile
    import json as _json
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    jd = (body.get("job_description") or "").strip()
    if not jd:
        raise HTTPException(status_code=400, detail="job_description required")

    with get_session() as session:
        rp = session.exec(
            select(RecruiterProfile).where(RecruiterProfile.user_id == user_id_arg)
        ).first()
        if settings.use_supabase and (not rp or not rp.verified):
            raise HTTPException(status_code=403, detail="Verified recruiter account required.")
        # Verified candidate pool only (the bait): tier set, candidates (not
        # recruiters), not the searcher themselves.
        candidates = session.exec(
            select(UserProfile).where(
                UserProfile.trust_tier != "",
                UserProfile.account_type != "recruiter",
                UserProfile.user_id != user_id_arg,
            )
        ).all()

    if not candidates:
        return {"results": []}

    # Rank by embedding similarity of the JD vs each candidate's profile text.
    try:
        from app.matching.matcher import _get_embed_model
        import numpy as np
        model = _get_embed_model()
        def _ctext(p):
            return f"{p.current_title}\n{p.key_skills}\n{p.professional_summary}"[:1500]
        texts = [_ctext(p) for p in candidates]
        jd_emb = model.encode([jd], normalize_embeddings=True)[0]
        c_embs = model.encode(texts, normalize_embeddings=True)
        sims = (c_embs @ jd_emb).tolist()
    except Exception as e:
        log.warning("recruiter_search embedding failed: %s", e)
        sims = [0.0] * len(candidates)

    ranked = sorted(zip(candidates, sims), key=lambda x: x[1], reverse=True)[:25]
    results = []
    with get_session() as rsession:
      for p, sim in ranked:
        evidence = {}
        if p.trust_evidence:
            try:
                evidence = _json.loads(p.trust_evidence)
            except (ValueError, TypeError):
                evidence = {}
        results.append({
            "candidate_user_id": p.user_id,
            "handle": p.public_handle,
            "name": f"{p.first_name} {p.last_name}".strip() or "Candidate",
            "title": p.current_title or "",
            "match": round(max(0.0, min(1.0, sim)) * 100),
            "tier": p.trust_tier or "",
            "skills": [s.strip() for s in (p.key_skills or "").split(",") if s.strip()][:8],
            "availability": p.availability or "",
            "work_auth": p.work_authorization or "",
            "needs_sponsorship": bool(p.requires_sponsorship),
            "share_url": f"/u/{p.public_handle}" if p.public_handle else None,
            "rating": _avg_rating(rsession, p.user_id)["avg"],
        })
    return {"results": results}


@app.post("/api/recruiter/intro")
def recruiter_request_intro(request: Request, body: dict) -> dict:
    """Verified recruiter requests an intro to a candidate (candidate must accept
    before contact opens — no resume dump)."""
    from app.db.models import RecruiterProfile, CandidateIntro
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    candidate_id = body.get("candidate_user_id")
    if not candidate_id:
        raise HTTPException(status_code=400, detail="candidate_user_id required")
    with get_session() as session:
        rp = session.exec(
            select(RecruiterProfile).where(RecruiterProfile.user_id == user_id_arg)
        ).first()
        if settings.use_supabase and (not rp or not rp.verified):
            raise HTTPException(status_code=403, detail="Verified recruiter account required.")
        existing = session.exec(
            select(CandidateIntro).where(
                CandidateIntro.recruiter_user_id == user_id_arg,
                CandidateIntro.candidate_user_id == candidate_id,
            )
        ).first()
        if existing:
            return {"status": existing.status, "duplicate": True}
        intro = CandidateIntro(
            recruiter_user_id=user_id_arg, candidate_user_id=candidate_id,
            job_context=(body.get("job_context") or "")[:300],
        )
        session.add(intro)
        session.commit()
        return {"status": "requested"}


@app.get("/api/intros")
def list_intros(request: Request) -> dict:
    """Candidate's incoming intro requests from verified recruiters."""
    from app.db.models import CandidateIntro, RecruiterProfile
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    out = []
    with get_session() as session:
        intros = session.exec(
            select(CandidateIntro).where(CandidateIntro.candidate_user_id == user_id_arg)
            .order_by(CandidateIntro.created_at.desc())
        ).all()
        for i in intros:
            rp = session.exec(
                select(RecruiterProfile).where(RecruiterProfile.user_id == i.recruiter_user_id)
            ).first()
            rating = _avg_rating(session, i.recruiter_user_id)
            out.append({
                "id": i.id, "status": i.status, "job_context": i.job_context,
                "recruiter": (rp.full_name if rp else "") or "Recruiter",
                "company": (rp.company_name if rp else "") or "",
                "verified": bool(rp and rp.verified),
                "rating": rating["avg"], "rating_count": rating["count"],
            })
    return {"intros": out}


@app.post("/api/intros/{intro_id}/respond")
def respond_intro(intro_id: int, request: Request, accept: bool) -> dict:
    """Candidate accepts/declines an intro request."""
    from app.db.models import CandidateIntro
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    with get_session() as session:
        intro = session.get(CandidateIntro, intro_id)
        if not intro or intro.candidate_user_id != user_id_arg:
            raise HTTPException(status_code=404, detail="Intro not found")
        intro.status = "accepted" if accept else "declined"
        session.add(intro)
        session.commit()
    return {"status": intro.status}


def _avg_rating(session, user_id) -> dict:
    """Average earned rating for a user (None if no ratings yet)."""
    from app.db.models import IntroRating
    if not user_id:
        return {"avg": None, "count": 0}
    rows = session.exec(
        select(IntroRating.stars).where(IntroRating.ratee_user_id == user_id)
    ).all()
    if not rows:
        return {"avg": None, "count": 0}
    return {"avg": round(sum(rows) / len(rows), 1), "count": len(rows)}


def _intro_participant(session, intro_id: int, user_id):
    """Return the intro if user_id is a participant (recruiter or candidate)."""
    from app.db.models import CandidateIntro
    intro = session.get(CandidateIntro, intro_id)
    if not intro:
        return None
    if user_id not in (intro.recruiter_user_id, intro.candidate_user_id):
        return None
    return intro


@app.get("/api/conversations")
def list_conversations(request: Request) -> dict:
    """Accepted intros the user is part of (either role) — the inbox for chat."""
    from app.db.models import CandidateIntro, RecruiterProfile, UserProfile
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    out = []
    with get_session() as session:
        intros = session.exec(
            select(CandidateIntro).where(
                CandidateIntro.status == "accepted",
                (CandidateIntro.recruiter_user_id == user_id_arg)
                | (CandidateIntro.candidate_user_id == user_id_arg),
            ).order_by(CandidateIntro.created_at.desc())
        ).all()
        for i in intros:
            am_recruiter = i.recruiter_user_id == user_id_arg
            other_id = i.candidate_user_id if am_recruiter else i.recruiter_user_id
            if am_recruiter:
                up = session.exec(select(UserProfile).where(UserProfile.user_id == other_id)).first()
                other = f"{up.first_name} {up.last_name}".strip() if up else "Candidate"
            else:
                rp = session.exec(select(RecruiterProfile).where(RecruiterProfile.user_id == other_id)).first()
                other = (rp.full_name or rp.company_name) if rp else "Recruiter"
            out.append({"intro_id": i.id, "with": other or "Contact", "job_context": i.job_context})
    return {"conversations": out}


@app.get("/messages", response_class=HTMLResponse)
def messages_page(request: Request):
    """Role-agnostic chat inbox for accepted intros (candidate & recruiter)."""
    return templates.TemplateResponse(request=request, name="messages.html", context={})


@app.get("/api/intros/{intro_id}/messages")
def list_messages(intro_id: int, request: Request) -> dict:
    """Messages on an accepted intro (participants only)."""
    from app.db.models import IntroMessage
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    with get_session() as session:
        intro = _intro_participant(session, intro_id, user_id_arg)
        if not intro:
            raise HTTPException(status_code=404, detail="Intro not found")
        msgs = session.exec(
            select(IntroMessage).where(IntroMessage.intro_id == intro_id)
            .order_by(IntroMessage.created_at)
        ).all()
        return {"status": intro.status, "messages": [
            {"mine": m.sender_user_id == user_id_arg, "body": m.body,
             "at": m.created_at.isoformat()} for m in msgs]}


@app.post("/api/intros/{intro_id}/messages")
def send_message(intro_id: int, request: Request, body: dict) -> dict:
    """Send a message on an accepted intro (participants only, must be accepted)."""
    from app.db.models import IntroMessage
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    text = (body.get("body") or "").strip()[:2000]
    if not text:
        raise HTTPException(status_code=400, detail="Empty message")
    with get_session() as session:
        intro = _intro_participant(session, intro_id, user_id_arg)
        if not intro:
            raise HTTPException(status_code=404, detail="Intro not found")
        if intro.status != "accepted":
            raise HTTPException(status_code=403, detail="Intro must be accepted before messaging")
        session.add(IntroMessage(intro_id=intro_id, sender_user_id=user_id_arg, body=text))
        session.commit()
    return {"ok": True}


@app.post("/api/intros/{intro_id}/rate")
def rate_intro(intro_id: int, request: Request, body: dict) -> dict:
    """Two-way rating after an interaction (1-5 stars). Participants only."""
    from app.db.models import IntroRating
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None
    stars = int(body.get("stars") or 0)
    if not (1 <= stars <= 5):
        raise HTTPException(status_code=400, detail="stars must be 1-5")
    with get_session() as session:
        intro = _intro_participant(session, intro_id, user_id_arg)
        if not intro:
            raise HTTPException(status_code=404, detail="Intro not found")
        ratee = intro.candidate_user_id if user_id_arg == intro.recruiter_user_id else intro.recruiter_user_id
        existing = session.exec(
            select(IntroRating).where(IntroRating.intro_id == intro_id,
                                      IntroRating.rater_user_id == user_id_arg)
        ).first()
        if existing:
            existing.stars = stars
            existing.note = (body.get("note") or "")[:500]
            session.add(existing)
        else:
            session.add(IntroRating(intro_id=intro_id, rater_user_id=user_id_arg,
                                    ratee_user_id=ratee, stars=stars,
                                    note=(body.get("note") or "")[:500]))
        session.commit()
    return {"ok": True}


@app.post("/api/verify/identity")
def verify_identity(request: Request) -> dict:
    """Sync email/phone verification status from Supabase Auth (email is
    confirmed at signup via the magic link). Sets the flags that feed the
    Identity trust dimension, then recomputes the Trust Profile."""
    from app.db.models import UserProfile
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None

    email_ok, phone_ok = False, False
    if settings.use_supabase:
        auth = request.headers.get("Authorization", "")
        token = auth.split(" ", 1)[1] if auth.startswith("Bearer ") else request.cookies.get("sb_token")
        if token:
            try:
                from app.db.supabase_client import verify_jwt
                payload = verify_jwt(token) or {}
                email_ok = bool(payload.get("email_confirmed"))
                phone_ok = bool(payload.get("phone_confirmed"))
            except Exception as e:
                log.debug("identity verify: %s", e)
    else:
        email_ok = True  # local/dev

    with get_session() as session:
        profile = session.exec(
            select(UserProfile).where(UserProfile.user_id == user_id_arg)
        ).first()
        if not profile:
            raise HTTPException(status_code=404, detail="Profile not found")
        if email_ok:
            profile.email_verified = True
        if phone_ok:
            profile.phone_verified = True
        session.add(profile)
        session.commit()

    from app.intelligence.trust_service import compute_and_store
    compute_and_store(user_id_arg)
    return {"email_verified": email_ok, "phone_verified": phone_ok}


@app.get("/api/trust")
def get_trust_profile(request: Request, recompute: bool = False) -> dict:
    """The signed-in user's Trust Profile — five evidence-backed dimensions.

    Pass ?recompute=1 to refresh now (re-harvests GitHub); otherwise returns the
    last stored result. Returns empty dimensions for a brand-new profile.
    """
    import json as _json
    from app.db.models import UserProfile
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None

    if recompute:
        from app.intelligence.trust_service import compute_and_store
        compute_and_store(user_id_arg)

    with get_session() as session:
        profile = session.exec(
            select(UserProfile).where(UserProfile.user_id == user_id_arg)
        ).first()
        if not profile:
            return {"tier": "", "dimensions": [], "computed_at": None, "share_url": None}
        handle = _ensure_public_handle(profile, session)
        evidence = {}
        if profile.trust_evidence:
            try:
                evidence = _json.loads(profile.trust_evidence)
            except (ValueError, TypeError):
                evidence = {}
        base = str(request.base_url).rstrip("/")
        # Momentum — first vs latest snapshot, so candidates see growth.
        from app.db.models import TrustHistory
        hist = session.exec(
            select(TrustHistory).where(TrustHistory.user_id == user_id_arg)
            .order_by(TrustHistory.created_at)
        ).all()
        momentum = None
        if len(hist) >= 2 and hist[-1].overall != hist[0].overall:
            momentum = {"from": hist[0].overall, "to": hist[-1].overall,
                        "delta": hist[-1].overall - hist[0].overall,
                        "since": hist[0].created_at.isoformat()}
        return {
            "tier": profile.trust_tier or "",
            "computed_at": profile.trust_computed_at.isoformat() if profile.trust_computed_at else None,
            "dimensions": [evidence[k] for k in
                           ("identity", "technical", "consistency", "activity", "completeness")
                           if k in evidence],
            "public_handle": handle,
            "share_url": f"{base}/u/{handle}" if handle else None,
            "momentum": momentum,
            "risk_flags": _trust_risk_flags(profile, evidence),
        }


def _ensure_public_handle(profile, session) -> Optional[str]:
    """Mint a stable, unique public handle (hirepath.dev/u/<handle>) once."""
    from app.db.models import UserProfile
    if profile.public_handle:
        return profile.public_handle
    import re as _re, secrets
    base = _re.sub(r"[^a-z0-9]+", "-",
                   f"{profile.first_name}-{profile.last_name}".lower()).strip("-") or "user"
    base = base[:24]
    for _ in range(6):
        cand = f"{base}-{secrets.token_hex(2)}"
        exists = session.exec(
            select(UserProfile).where(UserProfile.public_handle == cand)
        ).first()
        if not exists:
            profile.public_handle = cand
            session.add(profile)
            session.commit()
            return cand
    return None


@app.get("/application/{application_id}/sponsorship")
def application_sponsorship(application_id: int, request: Request) -> dict:
    """Sponsorship Reality for a job — grounded in PUBLIC H-1B/LCA data (not
    self-reported claims). Combines the employer's USCIS filing record with a
    legal, explainable sponsorship assessment of the posting."""
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        job = session.get(Job, application.job_id)
    company = (job.company if job else "") or ""
    out = {"company": company, "h1b": None, "assessment": None}
    try:
        from app.intelligence.h1b_data import lookup
        rec = lookup(company)
        if rec:
            total = (rec.get("approvals", 0) or 0) + (rec.get("denials", 0) or 0)
            out["h1b"] = {
                "approvals": rec.get("approvals", 0),
                "denials": rec.get("denials", 0),
                "approval_rate": round((rec.get("rate", 0) or 0) * 100),
                "year": rec.get("year"),
                "wage_level": rec.get("wage", ""),
                "total_filings": total,
            }
    except Exception as e:
        log.debug("sponsorship h1b lookup failed: %s", e)
    try:
        from app.intelligence.sponsorship import assess
        a = assess(company=company, description=(job.description if job else "") or "",
                   url=(job.url if job else "") or "")
        out["assessment"] = {"badge": a.badge, "reason": a.reason, "tone": a.tone,
                             "cap_exempt": a.cap_exempt}
    except Exception as e:
        log.debug("sponsorship assess failed: %s", e)
    return out


@app.get("/recruiter", response_class=HTMLResponse)
def recruiter_portal(request: Request):
    """Demand-side portal — register/verify + reverse-search the verified pool."""
    return templates.TemplateResponse(request=request, name="recruiter.html", context={})


@app.get("/u/{handle}", response_class=HTMLResponse)
def public_trust_profile(handle: str, request: Request):
    """Public, candidate-owned profile page — evidence-backed, no raw PII.

    Shows the Trust Profile (dimensions + evidence), skills, projects, work-auth
    and availability, but never the email/phone (anti-harvest). Watermarked.
    """
    import json as _json
    from app.db.models import UserProfile
    with get_session() as session:
        profile = session.exec(
            select(UserProfile).where(UserProfile.public_handle == handle)
        ).first()
        if not profile:
            raise HTTPException(status_code=404, detail="Profile not found")
        evidence = {}
        if profile.trust_evidence:
            try:
                evidence = _json.loads(profile.trust_evidence)
            except (ValueError, TypeError):
                evidence = {}
        # City only — never the full street/PII.
        city = (profile.location or "").split(",")[0].strip()
        skills = [s.strip() for s in (profile.key_skills or "").split(",") if s.strip()][:18]
        ctx = {
            "request": request,
            "name": f"{profile.first_name} {profile.last_name}".strip() or "Candidate",
            "title": profile.current_title or "",
            "city": city,
            "summary": profile.professional_summary or "",
            "tier": profile.trust_tier or "",
            "dimensions": [evidence[k] for k in
                           ("identity", "technical", "consistency", "activity", "completeness")
                           if k in evidence],
            "skills": skills,
            "years_experience": profile.years_experience or 0,
            "work_auth": profile.work_authorization or "",
            "github_url": profile.github_url or "",
            "linkedin_url": profile.linkedin_url or "",
            "portfolio_url": profile.portfolio_url or "",
            "handle": handle,
            # Work Readiness Passport — the recruiter "can they start?" answers
            "availability": profile.availability or "",
            "requires_sponsorship": bool(profile.requires_sponsorship),
            "remote_ok": bool(profile.remote_ok),
            "open_to_relocation": bool(profile.open_to_relocation),
            "preferred_country": profile.preferred_country or "",
            "salary_min": profile.salary_min or 0,
            "salary_max": profile.salary_max or 0,
            "articulation_video_url": profile.articulation_video_url or "",
            "risk_flags": _trust_risk_flags(profile, evidence),
        }
    ctx.pop("request", None)
    return templates.TemplateResponse(request=request, name="public_profile.html", context=ctx)


def _trust_risk_flags(profile, evidence: dict) -> list:
    """Recruiter-facing risk summary — derived from existing signals, each with a
    clear status so a recruiter can scan concerns in seconds."""
    flags = []
    def add(label, concern, note=""):
        flags.append({"label": label, "concern": bool(concern), "note": note})
    cons = (evidence.get("consistency") or {}).get("score", 0)
    add("Resume consistency", cons and cons < 60, "Some claims not yet grounded" if cons and cons < 60 else "Claims grounded")
    act = (evidence.get("activity") or {}).get("score", 0)
    add("Open-source activity", act == 0, "No recent public activity" if act == 0 else "Recently active")
    add("Portfolio", not (profile.portfolio_url or "").strip(), "No portfolio linked" if not (profile.portfolio_url or "").strip() else "Portfolio provided")
    add("Identity", not profile.email_verified, "Email not verified" if not profile.email_verified else "Identity verified")
    return flags


# ── Target Roles endpoints ──────────────────────────────────────────────────
# Target Roles are the job titles we search & rank for. They live separately
# from the profile so a user can target roles different from their current one,
# and so discovery can be gated on "roles set + resume uploaded".

# A small curated pool used to seed suggestions when we can't infer much.
_ROLE_SUGGESTION_POOL = [
    "Software Engineer", "Senior Software Engineer", "Backend Engineer",
    "Frontend Engineer", "Full Stack Engineer", "Machine Learning Engineer",
    "AI Engineer", "Data Scientist", "Data Engineer", "Data Analyst",
    "DevOps Engineer", "Platform Engineer", "Product Manager",
    "Engineering Manager", "Cloud Engineer", "Security Engineer",
    "Mobile Engineer", "QA Engineer", "Site Reliability Engineer",
]

# Department / industry → role pool, so suggestions + discovery work for
# non-CS fields. Keyed by lowercase signals found in degree/industry/title.
_DEPARTMENT_ROLES = {
    "civil": ["Civil Engineer", "Structural Engineer", "Project Engineer",
              "Transportation Engineer", "Geotechnical Engineer", "Construction Engineer"],
    "mechanical": ["Mechanical Engineer", "Design Engineer", "Manufacturing Engineer",
                   "HVAC Engineer", "Mechanical Design Engineer", "Product Engineer"],
    "aerospace": ["Aerospace Engineer", "Mechanical Engineer", "Systems Engineer",
                  "Propulsion Engineer", "Flight Test Engineer"],
    "electrical": ["Electrical Engineer", "Power Systems Engineer", "Controls Engineer",
                   "Hardware Engineer", "Electronics Engineer", "Embedded Engineer"],
    "chemical": ["Chemical Engineer", "Process Engineer", "Process Safety Engineer",
                 "Manufacturing Engineer", "Production Engineer"],
    "biomedical": ["Biomedical Engineer", "R&D Engineer", "Clinical Engineer",
                   "Quality Engineer", "Validation Engineer"],
    "industrial": ["Industrial Engineer", "Process Engineer", "Operations Engineer",
                   "Supply Chain Analyst", "Manufacturing Engineer"],
    "environmental": ["Environmental Engineer", "Civil Engineer", "Sustainability Analyst",
                      "Water Resources Engineer"],
    "finance": ["Financial Analyst", "Investment Analyst", "Financial Planning Analyst",
                "Corporate Finance Analyst"],
    "accounting": ["Accountant", "Staff Accountant", "Financial Analyst", "Auditor"],
    "marketing": ["Marketing Analyst", "Digital Marketing Specialist", "Brand Manager",
                  "Marketing Coordinator"],
    "data": ["Data Analyst", "Data Scientist", "Business Analyst", "Data Engineer"],
    "biology": ["Research Associate", "Lab Technician", "Scientist", "Research Scientist"],
    "chemistry": ["Chemist", "Research Associate", "Analytical Chemist", "QC Chemist"],
}

_DEPARTMENT_SIGNALS = {
    "civil": ("civil",),
    "mechanical": ("mechanical", "mech eng"),
    "aerospace": ("aerospace", "aeronautic", "aero eng"),
    "electrical": ("electrical", "electronic", "power systems", "ece"),
    "chemical": ("chemical eng", "chem eng", "chemical engineering"),
    "biomedical": ("biomedical", "bioengineer", "bme"),
    "industrial": ("industrial eng", "industrial engineering", "operations research"),
    "environmental": ("environmental eng", "environmental engineering"),
    "finance": ("finance", "financial"),
    "accounting": ("accounting", "accountant"),
    "marketing": ("marketing",),
    "data": ("data science", "statistics", "analytics"),
    "biology": ("biology", "biological", "biotech"),
    "chemistry": ("chemistry", "chemist"),
}


def _detect_department(profile) -> str | None:
    """Detect the user's department/field from industry + degree + current title."""
    if not profile:
        return None
    blob = " ".join([
        (getattr(profile, "industry", "") or ""),
        (getattr(profile, "degree", "") or ""),
        (getattr(profile, "current_title", "") or ""),
    ]).lower()
    if not blob.strip():
        return None
    for dept, signals in _DEPARTMENT_SIGNALS.items():
        if any(s in blob for s in signals):
            return dept
    return None


def _department_keywords(profile) -> list[str]:
    """Discovery keyword fallback adapted to the user's department (empty if CS/unknown)."""
    dept = _detect_department(profile)
    return list(_DEPARTMENT_ROLES.get(dept, [])) if dept else []


def _suggest_roles(profile, limit: int = 6) -> list[str]:
    """Suggest target roles from the user's department + current title + skills,
    padded with sensible defaults. Pure string work — no LLM, instant + free."""
    out: list[str] = []
    seen: set[str] = set()

    def _add(r: str):
        r = (r or "").strip()
        if r and r.lower() not in seen:
            seen.add(r.lower())
            out.append(r)

    if profile:
        _add(profile.current_title)
        # Skills that read like roles (contain "engineer/scientist/developer/...")
        role_words = ("engineer", "scientist", "developer", "manager", "analyst", "designer")
        for s in (profile.key_skills or "").split(","):
            s = s.strip()
            if s and any(w in s.lower() for w in role_words):
                _add(s)

    # Department-specific pool first (non-CS fields), then the generic pool.
    pool = _department_keywords(profile) + _ROLE_SUGGESTION_POOL
    for r in pool:
        if len(out) >= limit:
            break
        _add(r)
    return out[:limit]


@app.get("/api/target-roles")
def get_target_roles(request: Request) -> dict:
    """Return the user's saved target roles + suggestions to pick from."""
    from app.autofill.answer_pack import _get_or_create_profile
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    profile = _get_or_create_profile(user_id=uid if uid != "local" else None)
    roles = [r.strip() for r in (profile.target_roles or "").split(",") if r.strip()]
    return {
        "roles": roles,
        "suggestions": [s for s in _suggest_roles(profile) if s not in roles],
        "has_resume": _user_has_resume(uid),
    }


class TargetRolesUpdate(BaseModel):
    roles: list[str]


@app.put("/api/target-roles")
def update_target_roles(request: Request, body: TargetRolesUpdate) -> dict:
    """Save the user's target roles (deduped, trimmed, max 12)."""
    from app.db.models import UserProfile
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid != "local" else None

    # Clean + dedupe (case-insensitive), preserve order, cap the list.
    cleaned: list[str] = []
    seen: set[str] = set()
    for r in (body.roles or []):
        r = (r or "").strip()
        if r and r.lower() not in seen:
            seen.add(r.lower())
            cleaned.append(r)
    cleaned = cleaned[:12]

    with get_session() as session:
        q = select(UserProfile)
        if user_id_arg:
            q = q.where(UserProfile.user_id == user_id_arg)
        db_profile = session.exec(q).first()
        if not db_profile:
            db_profile = UserProfile(user_id=user_id_arg)
            session.add(db_profile)
            session.flush()
        db_profile.target_roles = ", ".join(cleaned)
        db_profile.updated_at = _dt.utcnow()
        session.add(db_profile)
        session.commit()
    return {"success": True, "roles": cleaned}


@app.get("/api/profile/memory")
def get_profile_memory(request: Request) -> dict:
    """Latest weekly recruiter brief(s) for the current user (own data only)."""
    from app.config import settings
    from app.db.models import UserPersonalMemory
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid and uid != "local" else None
    with get_session() as session:
        q = select(UserPersonalMemory)
        if user_id_arg:
            q = q.where(UserPersonalMemory.user_id == user_id_arg)
        else:
            q = q.where(UserPersonalMemory.user_id == None)  # noqa: E711
        q = q.order_by(UserPersonalMemory.created_at.desc())
        rows = session.exec(q).all()[:5]
    return {
        "entries": [
            {"id": r.id, "source": r.source, "created_at": r.created_at.isoformat(),
             "recommendations": r.recommendations, "parsed_updates": r.parsed_updates}
            for r in rows
        ]
    }


@app.delete("/api/profile/memory/{entry_id}")
def delete_profile_memory(entry_id: int, request: Request) -> dict:
    """Delete one of the current user's OWN recruiter-memory entries."""
    from app.config import settings
    from app.db.models import UserPersonalMemory
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id_arg = uid if uid and uid != "local" else None
    with get_session() as session:
        row = session.get(UserPersonalMemory, entry_id)
        if not row:
            raise HTTPException(status_code=404, detail="Entry not found")
        # Ownership check — never delete another tenant's memory.
        if row.user_id != user_id_arg:
            raise HTTPException(status_code=403, detail="Not allowed")
        session.delete(row)
        session.commit()
    return {"ok": True}


@app.post("/api/profile/memory/refresh")
@_rate_limit("3/minute")
def refresh_profile_memory(request: Request) -> dict:
    """Trigger a GitHub harvest + recruiter brief now (own data only)."""
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    from app.intelligence.harvester import run_harvest
    try:
        return run_harvest(user_id=uid if uid and uid != "local" else None, notify=False)
    except Exception as e:
        log.exception("Profile memory refresh failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


# ── Admin: one-off H-1B CSV upload (browser-based; gated by ADMIN_TOKEN) ──────
_H1B_UPLOAD_HTML = """<!doctype html><html><head><meta charset=utf-8>
<meta name=viewport content="width=device-width,initial-scale=1">
<title>H-1B Data Upload</title><style>
body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;background:#F4F1EA;color:#2E2A24;
display:flex;min-height:100vh;align-items:center;justify-content:center;margin:0}
.card{background:#FCFAF5;border:1px solid #E6E0D3;border-radius:20px;padding:28px;max-width:460px;width:92%;
box-shadow:0 12px 40px rgba(46,42,36,.10)}
h1{font-size:18px;margin:0 0 6px}p{font-size:13px;color:#8C857A;line-height:1.5}
label{font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:.05em;color:#8C857A;display:block;margin:14px 0 4px}
input{width:100%;box-sizing:border-box;padding:10px;border:1px solid #E6E0D3;border-radius:10px;font-size:14px;background:#fff}
button{margin-top:18px;width:100%;padding:12px;border:0;border-radius:9999px;font-weight:700;color:#fff;
background:linear-gradient(135deg,#2FB4A6,#1F9C8F);cursor:pointer;font-size:14px}
button:disabled{opacity:.6}#msg{margin-top:14px;font-size:13px}
</style></head><body><div class=card>
<h1>🛂 H-1B Employer Data Upload</h1>
<p>Pick the USCIS H-1B Employer Data Hub CSV from your computer. It loads the public
approval data so JobAgent can show real sponsorship numbers. One-time (re-run yearly).</p>
<label>Admin token</label><input id=token type=password placeholder="ADMIN_TOKEN value">
<label>CSV file</label><input id=file type=file accept=".csv">
<button id=go onclick=up()>Upload &amp; ingest</button>
<div id=msg></div>
<hr style="border:0;border-top:1px solid #E6E0D3;margin:22px 0 4px">
<label>Verify a company</label>
<input id=lkco placeholder="e.g. Google, Stripe, Deloitte…" onkeydown="if(event.key==='Enter')lk()">
<button onclick=lk() style="background:#EFEADF;color:#2E2A24;margin-top:10px">Look up H-1B record</button>
<div id=lkmsg style="margin-top:10px;font-size:13px"></div></div><script>
async function lk(){const t=document.getElementById('token').value;const c=document.getElementById('lkco').value.trim();
const m=document.getElementById('lkmsg');if(!t||!c){m.textContent='Enter token + company.';return;}m.textContent='…';
function _esc(s){return (s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;').replace(/'/g,'&#39;');}
try{const r=await fetch('/api/admin/h1b-lookup?token='+encodeURIComponent(t)+'&company='+encodeURIComponent(c));
const d=await r.json();if(!r.ok){m.textContent='❌ '+(d.detail||'error');return;}
if(d.record){m.innerHTML='✅ <b>'+_esc(d.record.name||c)+'</b><br>'+d.record.approvals+' approvals · '+d.record.denials+' denials · '+Math.round((d.record.rate||0)*100)+'% rate (FY'+d.record.year+')';}
else{m.innerHTML='⚠️ No H-1B record for "'+_esc(c)+'" (normalized: <code>'+_esc(d.normalized)+'</code>). Not all employers sponsor.';}}
catch(e){m.textContent='❌ '+e;}}
async function poll(t){try{const r=await fetch('/api/admin/h1b-status?token='+encodeURIComponent(t));
if(r.ok){const d=await r.json();const m=document.getElementById('msg');
function _esc(s){return (s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;').replace(/'/g,'&#39;');}
if(d.last_error){m.innerHTML='<b style=color:#b91c1c>❌ Ingest error:</b> '+_esc(d.last_error)+
'<br><span style=color:#8C857A;font-size:11px>Columns found: '+((d.headers||[]).map(_esc).join(', ')||'none')+'</span>';}
else{m.textContent='✅ Employers in database: '+d.employers+(d.last_rows?(' ('+d.last_rows+' rows loaded)'):'');}}}catch(e){}}
async function up(){const t=document.getElementById('token').value;const f=document.getElementById('file').files[0];
const m=document.getElementById('msg');const b=document.getElementById('go');
if(!t||!f){m.textContent='Enter the token and choose a CSV.';return;}
b.disabled=true;b.textContent='Uploading…';m.textContent='';
const fd=new FormData();fd.append('token',t);fd.append('file',f);
try{const r=await fetch('/api/admin/h1b-upload',{method:'POST',body:fd});
const d=await r.json();
if(r.ok){m.textContent='⏳ '+(d.note||'Ingesting…')+' ('+Math.round((d.size_bytes||0)/1048576)+' MB)';
let n=0;const iv=setInterval(()=>{poll(t);if(++n>20)clearInterval(iv);},5000);}
else{m.textContent='❌ '+(d.detail||'Failed');}}
catch(e){m.textContent='❌ '+e;}b.disabled=false;b.textContent='Upload & ingest';}
</script></body></html>"""


def _require_admin(token: str) -> None:
    """Gate admin actions. Distinct messages so the user knows which is wrong."""
    from app.config import settings
    if not (settings.admin_token or "").strip():
        raise HTTPException(
            status_code=403,
            detail="H-1B upload is disabled: ADMIN_TOKEN is not set on the server. "
                   "Add it as an environment variable and redeploy, then reload.",
        )
    if (token or "").strip() != settings.admin_token.strip():
        raise HTTPException(
            status_code=403,
            detail="Wrong admin token — it must exactly match the ADMIN_TOKEN "
                   "environment variable on the server (watch for extra spaces/quotes).",
        )


@app.get("/admin/h1b", response_class=HTMLResponse)
def admin_h1b_page(request: Request):
    """Browser upload page for the USCIS H-1B CSV (gated by ADMIN_TOKEN env)."""
    from app.config import settings
    if not (settings.admin_token or "").strip():
        return HTMLResponse(
            "<body style='font-family:sans-serif;max-width:520px;margin:60px auto;color:#2E2A24'>"
            "<h2>H-1B upload is disabled</h2>"
            "<p>The <code>ADMIN_TOKEN</code> environment variable is not set on the server "
            "(or the server hasn't redeployed since you set it).</p>"
            "<p><b>Fix:</b> add an <code>ADMIN_TOKEN</code> variable on your host → redeploy → reload this page.</p>"
            "</body>",
            status_code=403,
        )
    return HTMLResponse(_H1B_UPLOAD_HTML)


@app.get("/api/admin/h1b-status")
def admin_h1b_status(token: str = "") -> dict:
    from app.db.models import H1BSponsor
    from app.intelligence import h1b_data as _h
    _require_admin(token)
    with get_session() as session:
        count = session.exec(select(func.count(H1BSponsor.id))).one()
    li = _h.LAST_INGEST
    return {
        "employers": int(count if not isinstance(count, (list, tuple)) else count[0]),
        "last_rows": li.get("rows", 0),
        "last_error": li.get("error", ""),
        "headers": li.get("headers", []),
    }


@app.post("/api/admin/h1b-upload")
async def admin_h1b_upload(bg: BackgroundTasks, token: str = Form(""), file: UploadFile = File(...)) -> dict:
    """Accept the USCIS CSV from the browser and ingest it in the background."""
    _require_admin(token)
    import tempfile, os as _os
    data = await file.read()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv")
    tmp.write(data)
    tmp.close()

    def _do():
        try:
            from app.intelligence.h1b_data import ingest_csv
            n = ingest_csv(tmp.name)
            log.info("H-1B upload ingested %d employer-year rows", n)
        except Exception as e:
            log.exception("H-1B ingest failed: %s", e)
            try:
                from app.intelligence import h1b_data as _h
                _h.LAST_INGEST.update(error=str(e))
            except Exception:
                pass
        finally:
            try:
                _os.unlink(tmp.name)
            except Exception:
                pass

    bg.add_task(_do)
    return {"started": True, "size_bytes": len(data),
            "note": "Ingesting in the background — employer count will update below shortly."}


class LinkedInPasteBody(BaseModel):
    text: str


@app.post("/api/profile/memory/linkedin")
@_rate_limit("6/minute")
def ingest_linkedin_paste(request: Request, body: LinkedInPasteBody) -> dict:
    """Legal LinkedIn path — store the user's OWN pasted profile text (no scraping)."""
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    from app.intelligence.harvester import ingest_linkedin_text
    try:
        return ingest_linkedin_text(uid if uid and uid != "local" else None, body.text)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        log.exception("LinkedIn paste failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


def _extract_text_from_upload(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded PDF/DOCX/TXT/MD file.

    Uses the same extractors as the résumé loader (pypdf / python-docx), so a
    LinkedIn 'Save to PDF' export is read the same way a résumé PDF is.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "pdf"
    if ext not in ("pdf", "docx", "txt", "md"):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, TXT, MD allowed")
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")
    try:
        if ext == "pdf":
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        if ext == "docx":
            import io
            from docx import Document
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return data.decode("utf-8", errors="ignore")  # txt / md
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read that file: {e}")


@app.post("/api/profile/memory/linkedin/pdf")
@_rate_limit("6/minute")
async def ingest_linkedin_pdf(request: Request) -> dict:
    """Legal LinkedIn path (PDF): the user uploads their OWN profile exported via
    LinkedIn's 'Save to PDF' (Profile → More → Save to PDF). We extract the text
    and store it — no scraping, no automation. Captures the full profile, unlike
    the on-page extension read which only sees rendered sections."""
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")

    form = await request.form()
    file = form.get("file")
    if not file:
        raise HTTPException(status_code=400, detail="No file provided")

    data = await file.read()
    text = _extract_text_from_upload(file.filename or "profile.pdf", data).strip()
    if len(text) < 40:
        raise HTTPException(
            status_code=400,
            detail="Couldn't read enough text from that file. Use LinkedIn's 'Save to PDF' export.",
        )

    from app.intelligence.harvester import ingest_linkedin_text
    try:
        return ingest_linkedin_text(uid if uid and uid != "local" else None, text)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        log.exception("LinkedIn PDF import failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


def _user_needs_sponsorship(uid) -> bool:
    """True when this user will need visa sponsorship (drives cap-exempt boost)."""
    try:
        from app.autofill.answer_pack import _get_or_create_profile
        from app.intelligence.work_auth import assess_profile
        p = _get_or_create_profile(user_id=uid if uid and uid != "local" else None)
        return bool(assess_profile(p).needs_future_sponsorship)
    except Exception:
        return False


def _get_target_roles(uid) -> list[str]:
    """Load saved target roles for a user as a list (empty if none)."""
    from app.db.models import UserProfile
    user_id_arg = uid if uid and uid != "local" else None
    with get_session() as session:
        q = select(UserProfile)
        if user_id_arg:
            q = q.where(UserProfile.user_id == user_id_arg)
        p = session.exec(q).first()
    if not p or not p.target_roles:
        return []
    return [r.strip() for r in p.target_roles.split(",") if r.strip()]


# ── Profile Avatar endpoints ────────────────────────────────────────────────

def _sign_avatar(bucket, path: str) -> str | None:
    """Return a signed URL for an avatar object, or None if it doesn't exist."""
    try:
        signed = bucket.create_signed_url(path, 3600)
    except Exception:
        return None
    return (signed or {}).get("signedURL") or (signed or {}).get("signedUrl")


@app.get("/api/profile/avatar")
def get_avatar(request: Request) -> dict:
    """Return the signed URL for the user's profile avatar, or null."""
    from app.config import settings
    uid = _get_user_id(request)
    if settings.use_supabase and uid and uid != "local":
        try:
            from app.db.supabase_client import service_client
            bucket = service_client().storage.from_("avatars")
            # Upload always stores {uid}/avatar.<ext> — try known extensions
            # directly. This avoids relying on list(), which can return empty.
            for ext in ("jpg", "jpeg", "png", "webp", "gif"):
                url = _sign_avatar(bucket, f"{uid}/avatar.{ext}")
                if url:
                    return {"url": url}
            # Fallback: list the folder for any unexpected filename.
            try:
                files = bucket.list(uid) or []
                names = [f.get("name", "") for f in files if f.get("name", "").startswith("avatar.")]
                if names:
                    url = _sign_avatar(bucket, f"{uid}/{names[0]}")
                    if url:
                        return {"url": url}
            except Exception:
                pass
        except Exception:
            pass
    return {"url": None}


@app.post("/api/profile/avatar")
async def upload_avatar(request: Request, file: UploadFile = File(...)) -> dict:
    """Upload a profile photo and store in Supabase storage (avatars bucket)."""
    from app.config import settings
    uid = _get_user_id(request)
    ext = (file.filename or "avatar.jpg").rsplit(".", 1)[-1].lower()
    if ext not in {"jpg", "jpeg", "png", "webp", "gif"}:
        raise HTTPException(status_code=400, detail="Unsupported image type")
    content = await file.read()
    if len(content) > 3 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large — max 3 MB")

    if settings.use_supabase and uid and uid != "local":
        try:
            from app.db.supabase_client import service_client
            sb = service_client()
            path = f"{uid}/avatar.{ext}"
            mime = file.content_type or "image/jpeg"
            # Try upsert first; fall back to remove+upload for older supabase-py
            try:
                sb.storage.from_("avatars").upload(path, content, {"content-type": mime, "upsert": "true"})
            except Exception:
                try:
                    sb.storage.from_("avatars").remove([path])
                except Exception:
                    pass
                sb.storage.from_("avatars").upload(path, content, {"content-type": mime})
            signed = sb.storage.from_("avatars").create_signed_url(path, 3600)
            url = (signed or {}).get("signedURL") or (signed or {}).get("signedUrl")
            return {"url": url}
        except Exception as exc:
            log.exception("Avatar upload failed: %s", exc)
            raise HTTPException(status_code=500, detail=f"Upload failed: {exc}")
    return {"url": None}


# ── Answer Pack endpoint ────────────────────────────────────────────────────

@app.get("/application/{application_id}/referral")
@_rate_limit("20/minute")
def get_referral_drafts(application_id: int, request: Request) -> dict:
    """Draft referral / hiring-manager / visa-alumni outreach for one application.
    Drafts only — the user sends them. Strictly scoped to the owning user."""
    _require_owned_application(request, application_id)
    uid = _get_user_id(request)
    from app.intelligence.referral import generate_referral_drafts
    try:
        return generate_referral_drafts(application_id, user_id=uid if uid != "local" else None)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        log.exception("Referral draft generation failed for app %d: %s", application_id, e)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/application/{application_id}/connections")
@_rate_limit("10/minute")
def get_job_connections(application_id: int, request: Request) -> dict:
    """Find public LinkedIn referrers/champions for this job via Google X-Ray
    (SerpAPI — no LinkedIn login or scraping). Owner-scoped + plan-gated."""
    _require_owned_application(request, application_id)
    uid = _get_user_id(request)
    from app.config import settings
    # Plan gate: Pro feature. Free users get an upsell; local/dev is unlimited.
    if settings.use_supabase and uid and uid != "local":
        if _get_user_plan(uid) == PlanTier.FREE:
            raise HTTPException(
                status_code=402,
                detail="Finding LinkedIn referrers is a Pro feature — upgrade to unlock champion search.",
            )
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        job = session.get(Job, application.job_id)
    needs = _user_needs_sponsorship(uid if uid and uid != "local" else None)
    from app.intelligence.linkedin_xray import find_champions
    return find_champions(job.company or "", job.title or "", visa=needs)


@app.get("/api/admin/h1b-lookup")
def admin_h1b_lookup(company: str = "", token: str = "") -> dict:
    """Verify the ingested H-1B data for a company (admin only)."""
    _require_admin(token)
    from app.intelligence.h1b_data import lookup, normalize
    rec = lookup(company)
    return {"company": company, "normalized": normalize(company), "record": rec}


@app.get("/application/{application_id}/answer-pack")
def get_answer_pack(application_id: int, request: Request) -> dict:
    """Generate (or return cached) answer pack for one application."""
    from app.autofill.answer_pack import generate_answer_pack
    _require_owned_application(request, application_id)
    try:
        return generate_answer_pack(application_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        log.exception("Answer pack generation failed for app %d: %s", application_id, e)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/application/{application_id}/download-resume")
def download_tailored_resume(application_id: int, request: Request):
    """Serve the tailored DOCX resume file as a direct download."""
    _require_owned_application(request, application_id)
    with get_session() as session:
        application = session.get(Application, application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        if not application.tailored_resume_path:
            raise HTTPException(status_code=400, detail="No tailored resume found for this application")
            
        import os
        from fastapi.responses import FileResponse
        path = application.tailored_resume_path
        if not os.path.exists(path):
            raise HTTPException(status_code=404, detail="Resume file not found on disk")
            
        filename = os.path.basename(path)
        return FileResponse(
            path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )


class AskCopilotRequest(BaseModel):
    question: str


@app.post("/application/{application_id}/ask")
async def ask_copilot(application_id: int, req: AskCopilotRequest, request: Request) -> dict:
    """Ask custom question grounded in the JD and resume context."""
    _require_owned_application(request, application_id)
    
    from app.db.init_db import get_session
    from app.db.models import Application, Job
    
    with get_session() as session:
        app = session.get(Application, application_id)
        if not app:
            raise HTTPException(status_code=404, detail="Application not found")
        job = session.get(Job, app.job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
            
        from app.autofill.answer_pack import _load_resume_text_from_path, _get_or_create_profile
        resume_text = _load_resume_text_from_path(app.tailored_resume_path)
        if not resume_text:
            from app.matching.pipeline import _load_resume
            resume_text = _load_resume(user_id=app.user_id)
            
        profile = _get_or_create_profile(user_id=app.user_id)
        
    from app.config import settings
    
    prompt = f"""You are a helpful, professional career assistant and recruiter.
The candidate is applying for the following role:
Company: {job.company}
Title: {job.title}
Job Description:
{job.description[:4000] if job.description else ""}

Candidate Profile:
- Name: {profile.first_name} {profile.last_name}
- Title: {profile.current_title}
- Experience: {profile.years_experience} years
- Summary: {profile.professional_summary[:600] if profile.professional_summary else ""}

Candidate Resume Details:
{resume_text[:6000]}

User Question:
"{req.question}"

Write a concise, professional response to the user's question. 
- If this is a job application essay/form question, write the answer in the first-person ("I") as the candidate, ready to be copy-pasted, matching the candidate's actual background honestly (do not fabricate experience). Keep it under 150 words.
- If this is an interview prep or advice question, give clear, blunt feedback grounded in the candidate's real resume bullets.
Do NOT use markdown headers or introduction/conversational prefix. Return only the response/answer text."""

    system_prompt = "You write clear, professional, and honest answers grounded strictly in the candidate's resume and job description."
    
    answer = ""
    if settings.anthropic_api_key:
        try:
            from anthropic import Anthropic
            client = Anthropic(api_key=settings.anthropic_api_key)
            resp = client.messages.create(
                model=settings.cover_letter_model,
                max_tokens=500,
                system=system_prompt,
                messages=[{"role": "user", "content": prompt}],
            )
            answer = resp.content[0].text.strip()
        except Exception as e:
            log.warning("Ask copilot Anthropic failed: %s", e)
            
    if not answer and settings.openai_api_key:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=settings.openai_api_key)
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=500,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ]
            )
            answer = resp.choices[0].message.content.strip()
        except Exception as e:
            log.warning("Ask copilot OpenAI failed: %s", e)
            
    if not answer:
        raise HTTPException(status_code=500, detail="No LLM backend available to answer the question")
        
    return {"answer": answer}


class ExtractLinkRequest(BaseModel):
    url: str


@app.post("/run/extract-link")
@_rate_limit("10/minute")
async def trigger_extract_link(req: ExtractLinkRequest, request: Request, bg: BackgroundTasks) -> dict:
    from app.discovery.extractor import extract_and_rank_job
    from app.tailoring.tailor import tailor_for_application

    uid = _require_user(request)
    log.info("Extracting manual link: %s", req.url)
    try:
        app_id = await extract_and_rank_job(req.url, user_id=uid if uid != "local" else None)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))

    # Tailor in the background
    bg.add_task(tailor_for_application, app_id)

    return {"success": True, "application_id": app_id}


# ── GDPR / Account deletion ──────────────────────────────────────────────────

@app.delete("/api/account")
def delete_account(request: Request) -> dict:
    """Permanently delete all data for the authenticated user."""
    uid = _require_user(request)
    from app.db.models import (
        UserProfile, PendingQuestion, AnswerMemory, DiscoveryRun,
        UserSubscription, UserUsage,
    )
    from app.db.init_db import get_session
    from sqlmodel import select, delete as sql_delete

    with get_session() as session:
        # Collect application IDs for this user
        if uid != "local":
            app_ids = [
                a.id for a in session.exec(
                    select(Application).where(Application.user_id == uid)
                ).all()
            ]
            # Delete pending questions linked to those applications
            if app_ids:
                session.exec(sql_delete(PendingQuestion).where(PendingQuestion.application_id.in_(app_ids)))
            # Delete every user-scoped row so no orphaned data remains (GDPR).
            session.exec(sql_delete(Application).where(Application.user_id == uid))
            session.exec(sql_delete(Job).where(Job.user_id == uid))
            session.exec(sql_delete(UserProfile).where(UserProfile.user_id == uid))
            session.exec(sql_delete(AnswerMemory).where(AnswerMemory.user_id == uid))
            session.exec(sql_delete(DiscoveryRun).where(DiscoveryRun.user_id == uid))
            session.exec(sql_delete(UserSubscription).where(UserSubscription.user_id == uid))
            session.exec(sql_delete(UserUsage).where(UserUsage.user_id == uid))
            session.commit()

    # Delete resume files from Supabase Storage
    from app.config import settings
    if settings.use_supabase and uid and uid != "local":
        try:
            from app.db.supabase_client import service_client
            sb = service_client()
            files = sb.storage.from_("resume").list(uid) or []
            paths = [f"{uid}/{f['name']}" for f in files if f.get("name")]
            if paths:
                sb.storage.from_("resume").remove(paths)
            # Also delete Supabase Auth user
            sb.auth.admin.delete_user(uid)
        except Exception:
            pass

    return {"success": True, "message": "All account data deleted."}


# ── Email sync (browser extension) ───────────────────────────────────────────

class SyncEmailPayload(BaseModel):
    emails: list  # List of dicts with keys: subject, sender, body, date, company_guess
    source: str = "inbox"  # "gmail" | "outlook"
    day_range: Optional[int] = None


_REJECTION_KEYWORDS = [
    'unfortunately', 'not moving forward', 'other candidates',
    'other applicants', 'not selected', 'regret to inform',
    'we regret', 'decided not to', 'will not be proceeding',
    'not be proceeding', 'position has been filled', 'role has been filled',
    'no longer under consideration', 'not be moving forward',
    "won't be moving forward", 'will not be moving forward',
    'decided to move forward with other', 'decided to proceed with other',
    'pursue other candidates', 'not to move forward',
    'after careful consideration', 'we have chosen', 'were not selected',
    'wish you the best', 'wish you success', 'wish you well',
    'not a match at this time', 'not be advancing', 'will not be advancing',
    'unable to offer', 'not be extending', 'application was unsuccessful',
    'were unsuccessful', 'thank you for your interest, however',
    'not progressing', 'will not progress',
]

# Phrases that strongly indicate a real interview invite / scheduling request.
# Keep these precise — loose words like bare 'interview' or 'next steps' cause
# false positives on acknowledgment auto-replies that mention the job title or
# describe a multi-step review process.
_POSITIVE_KEYWORDS = [
    'invitation to interview',
    'schedule an interview',
    'schedule a call',
    'schedule a time',
    'please schedule',
    'book a time',
    'set up a call',
    'set up an interview',
    'pleased to invite you',
    'invite you to interview',
    'move forward with your candidacy',
    'move forward with your application and',
    'selected for an interview',
    'selected to interview',
    'advance to the interview',
    'advance to the next round',
    'phone screen',
    'recruiter screen',
    'video call',
    'video interview',
    'meet the team',
    'like to speak with you about the role',
    'like to speak with you about this role',
    'love to chat about the role',
    'next round',
    'technical assessment',
    'coding challenge',
    'take-home assessment',
    'technical interview',
    'onsite interview',
    'on-site interview',
    'final round',
]

# Phrases that mark an email as an acknowledgment / under-review auto-reply.
# If any of these appear, we suppress a false-positive interview signal so that
# plain confirmation emails do not flip the card to INTERVIEWING.
_ACKNOWLEDGMENT_KEYWORDS = [
    'thank you for applying',
    'thank you for your application',
    'we have received your application',
    'we received your application',
    'your application has been received',
    'application is currently under review',
    'application is under review',
    'currently under review',
    'profile is currently under review',
    'profile is under review',
    'under consideration',
    'our team will review',
    'our recruiting team will',
    'will be in touch',
    'will reach out',
    'will contact you',
    'will get back to you',
    'if your experience matches',
    'if selected for',
    'if you are selected',
    'reviewing applications',
    'reviewing all applications',
    'application has been submitted',
    'application was submitted',
    'confirmation of your application',
    'application confirmation',
]


@app.post("/api/sync-emails")
@_rate_limit("10/minute")
async def sync_emails(payload: SyncEmailPayload, request: Request, bg: BackgroundTasks) -> dict:
    """Ingest emails from the browser extension, match to applications,
    and auto-detect rejections / interview invitations."""
    from datetime import datetime
    import json as _json

    uid = _require_user(request)
    _uid_filter = uid and uid != "local"

    # Load all user applications joined with their jobs
    with get_session() as session:
        q = select(Application, Job).join(Job)
        if _uid_filter:
            q = q.where(Application.user_id == uid)
        rows = session.exec(q).all()

    matched = 0
    rejections = 0
    interviews = 0
    imported = 0
    unmatched_list: list[dict] = []

    def _email_status(is_rej: bool, is_int: bool) -> "ApplicationStatus":
        if is_rej:
            return ApplicationStatus.REJECTED
        if is_int:
            return ApplicationStatus.INTERVIEWING
        return ApplicationStatus.SUBMITTED

    def _clean_email_title(subj: str, company: str) -> str:
        """Best-effort job title from an email subject line."""
        import re as _re
        t = subj or ""
        # Strip common boilerplate prefixes/suffixes
        t = _re.sub(r"(?i)^\s*(re|fwd?)\s*:\s*", "", t)
        t = _re.sub(r"(?i)\b(thank you for (applying|your application|your interest)( (to|in))?|"
                    r"your application( (for|to))?|application (received|update|status|next steps)|"
                    r"we received your application|application to|application for)\b", "", t)
        if company:
            t = _re.sub(_re.escape(company), "", t, flags=_re.IGNORECASE)
        t = _re.sub(r"[\-–—|@:]+", " ", t)
        t = _re.sub(r"\s{2,}", " ", t).strip(" -–—|,.")
        # If nothing meaningful survived, fall back to the company name.
        if len(t) < 3:
            return f"Application — {company}" if company else "Application (from email)"
        return t[:120]

    for email in payload.emails:
        company_guess = (email.get("company_guess") or "").strip().lower()
        body_lower = (email.get("body") or "").lower()
        subject_lower = (email.get("subject") or "").lower()
        combined_text = subject_lower + " " + body_lower

        # ── Fuzzy match to an existing application ──────────────────────
        matched_app = None
        matched_job = None
        if company_guess:
            for app_row, job_row in rows:
                if company_guess in (job_row.company or "").lower():
                    matched_app = app_row
                    matched_job = job_row
                    break

        # ── Detect rejection / positive signal ──────────────────────────
        is_rejection = any(kw in combined_text for kw in _REJECTION_KEYWORDS)
        # An acknowledgment email ("under review", "we received your application")
        # suppresses a positive signal so confirmation auto-replies don't flip
        # the card to INTERVIEWING incorrectly.
        is_acknowledgment = any(kw in combined_text for kw in _ACKNOWLEDGMENT_KEYWORDS)
        is_interview = (
            not is_acknowledgment
            and any(kw in combined_text for kw in _POSITIVE_KEYWORDS)
        )

        if matched_app and matched_job:
            matched += 1

            if is_rejection:
                rejections += 1
                with get_session() as session:
                    app_obj = session.get(Application, matched_app.id)
                    if app_obj:
                        app_obj.status = ApplicationStatus.REJECTED
                        app_obj.updated_at = datetime.utcnow()
                        app_obj.notes = (
                            (app_obj.notes or "")
                            + f"\nAuto-detected rejection from email on {datetime.utcnow():%Y-%m-%d}."
                        )
                        session.add(app_obj)
                        session.commit()

                # Build resume text for background analysis
                try:
                    from app.autofill.answer_pack import _load_resume_text_from_path
                    resume_text = _load_resume_text_from_path(matched_app.tailored_resume_path)
                except Exception:
                    resume_text = ""
                if not resume_text:
                    try:
                        from app.matching.pipeline import _load_resume
                        resume_text = _load_resume(user_id=uid)
                    except Exception:
                        resume_text = ""

                def _run_analysis(app_id: int, jd: str, resume_md: str, email_body: str):
                    from app.tailoring.analyzer import RejectionAnalyzer
                    import json
                    result = RejectionAnalyzer().analyze(jd, resume_md, email_body)
                    with get_session() as s:
                        a = s.get(Application, app_id)
                        if a:
                            a.rejection_analysis = json.dumps(result)
                            s.add(a)
                            s.commit()

                bg.add_task(
                    _run_analysis,
                    matched_app.id,
                    matched_job.description or "",
                    resume_text,
                    email.get("body") or "",
                )

            elif is_interview:
                interviews += 1
                with get_session() as session:
                    app_obj = session.get(Application, matched_app.id)
                    if app_obj and app_obj.status == ApplicationStatus.SUBMITTED:
                        app_obj.status = ApplicationStatus.INTERVIEWING
                        app_obj.updated_at = datetime.utcnow()
                        app_obj.notes = (
                            (app_obj.notes or "")
                            + f"\nAuto-detected interview signal from email on {datetime.utcnow():%Y-%m-%d}."
                        )
                        session.add(app_obj)
                        session.commit()
        else:
            # ── No existing application matched — import as a tracked one ────
            # These are real job-related emails (rejections, acknowledgements,
            # interview invites) for jobs the user applied to OUTSIDE HirePath.
            # We surface them in the dashboard tracker so nothing is lost.
            raw_company = (email.get("company_guess") or "").strip()
            if not raw_company:
                unmatched_list.append({
                    "subject": email.get("subject"),
                    "sender": email.get("sender"),
                    "company_guess": email.get("company_guess"),
                })
                continue

            import hashlib as _hashlib
            title = _clean_email_title(email.get("subject") or "", raw_company)
            ext_id = "email:" + _hashlib.sha1(
                f"{(uid or 'local')}|{raw_company.lower()}|{title.lower()}".encode()
            ).hexdigest()[:20]

            try:
                with get_session() as session:
                    existing = session.exec(
                        select(Job).where(
                            Job.source == JobSource.MANUAL,
                            Job.external_id == ext_id,
                        ).where(
                            (Job.user_id == uid) if _uid_filter else (Job.user_id.is_(None) | (Job.user_id == "local"))
                        )
                    ).first()

                    if existing:
                        # Already imported — only upgrade status on a stronger signal.
                        app_obj = session.exec(
                            select(Application).where(Application.job_id == existing.id)
                        ).first()
                        if app_obj and is_rejection and app_obj.status != ApplicationStatus.REJECTED:
                            app_obj.status = ApplicationStatus.REJECTED
                            app_obj.updated_at = datetime.utcnow()
                            session.add(app_obj)
                            session.commit()
                        continue

                    job = Job(
                        user_id=uid if _uid_filter else None,
                        source=JobSource.MANUAL,
                        external_id=ext_id,
                        company=raw_company,
                        title=title,
                        url="",
                        description=email.get("body") or "",
                    )
                    session.add(job)
                    session.commit()
                    session.refresh(job)

                    app_obj = Application(
                        user_id=uid if _uid_filter else None,
                        job_id=job.id,
                        status=_email_status(is_rejection, is_interview),
                        apply_track="email_import",
                        submitted_at=datetime.utcnow(),
                        notes=f"Imported from {payload.source} email on {datetime.utcnow():%Y-%m-%d}.",
                    )
                    session.add(app_obj)
                    session.commit()
                imported += 1
                if is_rejection:
                    rejections += 1
                elif is_interview:
                    interviews += 1
            except Exception as e:
                print(f"[sync-emails] failed to import email '{title}': {e}")
                unmatched_list.append({
                    "subject": email.get("subject"),
                    "sender": email.get("sender"),
                    "company_guess": email.get("company_guess"),
                })

    return {
        "success": True,
        "processed": len(payload.emails),
        "matched": matched,
        "imported": imported,
        "rejections": rejections,
        "interviews": interviews,
        "unmatched": len(unmatched_list),
    }


# ── CSV export ───────────────────────────────────────────────────────────────

@app.get("/api/export/applications.csv")
def export_applications_csv(request: Request):
    """Download all applications as a CSV file."""
    import csv, io
    from fastapi.responses import StreamingResponse

    uid = _require_user(request)
    _uid_filter = uid and uid != "local"

    with get_session() as session:
        q = select(Application, Job).join(Job)
        if _uid_filter:
            q = q.where(Application.user_id == uid)
        results = session.exec(q).all()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "id", "company", "title", "location", "remote", "source",
        "status", "apply_track", "rerank_score", "blended_score",
        "url", "created_at", "updated_at", "submitted_at", "notes"
    ])
    for app_row, job_row in results:
        writer.writerow([
            app_row.id,
            job_row.company,
            job_row.title,
            job_row.location,
            job_row.remote,
            job_row.source.value if job_row.source else "",
            app_row.status.value if app_row.status else "",
            app_row.apply_track or "",
            job_row.rerank_score,
            job_row.blended_score,
            job_row.url,
            app_row.created_at.isoformat() if app_row.created_at else "",
            app_row.updated_at.isoformat() if app_row.updated_at else "",
            app_row.submitted_at.isoformat() if app_row.submitted_at else "",
            (app_row.notes or "").replace("\n", " "),
        ])

    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=applications.csv"},
    )
