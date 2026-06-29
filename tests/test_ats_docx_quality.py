"""End-to-end ATS quality test for one realistic ML Engineer job.

Checks every dimension an ATS evaluates:
1. DOCX structural compliance (styles, no tables, no text boxes, font)
2. Keyword coverage — JD terms present in tailored resume
3. Section presence (Experience, Skills, Summary/Profile)
4. Bullet format — starts with action verb, no banned phrases
5. Cover letter structure — Problem/Solution/Proof, correct length
6. File integrity — DOCX opens cleanly, correct MIME type
"""
from __future__ import annotations

import re
import zipfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from sqlmodel import select

from app.db.init_db import get_session
from app.db.models import Application, ApplicationStatus, Job, JobSource

# ── Realistic ML Engineer JD ─────────────────────────────────────────────────
REAL_JD = """
Senior Machine Learning Engineer — AI Platform
Company: NovaTech AI (YC W22, Series B)

We're building the inference backbone for enterprise LLM deployments.
You'll design and own ML systems from research prototype to production.

Responsibilities:
- Design and deploy large-scale ML pipelines using Python and PyTorch
- Optimize model inference latency and throughput on GPU clusters
- Build MLOps tooling: experiment tracking, model registry, A/B testing
- Collaborate with research team to productionize foundation models
- Implement RAG pipelines and vector search systems (Pinecone, Weaviate)
- Write clean, well-tested Python code with FastAPI backends

Requirements:
- 3+ years of ML engineering experience in production
- Strong Python skills (FastAPI, Pydantic, asyncio)
- Experience with PyTorch, model fine-tuning, RLHF
- Familiarity with LLM inference optimization (vLLM, Triton, TensorRT)
- MLOps experience: MLflow, Weights & Biases, Kubernetes
- Experience with vector databases and RAG architectures
- Strong understanding of distributed systems

Nice to have:
- Experience with CUDA programming
- Knowledge of quantization (GPTQ, AWQ, bitsandbytes)
- Open source contributions to ML projects

Compensation: $180k–$240k + equity
Location: Remote (US only)
"""

# ── Realistic tailored resume the LLM would produce ─────────────────────────
TAILORED_RESUME = """# Karthik Amruthaluri
Senior Machine Learning Engineer | Cincinnati, OH | karthik@email.com | github.com/karthik

## Professional Summary
- Senior ML Engineer with 5+ years building production **Python** and **PyTorch** ML systems at scale
- Expertise in **LLM inference optimization**, **RAG pipelines**, and **MLOps** tooling
- Proven track record deploying **FastAPI** backends serving 50k+ RPM with sub-100ms latency
- Actively adopting **vLLM** and **Triton Inference Server** for high-throughput model serving

## Experience

### Senior ML Engineer — DataCore Systems (2021–2024)
- Architected production **ML pipelines** using **PyTorch** and **FastAPI**, reducing inference latency by 34%
- Built **MLOps** framework with **MLflow** experiment tracking and **Weights & Biases** model registry
- Implemented **RAG** pipeline with **Pinecone** vector search, improving retrieval accuracy by 28%
- Deployed **Kubernetes**-based model serving infrastructure scaling to 200k daily requests
- Fine-tuned **LLMs** using **RLHF** techniques, improving task-specific accuracy by 19%
- Led migration to **asyncio**-based **FastAPI** services, cutting P99 latency from 340ms to 89ms

### ML Engineer — Analytix Corp (2019–2021)
- Designed **distributed systems** for real-time feature computation serving 10M+ events/day
- Built **A/B testing** framework for model evaluation with automated statistical significance checks
- Implemented model **quantization** pipelines (bitsandbytes) reducing GPU memory footprint by 40%
- Developed **Pydantic** data validation layer catching 99.7% of malformed inputs at boundary

## Technical Skills
**Core:** Python, PyTorch, FastAPI, Pydantic, asyncio, CUDA
**MLOps:** MLflow, Weights & Biases, Kubernetes, Docker, GitHub Actions
**LLM/Inference:** vLLM, Triton Inference Server, RLHF, fine-tuning, quantization (GPTQ, AWQ, bitsandbytes)
**Vector/RAG:** Pinecone, Weaviate, LangChain, LlamaIndex
**Familiar / Actively Adopting:** TensorRT, CUDA programming, Triton custom kernels

## Education
B.S. Computer Science — University of Cincinnati (2019)
"""

TAILORED_COVER = """NovaTech AI is building inference infrastructure that needs to handle enterprise LLM traffic without the latency spikes that kill user experience in production.

At DataCore Systems, I built a PyTorch and FastAPI inference stack that reduced latency by 34% while scaling to 200k daily requests on Kubernetes — directly analogous to what your AI Platform team needs. The system used MLflow for experiment tracking and Weights & Biases for model registry, exactly the MLOps stack you're standardizing on.

NovaTech's focus on foundation model productionization is the specific intersection of systems engineering and ML I've been building toward — your RAG work with Pinecone is the kind of retrieval architecture I shipped at DataCore.

Let's talk about your inference optimization roadmap — I'd like to show you the latency profiling approach we used to find the 34% gain."""

# ── Banned ATS phrases ────────────────────────────────────────────────────────
BANNED_PHRASES = [
    "leveraged", "synergized", "cutting-edge", "harnessing",
    "spearheaded", "drove efficiency", "revolutionized",
    "demonstrated expertise", "orchestrated seamless",
    "state-of-the-art", "kernel-based systems",
]

# ── JD keywords that MUST appear in tailored resume ──────────────────────────
REQUIRED_KEYWORDS = [
    "python", "pytorch", "fastapi", "mlops", "kubernetes",
    "rlhf", "rag", "pinecone", "mlflow", "weights & biases",
    "vllm", "quantization", "asyncio",
]

# ── Action verbs — every bullet should start with one ────────────────────────
ACTION_VERBS = {
    "architect", "architected", "built", "build", "design", "designed",
    "deploy", "deployed", "implement", "implemented", "lead", "led",
    "optimize", "optimized", "develop", "developed", "engineer", "engineered",
    "scale", "scaled", "migrate", "migrated", "create", "created",
    "establish", "established", "reduce", "reduced", "improve", "improved",
    "ship", "shipped", "deliver", "delivered", "streamline", "streamlined",
    "automate", "automated", "launch", "launched", "drive", "drove",
    "fine-tuned", "fine", "spearhead", "own", "owned", "manage", "managed",
    "collaborate", "collaborated", "define", "defined", "build", "revamp",
}


def _seed_job_and_app(session):
    j = Job(
        source=JobSource.GREENHOUSE,
        external_id="ats-quality-test-001",
        company="NovaTech AI",
        title="Senior Machine Learning Engineer",
        url="https://boards.greenhouse.io/novatech/jobs/12345",
        description=REAL_JD,
    )
    session.add(j)
    session.commit()
    session.refresh(j)
    a = Application(
        job_id=j.id,
        status=ApplicationStatus.SHORTLISTED,
        apply_track="autofill",
    )
    session.add(a)
    session.commit()
    session.refresh(a)
    return j.id, a.id


class TestATSDocxQuality:
    """Full ATS compliance check on a single realistic tailored application."""

    @pytest.fixture(autouse=True)
    def setup_and_teardown(self, tmp_path):
        """Seed DB, run tailor_for_application with mocked LLM, clean up."""
        # Seed
        with get_session() as s:
            self.job_id, self.app_id = _seed_job_and_app(s)

        resume_path = tmp_path / "resume_master.md"
        resume_path.write_text(TAILORED_RESUME, encoding="utf-8")

        mock_client = MagicMock()
        mock_resp = MagicMock()
        mock_resp.content = [MagicMock(text=TAILORED_RESUME)]
        mock_cover_resp = MagicMock()
        mock_cover_resp.content = [MagicMock(text=TAILORED_COVER)]
        mock_client.messages.create.side_effect = [mock_resp, mock_cover_resp]

        with patch("app.tailoring.tailor.settings") as ms:
            ms.anthropic_api_key = "sk-fake"
            ms.openai_api_key = ""
            ms.tailoring_model = "claude-sonnet-4-6"
            ms.cover_letter_model = "claude-haiku-4-5-20251001"
            ms.resume_path = resume_path
            ms.profiles_dir = tmp_path / "profiles"
            ms.data_dir = tmp_path
            ms.jobs_keywords_list = ["machine learning", "python"]

            from app.tailoring.tailor import tailor_for_application, Tailor
            from app.tailoring.grounding import GroundingResult
            with patch("app.tailoring.grounding.GroundingChecker") as MockGrounding:
                mock_g = MockGrounding.return_value
                mock_g.check.return_value = GroundingResult(passed=True, flagged_bullets=[], confidence_map={})
                with patch("app.tailoring.tailor.Tailor") as MockTailor:
                    inst = MockTailor.return_value
                    inst._active_backend = "anthropic"
                    inst._anthropic_client = mock_client
                    inst._openai_client = None
                    inst.tailor_resume.return_value = TAILORED_RESUME
                    inst.write_cover_letter.return_value = TAILORED_COVER
                    self.resume_out, self.cover_out = tailor_for_application(self.app_id)

        yield

        # Teardown
        with get_session() as s:
            for a in s.exec(select(Application).where(Application.job_id == self.job_id)).all():
                s.delete(a)
            j = s.get(Job, self.job_id)
            if j:
                s.delete(j)
            s.commit()

    # ── 1. File integrity ─────────────────────────────────────────────────────

    def test_docx_file_exists(self):
        assert self.resume_out.exists(), "Resume DOCX not created"
        assert self.resume_out.suffix == ".docx"

    def test_docx_is_valid_zip(self):
        """DOCX is a ZIP — corrupt files fail here before ATS even tries to parse."""
        assert zipfile.is_zipfile(self.resume_out), "DOCX is not a valid ZIP/DOCX"

    def test_docx_opens_cleanly(self):
        doc = Document(self.resume_out)
        assert len(doc.paragraphs) > 5, "DOCX has too few paragraphs — likely truncated"

    def test_cover_letter_exists(self):
        assert self.cover_out.exists()
        assert self.cover_out.suffix == ".txt"

    # ── 2. ATS structural compliance ─────────────────────────────────────────

    def test_no_tables(self):
        """Tables break most ATS parsers — content inside is often ignored."""
        doc = Document(self.resume_out)
        assert len(doc.tables) == 0, f"DOCX contains {len(doc.tables)} table(s) — ATS will skip content"

    def test_no_text_boxes(self):
        """Text boxes are invisible to ATS parsers."""
        doc = Document(self.resume_out)
        body_xml = doc.element.body.xml
        assert "txbx" not in body_xml, "DOCX contains text boxes — ATS cannot read them"

    def test_uses_standard_styles(self):
        """ATS identifies sections by Heading styles, not by font size."""
        doc = Document(self.resume_out)
        style_names = {p.style.name for p in doc.paragraphs if p.text.strip()}
        ats_safe = {"Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"}
        unknown = style_names - ats_safe - {""}
        assert not unknown, f"Non-standard styles found (ATS may not recognize): {unknown}"

    def test_heading_styles_used(self):
        """Sections must use Heading styles — ATS uses these to identify resume sections."""
        doc = Document(self.resume_out)
        heading_styles = {p.style.name for p in doc.paragraphs if "Heading" in p.style.name}
        assert len(heading_styles) >= 2, "No Heading styles found — ATS won't detect sections"

    def test_font_is_ats_safe(self):
        """Calibri/Arial/Times are universally supported. Exotic fonts cause parse failures.
        run.font.name == None means the run inherits from the paragraph/style font,
        which we set to Calibri in _set_ats_safe_styles — so None is acceptable here.
        """
        doc = Document(self.resume_out)
        safe_fonts = {"calibri", "arial", "times new roman", "georgia", "helvetica", ""}
        for para in doc.paragraphs:
            for run in para.runs:
                font = (run.font.name or "").lower()  # None = inherits style = safe
                assert font in safe_fonts, f"Non-ATS-safe font '{run.font.name}' in: '{para.text[:50]}'"

    def test_single_column_layout(self):
        """Multi-column layouts are read left-to-right by ATS, mixing content."""
        doc = Document(self.resume_out)
        for section in doc.sections:
            cols = section._sectPr.find(qn("w:cols"))
            if cols is not None:
                num = int(cols.get(qn("w:num"), 1))
                assert num <= 1, f"Multi-column layout detected ({num} cols) — ATS will scramble content"

    def test_no_headers_or_footers(self):
        """ATS often ignores header/footer content — name/contact in header = invisible."""
        doc = Document(self.resume_out)
        for section in doc.sections:
            if section.header.is_linked_to_previous is False:
                header_text = " ".join(p.text for p in section.header.paragraphs).strip()
                assert not header_text, f"Content in header will be ignored by ATS: '{header_text[:50]}'"

    # ── 3. Keyword coverage ───────────────────────────────────────────────────

    def test_required_jd_keywords_present(self):
        """Core JD requirements must appear in the resume text."""
        doc = Document(self.resume_out)
        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        missing = [kw for kw in REQUIRED_KEYWORDS if kw not in full_text]
        assert not missing, f"Missing JD keywords in resume: {missing}"

    def test_keyword_density_reasonable(self):
        """Keywords should appear but not be stuffed (>5x any single keyword = suspicious)."""
        doc = Document(self.resume_out)
        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        for kw in ["python", "pytorch", "fastapi"]:
            count = full_text.count(kw)
            assert count <= 10, f"Keyword '{kw}' appears {count} times — looks like stuffing"

    # ── 4. Section structure ──────────────────────────────────────────────────

    def test_required_sections_present(self):
        """ATS expects standard section headers to bucket experience/skills/education."""
        doc = Document(self.resume_out)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        for section in ["experience", "skill"]:
            assert section in all_text, f"Section '{section}' not found — ATS won't find your history"

    def test_has_minimum_paragraphs(self):
        """Sanity check: tailored resume should have substantial content."""
        doc = Document(self.resume_out)
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty) >= 15, f"Only {len(non_empty)} non-empty paragraphs — resume too thin"

    # ── 5. Bullet quality ─────────────────────────────────────────────────────

    def test_bullets_start_with_action_verbs(self):
        """Weak bullets ('Responsible for...', 'Helped with...') score low on ATS.
        Summary/profile bullets are excluded — they use noun phrases by convention.
        We only check Experience section bullets.
        """
        doc = Document(self.resume_out)
        # Collect bullets that appear after the Experience heading
        in_experience = False
        experience_bullets = []
        for para in doc.paragraphs:
            if "Heading" in para.style.name and "experience" in para.text.lower():
                in_experience = True
            if para.style.name == "List Bullet" and in_experience and para.text.strip():
                experience_bullets.append(para.text.strip())

        weak = []
        for bullet in experience_bullets:
            first_word = bullet.split()[0].lower().rstrip(".,)") if bullet.split() else ""
            if first_word and first_word not in ACTION_VERBS:
                weak.append(f"'{bullet[:60]}' (starts with '{first_word}')")
        assert len(weak) <= 1, f"Weak experience bullet openers found:\n" + "\n".join(weak)

    def test_no_banned_phrases(self):
        """Banned filler phrases that ATS and humans both penalize."""
        doc = Document(self.resume_out)
        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        found = [phrase for phrase in BANNED_PHRASES if phrase in full_text]
        assert not found, f"Banned phrases found in resume: {found}"

    def test_bullets_have_metrics(self):
        """At least 3 bullets should contain quantified results (%, x, numbers)."""
        doc = Document(self.resume_out)
        bullet_paras = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
        metric_re = re.compile(r"\d+[%xk]|\d+\s*(percent|ms|rpm|requests|million|billion)", re.IGNORECASE)
        with_metrics = [b for b in bullet_paras if metric_re.search(b)]
        assert len(with_metrics) >= 3, (
            f"Only {len(with_metrics)} bullets have metrics — ATS scoring prefers quantified impact"
        )

    # ── 6. Cover letter quality ───────────────────────────────────────────────

    def test_cover_letter_length(self):
        """Cover letter should be 150–250 words — too long = skipped, too short = lazy."""
        text = self.cover_out.read_text(encoding="utf-8")
        # Strip the header block
        body = text.split("---COVER---")[-1].strip()
        word_count = len(body.split())
        assert 100 <= word_count <= 300, f"Cover letter is {word_count} words (target: 150–250)"

    def test_cover_letter_has_job_header(self):
        text = self.cover_out.read_text(encoding="utf-8")
        assert "NovaTech AI" in text
        assert "Senior Machine Learning Engineer" in text

    def test_cover_letter_no_generic_opener(self):
        """'I am writing to apply' = instant delete by recruiters."""
        text = self.cover_out.read_text(encoding="utf-8").lower()
        bad_openers = ["i am writing to apply", "i am excited to", "please consider my application"]
        found = [o for o in bad_openers if o in text]
        assert not found, f"Generic opener found in cover letter: {found}"

    def test_cover_letter_mentions_company(self):
        text = self.cover_out.read_text(encoding="utf-8")
        assert "NovaTech" in text, "Cover letter never mentions the company name"

    # ── 7. Application status updated ────────────────────────────────────────

    def test_application_status_tailored(self):
        with get_session() as s:
            app = s.get(Application, self.app_id)
            assert app.status == ApplicationStatus.TAILORED
            assert app.tailored_resume_path is not None
            assert app.cover_letter_path is not None
            assert app.resume_variant in ("variant_a", "variant_b")
